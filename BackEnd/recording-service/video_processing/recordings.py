import os
import uuid
import shutil
import subprocess
import json
import re
import asyncio
from urllib.parse import quote_plus
import time
from datetime import datetime, timedelta
from tempfile import TemporaryDirectory
from typing import List
from django.db import connection
from pymongo import MongoClient
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.utils.decorators import method_decorator
from django.views import View
# Graphviz - wrap in try/except
try:
    from graphviz import Source
except ImportError:
    Source = None
    print("Warning: graphviz not available")

# =============================================================================
# LAZY LOADING FOR HEAVY IMPORTS - FIX FOR SLOW STARTUP
# =============================================================================
_torch = None
_MarianMTModel = None
_MarianTokenizer = None

def get_torch():
    """Lazy load PyTorch only when needed"""
    global _torch
    if _torch is None:
        try:
            import torch
            _torch = torch
            print(f"PyTorch loaded. GPU available: {torch.cuda.is_available()}")
        except ImportError:
            _torch = False
            print("Warning: torch not available - GPU features disabled")
    return _torch if _torch else None

def get_transformers():
    """Lazy load transformers only when needed"""
    global _MarianMTModel, _MarianTokenizer
    if _MarianMTModel is None:
        try:
            from transformers import MarianMTModel, MarianTokenizer
            _MarianMTModel = MarianMTModel
            _MarianTokenizer = MarianTokenizer
            print("Transformers loaded successfully")
        except ImportError:
            _MarianMTModel = False
            _MarianTokenizer = False
            print("Warning: transformers not available")
    return _MarianMTModel if _MarianMTModel else None, _MarianTokenizer if _MarianTokenizer else None

import logging
from urllib.parse import quote_plus
from django.http import StreamingHttpResponse
import openai
from asgiref.sync import sync_to_async
from django.http import StreamingHttpResponse, FileResponse
import requests
from io import BytesIO
from wsgiref.util import FileWrapper
import mimetypes
from bson import ObjectId
import boto3
from botocore.exceptions import NoCredentialsError
from pydub import AudioSegment
from deep_translator import GoogleTranslator
from django.utils import timezone
from openai import OpenAI
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from clients.notification_client import (
    ensure_notification_tables,
    _get_recording_meeting_info,
    _get_recording_participants,
    short_id
)

import pytz  # For timezone handling in notifications

# Status codes (previously imported from meetings.py)
BAD_REQUEST_STATUS = 400
NOT_FOUND_STATUS = 404
SERVER_ERROR_STATUS = 500
SUCCESS_STATUS = 200
TBL_MEETINGS = 'tbl_Meetings'

def create_meetings_table():
    """Stub — meetings table managed by Meeting Core Service"""
    pass

# Add this configuration after imports
RECORDING_SERVICE_URL = os.getenv(
    "RECORDING_SERVICE_URL", 
    "http://recording-service.imeetpro.svc.cluster.local:8080"
)

# === SSL CERTIFICATE FIX FOR HUGGINGFACE & GOOGLE TRANSLATE ===
import ssl
import certifi
import os
import urllib3

# Disable SSL warnings (for development/testing only)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Fix SSL certificate verification
os.environ['REQUESTS_CA_BUNDLE'] = certifi.where()
os.environ['SSL_CERT_FILE'] = certifi.where()
os.environ['CURL_CA_BUNDLE'] = certifi.where()

# Additional fix for transformers library
try:
    import requests
    requests.packages.urllib3.util.ssl_.DEFAULT_CIPHERS = 'ALL:@SECLEVEL=1'
    
    # Create a custom session with SSL verification disabled (for problematic environments)
    old_request = requests.Session.request
    def new_request(self, method, url, **kwargs):
        kwargs.setdefault('verify', False)
        return old_request(self, method, url, **kwargs)
    requests.Session.request = new_request
except Exception as e:
    print(f"SSL patch warning: {e}")

# Fix for httpx (used by some libraries)
try:
    import httpx
    httpx._config.DEFAULT_TIMEOUT_CONFIG = httpx.Timeout(30.0)
except:
    pass

# === CONFIGURATION ===
# AWS Configuration
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION = os.getenv("AWS_REGION", "ap-south-1")
AWS_S3_BUCKET = os.getenv("AWS_S3_BUCKET", "imeetpro-prod-recordings")

# S3 Client — fork-safe lazy initialization
# 
# WHY THIS IS NEEDED:
# Celery uses prefork workers. boto3's underlying urllib3 connection pool is
# NOT fork-safe. If we create s3_client at module import time, child processes
# inherit broken file descriptors and download_file() hangs silently forever.
# This proxy creates a fresh client per-process on first use.
from botocore.config import Config
import threading

_s3_client_local = threading.local()

def _build_s3_client():
    return boto3.client(
        "s3",
        aws_access_key_id=AWS_ACCESS_KEY_ID,
        aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
        region_name=AWS_REGION,
        config=Config(
            connect_timeout=15,
            read_timeout=300,
            retries={'max_attempts': 3, 'mode': 'standard'}
        )
    )

def get_s3_client():
    """Return a boto3 S3 client unique to this process. Recreates after fork()."""
    pid = os.getpid()
    if not hasattr(_s3_client_local, 'client') or getattr(_s3_client_local, 'pid', None) != pid:
        _s3_client_local.client = _build_s3_client()
        _s3_client_local.pid = pid
    return _s3_client_local.client

class _S3ClientProxy:
    """Proxy so existing code like `s3_client.download_file(...)` keeps working."""
    def __getattr__(self, name):
        return getattr(get_s3_client(), name)

s3_client = _S3ClientProxy()

S3_FOLDERS = {
    "videos": os.getenv("S3_FOLDER_VIDEOS", "videos"),
    "transcripts": os.getenv("S3_FOLDER_TRANSCRIPTS", "transcripts"),
    "summary": os.getenv("S3_FOLDER_SUMMARY", "summary"),
    "images": os.getenv("S3_FOLDER_IMAGES", "summary_image"),
    "subtitles": os.getenv("S3_FOLDER_SUBTITLES", "subtitles")
}

    
# ============================================
# OPENAI CLIENT INITIALIZATION - for transcription
# ============================================
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "").strip()
client = None
if OPENAI_API_KEY:
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
        print(f"✅ OpenAI client initialized successfully (key: {OPENAI_API_KEY[:8]}...)")
    except Exception as e:
        client = None
        print(f"❌ OpenAI client initialization FAILED: {e}")
else:
    print(f"❌ OPENAI_API_KEY is EMPTY or NOT SET in environment")
    print(f"❌ Transcription will NOT work!")


# ============================================
# GROQ CLIENT INITIALIZATION - for summary, translation, trainer eval, subtitles
# ============================================
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
groq_client = None

if GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        print(f"✅ Groq client initialized successfully (key: {GROQ_API_KEY[:8]}...)")
    except Exception as e:
        groq_client = None
        print(f"❌ Groq client initialization FAILED: {e}")
else:
    print(f"❌ GROQ_API_KEY is EMPTY or NOT SET in environment")
    print(f"❌ Summary, translation, and trainer evaluation will NOT work!")
    

openai.api_key = os.getenv("OPENAI_API_KEY")

# === FIXED FUNCTION 1: Enhanced S3 Key Extraction ===
def extract_s3_key_from_url(s3_url: str, s3_bucket: str) -> str:
    """
    Extract S3 key from URL with better handling for nested paths and special characters.
    Fixes issue with folder names like: 4b77cc81-9064-434c-8e79-0f9fab31cbdf_Updates
    """
    try:
        if not s3_url:
            return None
        
        logger.debug(f"Extracting S3 key from URL: {s3_url}")
        
        # Remove protocol if present
        if s3_url.startswith('s3://'):
            key = s3_url.replace(f's3://{s3_bucket}/', '')
            logger.debug(f"S3 protocol: Extracted key = {key}")
            return key
        
        # HTTP(S) URL handling - MAIN STRATEGY
        if 'http' in s3_url and '.amazonaws.com/' in s3_url:
            # Split on .amazonaws.com/ - everything after is the S3 key
            # This cleanly separates domain from path regardless of bucket name or special chars
            key = s3_url.split('.amazonaws.com/')[-1]
            # Remove query parameters if any
            key = key.split('?')[0]
            if key:
                logger.debug(f"✅ Extracted S3 key: {key}")
                return key
        
        # Fallback for non-standard URLs
        if 'http' in s3_url and s3_bucket in s3_url:
            # Find first occurrence of bucket
            idx = s3_url.find(s3_bucket)
            after_bucket = s3_url[idx + len(s3_bucket):]
            
            # Look for the path part (after domain suffixes)
            if '/' in after_bucket:
                # Take everything after first slash
                key = after_bucket.split('/', 1)[1]
                key = key.split('?')[0]
                if key:
                    logger.debug(f"✅ Extracted S3 key (fallback): {key}")
                    return key
        
        # If it looks like already a key (no URL patterns)
        if not s3_url.startswith('http') and '://' not in s3_url:
            logger.debug(f"Direct key format: {s3_url}")
            return s3_url
        
        logger.warning(f"⚠️ Could not extract S3 key from URL: {s3_url}")
        return None
        
    except Exception as e:
        logger.error(f"❌ Exception in extract_s3_key_from_url: {e}")
        logger.error(f"URL was: {s3_url}")
        return None

# === FIXED FUNCTION 2: Robust S3 Key Builder (DOCX) ===
def build_s3_key_from_parts(
    base_folder: str,
    meeting_id: str,
    user_id: str,
    meeting_type: str = None,
    doc_type: str = None,
    lang: str = None
) -> str:
    """
    Build complete S3 key path with proper validation (DOCX for documents).
    """
    try:
        # Sanitize inputs
        meeting_id = str(meeting_id).strip() if meeting_id else ""
        user_id = str(user_id).strip() if user_id else ""
        base_folder = str(base_folder).strip() if base_folder else "videos"

        if not meeting_id or not user_id:
            logger.error(f"Invalid input: meeting_id={meeting_id}, user_id={user_id}")
            return None

        # ================= Schedule Meeting =================
        if meeting_type and meeting_type.lower() == "schedulemeeting":
            try:
                schedule_meta = get_schedule_meeting_metadata(meeting_id)

                if schedule_meta.get("is_scheduled"):
                    folder_path = schedule_meta.get("folder_path", "")

                    if doc_type == "transcript":
                        return f"transcripts/{folder_path}/{meeting_id}_{user_id}_transcript.docx"

                    elif doc_type == "summary":
                        return f"summary/{folder_path}/{meeting_id}_{user_id}_summary.docx"

                    elif doc_type == "subtitles":
                        if lang:
                            return f"subtitles/{folder_path}/{meeting_id}_{user_id}_{lang}.srt"
                        return f"subtitles/{folder_path}"

                    # Video
                    return f"{base_folder}/{folder_path}/{meeting_id}_{user_id}_recording.mp4"

            except Exception as schedule_error:
                logger.warning(
                    f"Failed to get schedule metadata: {schedule_error}, falling back to default"
                )

        # ================= Default / Fallback =================
        if doc_type == "transcript":
            return f"transcripts/{meeting_id}_{user_id}_transcript.docx"

        elif doc_type == "summary":
            return f"summary/{meeting_id}_{user_id}_summary.docx"

        elif doc_type == "subtitles":
            if lang:
                return f"subtitles/{meeting_id}_{user_id}_{lang}.srt"
            return f"subtitles/{meeting_id}_{user_id}"

        # Video
        return f"{base_folder}/{meeting_id}_{user_id}_recording.mp4"

    except Exception as e:
        logger.error(f"Error building S3 key: {e}")
        return None


# === FIXED FUNCTION 3: Verify and Repair Video URL ===
def verify_and_repair_video_url(video_doc: dict) -> dict:
    """
    Verify video URL matches actual S3 content.
    If mismatch found, attempt to rebuild the correct path.
    
    Returns updated video document with corrected URL if needed.
    """
    try:
        meeting_id = video_doc.get("meeting_id")
        user_id = video_doc.get("user_id")
        stored_url = video_doc.get("video_url")
        meeting_type = video_doc.get("meeting_type", "InstantMeeting")
        
        if not all([meeting_id, user_id, stored_url]):
            logger.warning(f"Incomplete video metadata: meeting_id={meeting_id}, user_id={user_id}")
            return video_doc
        
        # Check if stored URL exists in S3
        try:
            s3_key = extract_s3_key_from_url(stored_url, AWS_S3_BUCKET)
            if s3_key:
                size = get_s3_object_size(s3_key)
                if size > 0:
                    logger.info(f"✅ Video URL verified: {s3_key} ({size} bytes)")
                    return video_doc  # URL is valid
                else:
                    logger.warning(f"⚠️ Video file not found at S3 key: {s3_key}")
        except Exception as check_error:
            logger.warning(f"⚠️ Failed to verify URL: {check_error}")
        
        # URL is invalid, attempt to rebuild correct path
        logger.info(f"🔧 Attempting to rebuild S3 path for meeting_id={meeting_id}, user_id={user_id}, type={meeting_type}")
        
        # CRITICAL FIX: Try file_path from MongoDB first (this is the merged file with audio)
        file_path_key = video_doc.get("file_path")
        if file_path_key:
            try:
                size = get_s3_object_size(file_path_key)
                if size > 0:
                    rebuilt_url = f"https://{AWS_S3_BUCKET}.s3.{AWS_REGION}.amazonaws.com/{file_path_key}"
                    logger.info(f"✅ Found video at file_path: {file_path_key} ({size} bytes)")
                    try:
                        collection.update_one(
                            {"_id": video_doc.get("_id")},
                            {"$set": {"video_url": rebuilt_url}}
                        )
                        logger.info(f"✅ Updated video URL from file_path in MongoDB")
                    except Exception as update_error:
                        logger.warning(f"⚠️ Failed to update MongoDB: {update_error}")
                    video_doc["video_url"] = rebuilt_url
                    video_doc["url_repaired"] = True
                    return video_doc
            except Exception as fp_error:
                logger.warning(f"⚠️ file_path check failed: {fp_error}")
        
        rebuilt_key = build_s3_key_from_parts(
            base_folder="videos",
            meeting_id=meeting_id,
            user_id=user_id,
            meeting_type=meeting_type
        )
        
        if not rebuilt_key:
            logger.error(f"Failed to rebuild S3 key")
            return video_doc
        
        # Check if rebuilt key exists
        try:
            size = get_s3_object_size(rebuilt_key)
            if size > 0:
                # Rebuilt path is valid! Update the document
                rebuilt_url = f"https://{AWS_S3_BUCKET}.s3.{AWS_REGION}.amazonaws.com/{rebuilt_key}"
                logger.info(f"✅ Found video at rebuilt path: {rebuilt_key}")
                
                # Update MongoDB
                try:
                    collection.update_one(
                        {"_id": video_doc.get("_id")},
                        {"$set": {"video_url": rebuilt_url}}
                    )
                    logger.info(f"✅ Updated video URL in MongoDB")
                except Exception as update_error:
                    logger.warning(f"⚠️ Failed to update MongoDB: {update_error}")
                
                video_doc["video_url"] = rebuilt_url
                video_doc["url_repaired"] = True
                return video_doc
            else:
                logger.warning(f"⚠️ Rebuilt path also not found: {rebuilt_key}")
                
        except Exception as rebuild_check_error:
            logger.warning(f"⚠️ Failed to check rebuilt path: {rebuild_check_error}")
        
        # Could not find video at any location
        logger.error(f"❌ Cannot locate video file for meeting_id={meeting_id}, user_id={user_id}")
        video_doc["url_error"] = "Video file not found in S3"
        return video_doc
        
    except Exception as e:
        logger.error(f"Error in verify_and_repair_video_url: {e}")
        return video_doc

# MongoDB Configuration - Updated to match FastAPI
mongo_user = os.getenv("MONGO_USER", "connectly")
mongo_password = os.getenv("MONGO_PASSWORD", "password")
mongo_host = os.getenv("MONGO_HOST", "mongodb.databases.svc.cluster.local")
mongo_port = os.getenv("MONGO_PORT", "27017")
mongo_db = os.getenv("MONGO_DB", "connectlydb")

MONGO_URI = os.getenv("MONGO_URI")
mongo_client = MongoClient(MONGO_URI)
db = mongo_client[mongo_db]
collection = db["test"]
TRASH_RETENTION_DAYS = 15
# === LOGGING SETUP ===
logger = logging.getLogger("video_processor")
logging.basicConfig(level=logging.INFO)

# === UTILITY FUNCTIONS ===
def upload_to_aws_s3(local_file_path: str, s3_key: str, timeout_sec: int = 1800) -> str:
    """
    Upload file to AWS S3 and return the URL.

    Wraps boto3.upload_file() in a worker thread with a hard 30-minute
    wall-clock timeout to prevent silent hangs (known boto3 fork-safety
    issue in Celery prefork workers when uploading files >400 MB).
    """
    import concurrent.futures

    def _do_upload():
        try:
            s3_client.upload_file(local_file_path, AWS_S3_BUCKET, s3_key)
            logger.info(f"File uploaded to s3://{AWS_S3_BUCKET}/{s3_key}")
            return f"https://{AWS_S3_BUCKET}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
        except NoCredentialsError:
            logger.error("AWS credentials not available.")
            return None
        except Exception as e:
            logger.error(f"S3 upload failed: {e}")
            return None

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(_do_upload)
            return future.result(timeout=timeout_sec)
    except concurrent.futures.TimeoutError:
        logger.error(f"❌ S3 upload HUNG after {timeout_sec}s for key: {s3_key}")
        return None

def download_from_s3(s3_key: str, local_file_path: str, timeout_sec: int = 1800) -> bool:
    """
    Download file from S3 using a single streaming connection.

    Why streaming get_object() instead of download_file():
    - download_file() uses a TransferManager with a 10-thread pool for
      multipart parallel download. In Celery prefork workers, those threads
      inherit a corrupted connection pool from the parent process. For files
      above ~400 MB the probability of every thread picking a broken connection
      approaches 1, causing a permanent silent hang.
    - get_object() uses ONE connection, ONE thread, no transfer manager.
      Works reliably for files of any size, including 1 GB+.

    Safety nets:
    - 30-minute wall-clock timeout (any hang fails loudly instead of forever)
    - Size verification (catches truncated downloads)
    - Partial file cleanup on failure
    - Progress logging at 25/50/75/100%
    """
    import concurrent.futures

    def _do_download():
        try:
            response = s3_client.get_object(Bucket=AWS_S3_BUCKET, Key=s3_key)
            content_length = response.get('ContentLength', 0)
            body = response['Body']

            chunk_size = 8 * 1024 * 1024  # 8 MB
            total_bytes = 0
            next_log_pct = 25  # Log at 25%, 50%, 75%, 100%

            with open(local_file_path, 'wb') as f:
                while True:
                    chunk = body.read(chunk_size)
                    if not chunk:
                        break
                    f.write(chunk)
                    total_bytes += len(chunk)

                    if content_length > 0:
                        pct = int((total_bytes / content_length) * 100)
                        if pct >= next_log_pct:
                            logger.info(
                                f"📥 Download progress: {pct}% "
                                f"({total_bytes / 1e6:.1f} / {content_length / 1e6:.1f} MB)"
                            )
                            next_log_pct = pct + 25

            try:
                body.close()
            except Exception:
                pass

            # Verify the download is complete
            actual_size = os.path.getsize(local_file_path)
            if content_length > 0 and actual_size != content_length:
                logger.error(
                    f"❌ Size mismatch: expected {content_length} bytes, "
                    f"got {actual_size} bytes"
                )
                return False

            logger.info(f"✅ Downloaded {actual_size / 1e6:.1f} MB to {local_file_path}")
            return True

        except Exception as e:
            logger.error(f"S3 streaming download failed: {e}")
            return False

    # Wall-clock timeout protection (catches any hang at the connection level)
    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(_do_download)
            return future.result(timeout=timeout_sec)
    except concurrent.futures.TimeoutError:
        logger.error(
            f"❌ S3 download HUNG after {timeout_sec}s for key: {s3_key} "
            f"— failing task to allow retry."
        )
        # Clean up any partial file
        try:
            if os.path.exists(local_file_path):
                os.remove(local_file_path)
                logger.info(f"🧹 Cleaned up partial download: {local_file_path}")
        except Exception:
            pass
        return False

def delete_from_s3(s3_key: str) -> bool:
    """Delete file from S3."""
    try:
        s3_client.delete_object(Bucket=AWS_S3_BUCKET, Key=s3_key)
        logger.info(f"Deleted from S3: {s3_key}")
        return True
    except Exception as e:
        logger.error(f"S3 delete failed: {e}")
        return False

def get_s3_object_size(s3_key: str) -> int:
    """Get object size from S3 - optimized version."""
    try:
        response = s3_client.head_object(Bucket=AWS_S3_BUCKET, Key=s3_key)
        return response['ContentLength']
    except Exception as e:
        logger.error(f"Failed to get S3 object size for {s3_key}: {e}")
        return 0

def stream_from_s3(s3_key: str, start: int = None, end: int = None) -> bytes:
    """Stream content from S3 - optimized version."""
    try:
        if start is not None and end is not None:
            range_header = f'bytes={start}-{end}'
            response = s3_client.get_object(Bucket=AWS_S3_BUCKET, Key=s3_key, Range=range_header)
        else:
            response = s3_client.get_object(Bucket=AWS_S3_BUCKET, Key=s3_key)
        
        return response['Body'].read()
    except Exception as e:
        logger.error(f"S3 streaming failed for {s3_key}: {e}")
        return None

def get_meeting_type(meeting_id: str) -> str:
    """Get meeting type from database - checks all meeting tables"""
    try:
        with connection.cursor() as cursor:
            # First check tbl_Meetings for Meeting_Type
            cursor.execute("SELECT Meeting_Type FROM tbl_Meetings WHERE ID = %s", [meeting_id])
            row = cursor.fetchone()
            
            if row and row[0]:
                meeting_type = row[0]
                logger.info(f"Meeting {meeting_id} type: {meeting_type}")
                return meeting_type
            
            # Fallback: Check specific tables if Meeting_Type is NULL
            # Check CalendarMeetings
            cursor.execute("SELECT id FROM tbl_CalendarMeetings WHERE id = %s", [meeting_id])
            if cursor.fetchone():
                logger.info(f"Meeting {meeting_id} found in CalendarMeetings")
                return "CalendarMeeting"
            
            # Check ScheduledMeetings
            cursor.execute("SELECT id FROM tbl_ScheduledMeetings WHERE id = %s", [meeting_id])
            if cursor.fetchone():
                logger.info(f"Meeting {meeting_id} found in ScheduledMeetings")
                return "ScheduleMeeting"
            
            # Default to InstantMeeting if found in tbl_Meetings but not in specific tables
            cursor.execute("SELECT ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
            if cursor.fetchone():
                logger.info(f"Meeting {meeting_id} defaulting to InstantMeeting")
                return "InstantMeeting"
            
            logger.warning(f"Meeting {meeting_id} not found in any table")
            return "InstantMeeting"  # Default fallback
            
    except Exception as e:
        logger.error(f"Error getting meeting type for {meeting_id}: {e}")
        return "InstantMeeting"  # Safe default

def get_schedule_meeting_metadata(meeting_id: str) -> dict:
    """
    Get schedule meeting metadata for organizing recordings into folders.
    Returns schedule_id, schedule_title, and folder path.
    
    This function retrieves metadata about a scheduled meeting so recordings
    can be organized in S3 folders by schedule.
    """
    try:
        with connection.cursor() as cursor:
            # Get the schedule info from tbl_ScheduledMeetings
            cursor.execute("""
                SELECT id, title 
                FROM tbl_ScheduledMeetings 
                WHERE id = %s
            """, [meeting_id])
            
            row = cursor.fetchone()
            if row:
                schedule_id, title = row[0], row[1]
                logger.info(f"✅ Found schedule meeting metadata: schedule_id={schedule_id}, title={title}")
                
                # Sanitize title for folder name (remove special chars, replace spaces with underscores)
                sanitized_title = re.sub(r'[^a-zA-Z0-9\s_-]', '', title)
                sanitized_title = re.sub(r'\s+', '_', sanitized_title.strip())
                
                return {
                    "is_scheduled": True,
                    "schedule_id": schedule_id,
                    "schedule_title": title,
                    "sanitized_folder_name": f"{schedule_id}_{sanitized_title}",
                    "folder_path": f"schedule_meetings/{schedule_id}_{sanitized_title}"
                }
            else:
                logger.warning(f"Meeting {meeting_id} not found in tbl_ScheduledMeetings")
                return {
                    "is_scheduled": False,
                    "schedule_id": None,
                    "schedule_title": None,
                    "sanitized_folder_name": None,
                    "folder_path": None
                }
                
    except Exception as e:
        logger.error(f"Error getting schedule meeting metadata for {meeting_id}: {e}")
        return {
            "is_scheduled": False,
            "schedule_id": None,
            "schedule_title": None,
            "sanitized_folder_name": None,
            "folder_path": None
        }

def build_s3_video_path(meeting_id: str, user_id: str, meeting_type: str, session_id: str = None) -> str:
    """
    Build S3 path for video based on meeting type.
    
    For ScheduleMeetings: 
        videos/schedule_meetings/{schedule_id}_{sanitized_title}/{meeting_id}_{user_id}_recording.mp4
    
    For InstantMeetings/CalendarMeetings with session_id: 
        videos/{meeting_id}/{session_id}/{meeting_id}_{session_id}_{timestamp}.mp4
    
    For InstantMeetings/CalendarMeetings without session_id: 
        videos/{meeting_id}_{user_id}_recording.mp4
    
    This organizes scheduled meeting recordings into folders automatically.
    """
    try:
        base_folder = S3_FOLDERS['videos']
        
        if meeting_type == "ScheduleMeeting":
            # Get schedule metadata and create folder structure
            schedule_meta = get_schedule_meeting_metadata(meeting_id)
            
            if schedule_meta.get("is_scheduled"):
                # Path: videos/schedule_meetings/schedule_id_title/meeting_id_user_id_recording.mp4
                s3_key = f"{base_folder}/{schedule_meta['folder_path']}/{meeting_id}_{user_id}_recording.mp4"
                logger.info(f"✅ ScheduleMeeting S3 path: {s3_key}")
                return s3_key
        
        # ✅ NEW: If session_id provided, use session-based structure
        if session_id:
            import time
            timestamp = int(time.time())
            s3_key = f"{base_folder}/{meeting_id}/{session_id}/{meeting_id}_{session_id}_{timestamp}.mp4"
            logger.info(f"✅ Session-based S3 path: {s3_key}")
            return s3_key
        
        # Default path for InstantMeeting, CalendarMeeting, or if schedule fetch fails
        s3_key = f"{base_folder}/{meeting_id}_{user_id}_recording.mp4"
        logger.info(f"Default S3 path: {s3_key}")
        return s3_key
        
    except Exception as e:
        logger.error(f"Error building S3 path: {e}")
        return f"{S3_FOLDERS['videos']}/{meeting_id}_{user_id}_recording.mp4"
      
def build_s3_document_path(
    meeting_id: str,
    user_id: str,
    meeting_type: str,
    doc_type: str,
    session_id: str = None
) -> str:
    """
    Build S3 path for transcript/summary based on meeting type.

    For ScheduleMeetings: organizes documents in schedule folder
    For others: flat structure

    doc_type can be: 'transcript', 'summary', or 'subtitles'
    """
    try:
        # Default meeting type
        if not meeting_type:
            meeting_type = "InstantMeeting"

        logger.debug(
            f"Building S3 path: type={doc_type}, meeting_type={meeting_type}"
        )

        # ================= Schedule Meeting =================
        if meeting_type == "ScheduleMeeting":
            try:
                schedule_meta = get_schedule_meeting_metadata(meeting_id)

                if schedule_meta and schedule_meta.get("is_scheduled"):
                    folder_prefix = schedule_meta.get("folder_path", "")

                    if doc_type == "transcript":
                        base_folder = S3_FOLDERS.get("transcripts", "transcripts")
                        return (
                            f"{base_folder}/{folder_prefix}/"
                            f"{meeting_id}_{user_id}_transcript.docx"
                        )

                    elif doc_type == "summary":
                        base_folder = S3_FOLDERS.get("summary", "summary")
                        return (
                            f"{base_folder}/{folder_prefix}/"
                            f"{meeting_id}_{user_id}_summary.docx"
                        )

                    elif doc_type == "subtitles":
                        base_folder = S3_FOLDERS.get("subtitles", "subtitles")
                        return f"{base_folder}/{folder_prefix}"  # lang appended later

                    elif doc_type == "trainer_evaluation":
                        return (
                            f"trainer_evaluations/{folder_prefix}/"
                            f"{meeting_id}_{user_id}_evaluation.docx"
                        )

                else:
                    logger.warning(
                        f"Schedule metadata not found or not scheduled for {meeting_id}"
                    )

            except Exception as schedule_error:
                logger.warning(
                    f"Error fetching schedule metadata: {schedule_error}, "
                    f"using default path"
                )

        # ================= Default / Fallback =================
        if doc_type == "transcript":
            if session_id:
                return (
                    f"{S3_FOLDERS.get('transcripts', 'transcripts')}/"
                    f"{meeting_id}/{session_id}/"
                    f"{meeting_id}_{session_id}_transcript.docx"
                )
            return (
                f"{S3_FOLDERS.get('transcripts', 'transcripts')}/"
                f"{meeting_id}_{user_id}_transcript.docx"
            )

        elif doc_type == "summary":
            if session_id:
                return (
                    f"{S3_FOLDERS.get('summary', 'summary')}/"
                    f"{meeting_id}/{session_id}/"
                    f"{meeting_id}_{session_id}_summary.docx"
                )
            return (
                f"{S3_FOLDERS.get('summary', 'summary')}/"
                f"{meeting_id}_{user_id}_summary.docx"
            )

        elif doc_type == "subtitles":
            if session_id:
                return f"{S3_FOLDERS.get('subtitles', 'subtitles')}/{meeting_id}/{session_id}"
            return f"{S3_FOLDERS.get('subtitles', 'subtitles')}"

        elif doc_type == "trainer_evaluation":
            if session_id:
                return (
                    f"trainer_evaluations/{meeting_id}/{session_id}/"
                    f"{meeting_id}_{session_id}_evaluation.docx"
                )
            return f"trainer_evaluations/{meeting_id}_{user_id}_evaluation.docx"

        logger.warning(f"Unknown doc_type: {doc_type}")
        return ""

    except Exception as e:
        logger.error(f"Error building document S3 path: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return ""

def is_user_allowed(meeting_id: str, email: str = "", user_id: str = "") -> bool:
    """Check if user is allowed to access meeting recording - FIXED VERSION"""
    try:
        if not email and not user_id:
            return False
        
        # Normalize email
        if email:
            email = email.strip().lower()
        
        # Check 1: MongoDB - Is user the video uploader or in visible_to?
        video = collection.find_one({"meeting_id": meeting_id})
        if video:
            if user_id and str(video.get("user_id", "")) == str(user_id):
                logger.info(f"✅ Access granted: User {user_id} is uploader")
                return True
            
            visible_to = video.get("visible_to", [])
            if email and visible_to:
                visible_emails = [e.strip().lower() for e in visible_to if e]
                if email in visible_emails:
                    logger.info(f"✅ Access granted: Email {email} in visible_to")
                    return True
            
        with connection.cursor() as cursor:
            # Check 2: Is user the meeting host?
            if user_id:
                cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                row = cursor.fetchone()
                if row and str(row[0]) == str(user_id):
                    logger.info(f"✅ Access granted: User {user_id} is host")
                    return True

            # Check 3: Is user a participant?
            if user_id:
                cursor.execute(
                    "SELECT COUNT(*) FROM tbl_Participants WHERE Meeting_ID = %s AND User_ID = %s",
                    [meeting_id, user_id]
                )
                if cursor.fetchone()[0] > 0:
                    logger.info(f"✅ Access granted: User {user_id} is participant")
                    return True

            # Check 4: ScheduledMeeting email
            if email:
                cursor.execute("SELECT email FROM tbl_ScheduledMeetings WHERE id = %s", [meeting_id])
                row = cursor.fetchone()
                if row and row[0]:
                    scheduled_emails = [e.strip().lower() for e in row[0].split(',') if e.strip()]
                    if email in scheduled_emails:
                        logger.info(f"✅ Access granted: Email in ScheduledMeeting")
                        return True

            # Check 5: CalendarMeeting emails
            if email:
                cursor.execute(
                    "SELECT email, guestEmails, attendees FROM tbl_CalendarMeetings WHERE id = %s",
                    [meeting_id]
                )
                row = cursor.fetchone()
                if row:
                    for field in row:
                        if field:
                            cal_emails = [e.strip().lower() for e in re.split(r'[;,]', field) if e.strip()]
                            if email in cal_emails:
                                logger.info(f"✅ Access granted: Email in CalendarMeeting")
                                return True

            # Check 6: Get email from user_id and recheck
            if user_id:
                cursor.execute("SELECT Email FROM tbl_Users WHERE ID = %s", [user_id])
                row = cursor.fetchone()
                if row and row[0]:
                    user_email = row[0].strip().lower()
                    # Check this email in visible_to
                    if video:
                        visible_to = video.get("visible_to", [])
                        visible_emails = [e.strip().lower() for e in visible_to if e]
                        if user_email in visible_emails:
                            logger.info(f"✅ Access granted: User email {user_email} in visible_to")
                            return True

        logger.warning(f"❌ Access denied for user_id={user_id}, email={email}")
        return False
        
    except Exception as e:
        logger.error(f"Error in is_user_allowed: {e}")
        return False

def is_user_allowed_debug(meeting_id: str, email: str = "", user_id: str = "") -> bool:
    """Debug version - more permissive for testing."""
    try:
        if email or user_id:
            return True
        return False
    except Exception:
        return False

def sanitize_dot_code(dot_code: str) -> str:
    """
    Ensure ONLY valid Graphviz DOT syntax reaches the renderer.
    Removes bash, markdown, comments, and explanations.
    IMPROVED: Better handling of various DOT code formats from LLMs.
    """
    if not dot_code:
        raise ValueError("Empty DOT content provided")
    
    lines = dot_code.strip().splitlines()
    cleaned = []
    found_graph_start = False
    brace_count = 0

    for line in lines:
        l = line.strip()

        # Skip empty lines, markdown, shell commands, explanations
        if (
            not l
            or l.startswith("```")
            or l.startswith("#!")
            or l.lower().startswith("bash")
            or l.lower().startswith("sh ")
            or l.lower().startswith("echo ")
            or l.lower().startswith("here is")
            or l.lower().startswith("here's")
            or l.lower().startswith("the following")
            or l.lower().startswith("this is")
            or l.lower().startswith("below is")
            or l.lower().startswith("note:")
        ):
            continue
        
        # Check if this line starts the graph
        if not found_graph_start:
            if l.lower().startswith(("digraph", "graph")):
                found_graph_start = True
                cleaned.append(line)
                brace_count += l.count('{') - l.count('}')
            continue
        
        # Once we found graph start, collect lines until braces balance
        cleaned.append(line)
        brace_count += l.count('{') - l.count('}')
        
        # Stop if braces are balanced (graph is complete)
        if brace_count <= 0 and found_graph_start:
            break

    final = "\n".join(cleaned).strip()

    # Validation
    if not final.lower().startswith(("digraph", "graph")):
        raise ValueError("Invalid DOT content: missing digraph/graph header")
    
    # Ensure closing brace exists
    if '{' in final and final.count('{') > final.count('}'):
        final += "\n}"
    
    return final

def enhance_dot_styling(dot_code: str) -> str:
    """
    Inject professional styling into LLM-generated DOT code.
    Adds colors, rounded boxes, fonts to make mind map visually appealing.
    """
    if not dot_code:
        return dot_code
 
    # Don't double-inject if already styled
    if "fillcolor" in dot_code and "fontname" in dot_code:
        return dot_code
 
    brace_idx = dot_code.find('{')
    if brace_idx == -1:
        return dot_code
 
    # Inject default node/edge styling after opening brace
    style_block = """
    graph [rankdir=LR, splines=ortho, nodesep=0.8, ranksep=1.2];
    node [
        shape=box,
        style="rounded,filled",
        fontname="Arial",
        fontsize=11,
        color="#2E86C1",
        fillcolor="#EBF5FB",
        penwidth=1.5
    ];
    edge [
        color="#5DADE2",
        penwidth=1.2,
        arrowsize=0.8
    ];
"""
    styled = dot_code[:brace_idx + 1] + style_block + dot_code[brace_idx + 1:]
 
    # FIX: Find the FIRST real node (skip graph/node/edge/digraph/subgraph keywords)
    import re
    skip_keywords = {'node', 'edge', 'graph', 'digraph', 'subgraph'}
    for match in re.finditer(r'(\w+)\s*\[', styled):
        node_name = match.group(1)
        if node_name.lower() not in skip_keywords:
            # Found the actual root node — style it with dark blue
            # Use exact match with word boundary to avoid replacing partial names
            old_pattern = f'{node_name}['
            new_pattern = f'{node_name}[fillcolor="#1A5276", fontcolor="white", fontsize=13, penwidth=2, '
            # Replace only the FIRST occurrence after style block
            pos = styled.find(old_pattern, brace_idx + len(style_block))
            if pos != -1:
                styled = styled[:pos] + new_pattern + styled[pos + len(old_pattern):]
            break  # Only style the first real node as root
 
    return styled


def get_meeting_participants_emails(meeting_id: str) -> list:
    """Get all participant emails for a meeting from different meeting types."""
    visible_to_emails = []
    
    try:
        with connection.cursor() as cursor:
            # Check Scheduled Meetings
            cursor.execute("SELECT email FROM tbl_ScheduledMeetings WHERE id = %s", [meeting_id])
            row = cursor.fetchone()
            if row and row[0]:
                visible_to_emails += [e.strip() for e in row[0].split(',') if e.strip()]

            # Check Calendar Meetings
            cursor.execute("SELECT email, guestEmails, attendees FROM tbl_CalendarMeetings WHERE id = %s", [meeting_id])
            row = cursor.fetchone()
            if row:
                for field in row:
                    if field:
                        emails = [e.strip() for e in re.split(r'[;,]', field) if e.strip()]
                        visible_to_emails += emails

            # Check Instant Meeting Participants
            cursor.execute("SELECT User_ID FROM tbl_Participants WHERE Meeting_ID = %s", [meeting_id])
            user_ids = [r[0] for r in cursor.fetchall() if r[0]]

            if user_ids:
                format_strings = ','.join(['%s'] * len(user_ids))
                cursor.execute(f"SELECT Email FROM tbl_Users WHERE ID IN ({format_strings})", user_ids)
                visible_to_emails += [r[0].strip() for r in cursor.fetchall() if r[0]]

        # Remove duplicates and empty strings
        visible_to_emails = list(set([email for email in visible_to_emails if email]))
        
        logger.info(f"Found {len(visible_to_emails)} participant emails for meeting {meeting_id}")
        return visible_to_emails
        
    except Exception as e:
        logger.error(f"Failed to get meeting participants emails for meeting {meeting_id}: {e}")
        return []


def get_user_details(user_id: str) -> dict:
    """Get user details (name, email) from database."""
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT ID, full_name, email 
                FROM tbl_Users 
                WHERE ID = %s
            """, [user_id])
            row = cursor.fetchone()
            
            if row:
                user_id_db, full_name, email = row
                
                # Priority: full_name > email prefix > User ID
                if full_name and full_name.strip():
                    display_name = full_name.strip()
                elif email and email.strip():
                    display_name = email.split('@')[0]
                else:
                    display_name = f"User {user_id}"
                
                logger.info(f"✅ Found user {user_id}: full_name='{full_name}', email='{email}', display='{display_name}'")
                
                return {
                    "user_id": str(user_id_db),
                    "user_name": display_name,
                    "user_email": email or ""
                }
            
            logger.warning(f"❌ User {user_id} not found in tbl_Users")
            return {
                "user_id": str(user_id),
                "user_name": f"User {user_id}",
                "user_email": ""
            }
            
    except Exception as e:
        logger.error(f"❌ Error getting user details for {user_id}: {e}")
        return {
            "user_id": str(user_id),
            "user_name": f"User {user_id}",
            "user_email": ""
        }

def clean_markdown(text: str) -> str:
    """
    Cleans markdown syntax from LLM summary while preserving hierarchy and readability.
    """
    import re

    # Allow code blocks. Only strip markdown syntax, not code.
    text = re.sub(r"```(\w+)?", "", text)
    text = text.replace("```", "")

    # remove bold
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    # inline code cleanup
    text = re.sub(r"`([^`]*)`", r"\1", text)

    # convert headings (#, ##, ###)
    text = re.sub(r"^### (.*)", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"^## (.*)", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"^# (.*)", r"\1", text, flags=re.MULTILINE)

    # normalize spacing
    text = re.sub(r"\n{2,}", "\n\n", text)

    return text.strip()

def format_srt_time(seconds: float) -> str:
    from datetime import timedelta

    td = timedelta(seconds=seconds)
    total_seconds = int(td.total_seconds())
    millis = int((td.total_seconds() - total_seconds) * 1000)

    return f"{str(timedelta(seconds=total_seconds)).zfill(8)},{millis:03}"

def create_srt_from_segments(segments: list, output_path: str):
    import logging
    logger = logging.getLogger("video_processor")

    if not segments:
        logger.warning(f"[WARN] No segments found. Writing fallback subtitle: {output_path}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(
                "1\n"
                "00:00:00,000 --> 00:00:05,000\n"
                "[No speech detected]\n\n"
            )
        return

    logger.info(f"[INFO] Writing {len(segments)} subtitle segments → {output_path}")

    with open(output_path, "w", encoding="utf-8") as f:
        for idx, seg in enumerate(segments, start=1):
            start = format_srt_time(float(seg.get("start", 0)))
            end = format_srt_time(float(seg.get("end", 0)))
            text = (seg.get("text") or "").strip()

            if not text:
                text = "[Silence]"

            f.write(f"{idx}\n")
            f.write(f"{start} --> {end}\n")
            f.write(f"{text}\n\n")

def generate_graph(dot_code: str, output_path: str):
    from graphviz import Source
    s = Source(dot_code)
    return s.render(filename=output_path, format="png", cleanup=True)

_NARRATION_VERBS = (
    r"discussed|explained|demonstrated|showed|emphasized|highlighted|covered|"
    r"talked about|touched on|touched upon|mentioned|described|outlined|"
    r"presented|introduced|elaborated on|elaborated|went over|walked through|"
    r"reviewed|noted|stated|pointed out|illustrated|compared|also discussed|"
    r"further explained|then explained|also explained|also mentioned|"
    r"also touched on|also demonstrated|also showed|continued to explain|"
    r"continued discussing"
)

_NARRATION_PATTERNS = [
    re.compile(rf"^the speaker (?:{_NARRATION_VERBS})\s+", re.IGNORECASE),
    re.compile(rf"^they (?:{_NARRATION_VERBS})\s+", re.IGNORECASE),
    re.compile(rf"^the (?:instructor|trainer|presenter|teacher|lecturer|host) (?:{_NARRATION_VERBS})\s+", re.IGNORECASE),
    re.compile(r"^in the (?:meeting|session|class|lecture|presentation|talk),?\s+", re.IGNORECASE),
    re.compile(r"^during the (?:meeting|session|class|lecture|presentation|talk),?\s+", re.IGNORECASE),
    re.compile(rf"^the speaker['\u2019]s (?:explanation|discussion|demonstration|description|presentation) of\s+", re.IGNORECASE),
]

def _strip_narration_sentence(sentence: str) -> str:
    """Strip leading 'The speaker X' / 'They X' narration from one sentence."""
    s = sentence.strip()
    if not s:
        return s
    for pat in _NARRATION_PATTERNS:
        m = pat.match(s)
        if m:
            rest = s[m.end():].strip()
            if rest:
                rest = rest[0].upper() + rest[1:] if len(rest) > 1 else rest.upper()
                return rest
    return s

def strip_narration_voice(text: str) -> str:
    """Apply narration stripping sentence-by-sentence over a block of text."""
    if not text:
        return text
    parts = re.split(r'(?<=[.!?])\s+', text.strip())
    cleaned = [_strip_narration_sentence(p) for p in parts if p.strip()]
    return " ".join(cleaned)

def save_trainer_evaluation_docx(evaluation: dict, path: str, meeting_title: str = ""):
    """Render trainer evaluation scores as a host-only DOCX."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    import pytz

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    h = doc.add_heading("Trainer Performance Evaluation", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    ist = pytz.timezone("Asia/Kolkata")
    date_str = datetime.now(ist).strftime("%B %d, %Y at %I:%M %p")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generated: {date_str}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    if meeting_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Session: {meeting_title}")
        r.font.size = Pt(11)
        r.font.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run(
        "CONFIDENTIAL — This evaluation is intended solely for the trainer/host. "
        "Do not distribute to participants."
    )
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_heading("Performance Scores", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Score"

    rows = [
        ("Technical Content",   f"{evaluation.get('technical_content', 0)}%"),
        ("Explanation Clarity", f"{evaluation.get('explanation_clarity', 0)}%"),
        ("Friendliness",        f"{evaluation.get('friendliness', 0)}%"),
        ("Communication",       f"{evaluation.get('communication', 0)}%"),
    ]
    for label, val in rows:
        c = table.add_row().cells
        c[0].text = label
        c[1].text = val

    doc.add_paragraph()

    feedback = (evaluation.get("overall_feedback") or "").strip()
    if feedback:
        doc.add_heading("Overall Feedback", level=2)
        doc.add_paragraph(feedback)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(
        "Scoring methodology: scores are derived from transcript-based analysis. "
        "Use as one input among many; do not treat as a sole performance measure."
    )
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.save(path)


def save_docx(content: str, path: str, image_path: str = None, title: str = ""):
    """
    Save content to DOCX with professional formatting.
    Handles: Executive Summary, Key Topics, Action Items, Detailed Discussion,
    Examples & Demonstrations (conditional), Key Takeaways, Concept Map
    """
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx import Document
    import re
    from datetime import datetime
 
    doc = Document()
 
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
 
    # === Document Header: Title ===
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
 
    # === Date line ONLY for summary documents, NOT for transcript ===
    is_summary = title and "summary" in title.lower()
    if is_summary:
        import pytz
        ist = pytz.timezone("Asia/Kolkata")
        now = datetime.now(ist)
        date_str = now.strftime("%B %d, %Y at %I:%M %p")
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Generated: {date_str}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
        date_run.font.italic = True
        doc.add_paragraph()  # spacer
 
    # Track block states separately:
    #   inside_dot_block   → ```dot / ```graphviz blocks (mind map source — skip, image renders elsewhere)
    #   inside_code_fence  → ```python / ```js / ```sql / etc. (real code — keep, render with monospace)
    inside_dot_block = False
    inside_code_fence = False
    current_section = ""

    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            # Preserve blank lines inside code blocks for readability
            if inside_code_fence:
                doc.add_paragraph()
            continue

        if stripped.startswith("```"):
            lang_hint = stripped[3:].strip().lower()
            if not inside_dot_block and not inside_code_fence:
                # Opening a fence — decide DOT vs code
                if lang_hint in ("dot", "graphviz"):
                    inside_dot_block = True
                else:
                    inside_code_fence = True
            elif inside_dot_block:
                inside_dot_block = False  # closing DOT block
            elif inside_code_fence:
                inside_code_fence = False  # closing code block
            continue

        # Inside DOT block — skip (mind map source code; image renders separately)
        if inside_dot_block:
            continue

        # Inside code fence — render with monospace font, preserving original indentation
        # (use `line` not `stripped` so Python indentation survives in the .docx)
        if inside_code_fence:
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip())
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            continue
 
        # Skip raw DOT syntax that leaked outside backticks
        if stripped.lower().startswith(("digraph", "graph ")) and "{" in stripped:
            inside_dot_block = True
            continue
        if stripped == "}" and inside_dot_block:
            inside_dot_block = False
            continue
        if "->" in stripped and ("[" in stripped or ";" in stripped):
            continue
        if stripped.startswith("node[") or stripped.startswith("edge["):
            continue
 
        # === Section Headers (## Executive Summary, etc.) ===
        if stripped.startswith("## "):
            section_name = stripped[3:].strip()
            current_section = section_name.upper()
 
            # Skip Concept Map header (image added separately at the end)
            if "CONCEPT MAP" in current_section or "MIND MAP" in current_section:
                continue

            # Skip "Action Items" heading if no real action items follow (only filler placeholder)
            if "ACTION ITEMS" in current_section:
                remaining = content.split(stripped, 1)[-1] if stripped in content else ""
                has_real_action_items = False
                for future_line in remaining.splitlines():
                    fl = future_line.strip()
                    if not fl:
                        continue
                    if fl.startswith("## "):
                        break
                    fl_lower = fl.lower()
                    # Skip filler placeholder lines
                    if any(p in fl_lower for p in [
                        "no specific action", "no action item", "no tasks were",
                        "were not identified", "not identified in this",
                        "no follow-up", "no next steps"
                    ]):
                        continue
                    # Real action item starts with bullet or bracket tag
                    if fl.startswith("- ") or fl.startswith("[") or fl.startswith("• "):
                        has_real_action_items = True
                        break
                if not has_real_action_items:
                    continue

            # Skip "Complete Code Example" heading if no real code follows
            if "COMPLETE CODE" in current_section:
                remaining = content.split(stripped, 1)[-1] if stripped in content else ""
                has_real_code = False
                for future_line in remaining.splitlines():
                    fl = future_line.strip()
                    if not fl:
                        continue
                    if fl.startswith("## "):
                        break
                    if fl.lower().startswith(("since no", "no specific", "no code", "no coding", "not applicable", "there is no")):
                        break
                    if fl.startswith(("import ", "from ", "def ", "class ", "print(", "    ", "\t")) or "=" in fl:
                        has_real_code = True
                        break
                if not has_real_code:
                    continue

            # Skip "Examples & Demonstrations" heading if no real examples follow
            if "EXAMPLES" in current_section and "DEMONSTRATIONS" in current_section:
                remaining = content.split(stripped, 1)[-1] if stripped in content else ""
                has_real_examples = False
                for future_line in remaining.splitlines():
                    fl = future_line.strip()
                    if not fl:
                        continue
                    if fl.startswith("## "):
                        break
                    if any(neg in fl.lower() for neg in ["no specific", "no example", "no demonstration", "no code", "not applicable", "were not", "was not", "did not"]):
                        continue
                    if fl.startswith("### ") or fl.startswith("- ") or len(fl) > 50:
                        has_real_examples = True
                        break
                if not has_real_examples:
                    continue

            heading = doc.add_heading(section_name, level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
                run.font.size = Pt(16)
            continue
 
        # === Sub-headers (### Topic Name) ===
        if stripped.startswith("### "):
            heading = doc.add_heading(stripped[4:].strip(), level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
                run.font.size = Pt(13)
            continue
 
        # === Bullet points (- text) ===
        if stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
 
            # Check for bold prefix like [Learners] or [Group]
            bold_match = re.match(r'\[([^\]]+)\]\s*(.*)', bullet_text)
            if bold_match:
                run_bold = p.add_run(f"[{bold_match.group(1)}] ")
                run_bold.bold = True
                run_bold.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
                run_bold.font.size = Pt(11)
                rest_text = bold_match.group(2)
                if is_summary:
                    rest_text = strip_narration_voice(rest_text)
                run_normal = p.add_run(rest_text)
                run_normal.font.size = Pt(11)
            else:
                p.text = ""
                cleaned_bullet = strip_narration_voice(bullet_text) if is_summary else bullet_text
                run = p.add_run(cleaned_bullet)
                run.font.size = Pt(11)
            continue
 
        # === Numbered items (1. text) ===
        if re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(stripped, style='List Number')
            for run in p.runs:
                run.font.size = Pt(11)
            continue
 
        # === Code lines (in code section or indented) ===
        if "COMPLETE CODE" in current_section or stripped.startswith("    ") or stripped.startswith("\t"):
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            continue
 
        # === Feedback line at the end ===
        if stripped.lower().startswith("was this summary helpful"):
            doc.add_paragraph()  # spacer
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
            run.font.italic = True
            continue
 
        # === Separator line ===
        if stripped == "---":
            continue
 
        # === Bold wrapped text (**text**) ===
        if stripped.startswith("**") and stripped.endswith("**"):
            clean_text = stripped.replace("**", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
            continue
 
        # === Skip "no content" filler — ONLY inside conditional summary sections.
        # These filters exist to clean up summary text like "no code example was provided".
        # They must NOT run on transcripts (where current_section is empty) or on regular
        # summary prose, because real sentences contain words like "not" + "example".
        if current_section in ("EXAMPLES & DEMONSTRATIONS", "COMPLETE CODE EXAMPLE"):
            lower = stripped.lower()
            has_negative = any(w in lower for w in ["no ", "not ", "without ", "didn't", "did not", "does not", "is not", "were not", "was not"])
            has_topic_word = any(w in lower for w in ["code example", "code is", "code was", "coding", "programming", "example", "demonstration"])
            if has_negative and has_topic_word:
                continue

            if ("no code" in lower or "no example" in lower or "no demonstration" in lower or "no specific example" in lower) and ("not" in lower or "no " in lower):
                continue

        lower_global = stripped.lower()
        global_filler_starters = (
            "no ", "no specific ", "since no ", "there is no ",
            "there are no ", "there were no ", "not applicable"
        )
        global_filler_topics = [
            "code", "coding", "programming",
            "example", "demonstration", "demo",
            "action item", "specific action",
            "tasks were", "task were",
            "follow-up", "next steps"
        ]
        if any(lower_global.startswith(s) for s in global_filler_starters):
            if any(t in lower_global for t in global_filler_topics):
                continue

        # === Regular paragraph ===
        para_text = strip_narration_voice(stripped) if is_summary else stripped
        p = doc.add_paragraph(para_text)
        for run in p.runs:
            run.font.size = Pt(11)
 
    # === Add Concept Map Image at end ===
    if image_path and os.path.exists(image_path):
        doc.add_page_break()
        heading = doc.add_heading("Concept Map", level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        doc.add_picture(image_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    doc.save(path)


def send_recording_completion_notifications(meeting_id, video_url, transcript_url=None, summary_url=None):
    """
    Send professional notifications to all meeting participants when recording is ready.
    Different messages for host vs participants.
    
    INCLUDES:
    - Duplicate prevention at function level (prevents double-calling)
    - Duplicate prevention at participant level (prevents same email twice)
    - Case-insensitive email handling
    - Professional HTML formatting (NO EMOJIS - MySQL utf8 compatible)
    """
    if not meeting_id or not video_url:
        logging.warning("Missing meeting_id or video_url for recording notifications")
        return 0

    # ✅ CRITICAL FIX 1: Check if notifications already sent for this meeting
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT COUNT(*) FROM tbl_Notifications 
                WHERE meeting_id = %s 
                  AND notification_type IN ('recording_completed', 'recording_completed_host')
            """, [str(meeting_id)])
            
            existing_count = cursor.fetchone()[0]
            
            if existing_count > 0:
                logging.warning(f"Recording notifications already sent for meeting {meeting_id} ({existing_count} exist). Skipping to prevent duplicates.")
                return 0
    except Exception as e:
        logging.warning(f"Could not check for existing notifications: {e}")
        # Continue anyway - better to risk duplicate than miss notification

    try:
        ensure_notification_tables()
    except Exception as e:
        logging.error(f"Failed to ensure notification tables: {e}")
        return 0

    ist = pytz.timezone("Asia/Kolkata")
    now = datetime.now(ist)
    sent_count = 0
    skipped_count = 0

    # ========== FETCH MEETING DETAILS ==========
    meeting_info = _get_recording_meeting_info(meeting_id)
    
    meeting_title = meeting_info.get('title') or f"Meeting {meeting_id[:8]}"
    meeting_type = meeting_info.get('meeting_type') or 'Meeting'
    host_name = meeting_info.get('host_name') or 'Host'
    host_email = meeting_info.get('host_email')
    video_name = meeting_info.get('custom_recording_name') or f"Recording - {meeting_title}"
    meeting_date = meeting_info.get('start_time')
    
    # Format date nicely
    if meeting_date:
        if isinstance(meeting_date, str):
            try:
                meeting_date = datetime.strptime(meeting_date, '%Y-%m-%d %H:%M:%S')
            except:
                meeting_date = now
        formatted_date = meeting_date.strftime('%B %d, %Y at %I:%M %p')
    else:
        formatted_date = now.strftime('%B %d, %Y')

    # ========== GET ALL PARTICIPANTS ==========
    participants = _get_recording_participants(meeting_id)
    
    if not participants:
        logging.warning(f"No participants found for meeting {meeting_id}")
        return 0

    # ✅ CRITICAL FIX 2: Deduplicate participants list (case-insensitive)
    unique_participants = []
    seen_emails = set()
    for p in participants:
        email = p.get('email', '') if isinstance(p, dict) else p
        if email and '@' in str(email):
            normalized_email = str(email).strip().lower()
            if normalized_email not in seen_emails:
                seen_emails.add(normalized_email)
                unique_participants.append(normalized_email)

    if len(unique_participants) != len(participants):
        logging.info(f"Deduplicated participants: {len(participants)} -> {len(unique_participants)}")

    logging.info(f"Sending recording notifications to {len(unique_participants)} unique participants")

    # ========== BUILD RESOURCE LIST (Professional - NO EMOJIS) ==========
    resources = ["Video Recording"]
    if transcript_url:
        resources.append("Full Transcript")
    if summary_url:
        resources.append("AI-Generated Summary & Mind Map")

    resources_text = "\n".join([f"  * {r}" for r in resources])

    # ========== SEND NOTIFICATIONS ==========
    for participant_email in unique_participants:
        if not participant_email or '@' not in participant_email:
            continue

        try:
            participant_email = participant_email.strip().lower()
            
            # ✅ CRITICAL FIX 3: Check if notification already exists for this specific user
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT COUNT(*) FROM tbl_Notifications 
                    WHERE meeting_id = %s 
                      AND LOWER(recipient_email) = %s
                      AND notification_type IN ('recording_completed', 'recording_completed_host')
                """, [str(meeting_id), participant_email])
                
                if cursor.fetchone()[0] > 0:
                    logging.info(f"Notification already exists for {participant_email}, skipping")
                    skipped_count += 1
                    continue

            is_host = (
                host_email and 
                participant_email == host_email.strip().lower()
            )

            # Calculate other participants count for host message
            other_participants_count = len(unique_participants) - 1

            if is_host:
                # ========== HOST NOTIFICATION (Professional Format - NO EMOJIS) ==========
                notification_type = "recording_completed_host"
                title = "Your Recording is Ready"
                message = (
                    f"<b>Great news!</b> Your meeting recording has been successfully processed and is now available.\n\n"
                    f"<b>Meeting Details</b>\n"
                    f"  * <b>Title:</b> {meeting_title}\n"
                    f"  * <b>Date:</b> {formatted_date}\n"
                    f"  * <b>Recording:</b> {video_name}\n\n"
                    f"<b>Available Resources</b>\n"
                    f"{resources_text}\n\n"
                    f"<b>Notification Status</b>\n"
                    f"  * {other_participants_count} participant(s) have been notified.\n\n"
                    f"<i>Click here to access your recording dashboard.</i>"
                )
                priority = "normal"
            else:
                # ========== PARTICIPANT NOTIFICATION (Professional Format - NO EMOJIS) ==========
                notification_type = "recording_completed"
                title = "Meeting Recording Available"
                message = (
                    f"<b>Hello!</b> A recording from your recent meeting is now available for you to view.\n\n"
                    f"<b>Meeting Details</b>\n"
                    f"  * <b>Title:</b> {meeting_title}\n"
                    f"  * <b>Host:</b> {host_name}\n"
                    f"  * <b>Date:</b> {formatted_date}\n"
                    f"  * <b>Recording:</b> {video_name}\n\n"
                    f"<b>You Have Access To</b>\n"
                    f"{resources_text}\n\n"
                    f"<i>Click here to view the recording in your dashboard.</i>"
                )
                priority = "normal"

            # ========== INSERT NOTIFICATION ==========
            notification_id = short_id()
            
            with connection.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO tbl_Notifications (
                        id, recipient_email, meeting_id, notification_type,
                        title, message, meeting_title, meeting_url,
                        is_read, priority, created_at
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, [
                    notification_id,
                    participant_email,
                    str(meeting_id),
                    notification_type,
                    title,
                    message,
                    meeting_title,
                    video_url,
                    False,
                    priority,
                    now
                ])
                
                # ✅ Explicit commit
                connection.commit()

                if cursor.rowcount > 0:
                    sent_count += 1
                    logging.info(f"Recording notification sent to {participant_email} ({'host' if is_host else 'participant'})")
                else:
                    logging.error(f"INSERT returned 0 rows for {participant_email}")

        except Exception as e:
            logging.error(f"Failed to send recording notification to {participant_email}: {e}", exc_info=True)

    logging.info(f"Recording notifications completed: {sent_count} sent, {skipped_count} skipped (duplicates)")
    return sent_count

# === NEW: CHUNKED TRANSCRIPTION FROM FASTAPI ===
async def transcribe_chunk(chunk_file: str, offset: float):
    try:
        with open(chunk_file, "rb") as f:
            result = groq_client.audio.translations.create(
                model="whisper-large-v3",
                file=f,
                response_format="verbose_json",
                prompt="Technical training session. Preserve technical terms, product names, commands, and acronyms exactly (e.g., SAP HANA, HDBLCM, HCMT, RPM, nslookup, hostname, IP address). Output strictly in English.",
                temperature=0.0
            )

        segments = []

        # CASE 1 — Whisper returned real timestamped segments
        if hasattr(result, "segments") and result.segments:
            for seg in result.segments:
                segments.append({
                    "start": offset + float(seg["start"]),
                    "end": offset + float(seg["end"]),
                    "text": seg["text"].strip()
                })

        # CASE 2 — No segments, fallback
        else:
            text = (result.text or "").strip()
            if not text:
                text = "[No speech detected]"

            segments.append({
                "start": offset,
                "end": offset + 5,
                "text": text
            })

        return segments

    except Exception as e:
        logger.error(f"[GROQ ERROR] {e}")
        return [{
            "start": offset,
            "end": offset + 5,
            "text": "[Transcription failed]"
        }]

def summarize_segment(transcript: str, context: str = ""):
    """
    Generate summary for transcript.
    - Short transcripts (<15K words): Single API call with full detailed prompt
    - Long transcripts (>=15K words): Split into chunks, summarize each, then combine
    Supports 1-4 hour meetings reliably.
    """
    if groq_client is None:
        logger.error("❌ Cannot generate summary - groq_client is None!")
        logger.error(f"❌ GROQ_API_KEY value: '{os.getenv('GROQ_API_KEY', 'NOT_SET')[:10]}...'")
        return "Summary unavailable - Groq API not configured. Please check GROQ_API_KEY environment variable."
    
    word_count = len(transcript.split())
    logger.info(f"📝 Transcript length: {word_count} words")
    
    # For short transcripts (< 1 hour meetings), use single-call approach with full detailed prompt
    if word_count < 15000:
        logger.info(f"📝 Short transcript - using single API call with full detailed prompt")
        return _generate_single_summary(transcript, context)
    
    # For long transcripts (2+ hour meetings), use chunked summarization
    logger.info(f"📚 Long transcript detected ({word_count} words), using chunked summarization")
    return _generate_chunked_summary(transcript, context)


def _generate_single_summary(transcript: str, context: str = "") -> str:
    """
    Detailed summarization for short transcripts (< 15K words / ~1 hour meetings).
    Merges old enterprise documentation quality rules with new structured format.
    """
    prompt = f"""
You are a senior documentation and technical writing expert. Your task is to convert the following raw transcript into a comprehensive, highly accurate, and professionally structured meeting summary document.

---------------------------------------------------------------------
GROUNDING RULES (HIGHEST PRIORITY — TOPIC-LEVEL FIDELITY):

The TRANSCRIPT decides WHAT topics to write about. WITHIN those topics, you may write naturally and clearly.

ALLOWED — when a topic IS discussed in the transcript:
- Explain it well with natural depth, related context, and helpful clarity for a reader who wasn't there.
- Cover what the speaker said and elaborate enough to make the discussion understandable.
- Use proportional length — short discussions get short sections, long discussions get long sections.

NOT ALLOWED — even if it would make the document "better":
- DO NOT introduce topics the speaker never mentioned. If the speaker talked about Python and automation, do not add sections about Java, AWS, web development, finance, or healthcare just because they're related.
- DO NOT fabricate action items. If no tasks were assigned, OMIT the entire Action Items section (no heading, no placeholder text). Same rule for Examples & Demonstrations and Complete Code Example.
- DO NOT invent examples or demonstrations. The Examples & Demonstrations section may only include things the speaker actually showed. If none, omit the section entirely.
- DO NOT invent code. Complete Code Example may only appear if the speaker actually wrote or showed code.
- DO NOT invent key takeaways. Takeaways must reflect what the speaker actually emphasized.
- DO NOT correct, modernize, or substitute tools/commands/APIs the speaker mentioned. Quote them as said.

THE TEST: Could a reader of the transcript point to where this topic / action item / example came from? If yes, include it. If no, leave it out.

---------------------------------------------------------------------
VOICE RULE — CONTENT-FIRST, NOT NARRATION (CRITICAL — APPLIES TO EVERY TOPIC):

This rule is universal. It applies whether the meeting is about Python, SAP HANA, Kubernetes, sales pipeline, Q3 revenue, recruitment process, GST calculation, blood pressure measurement, a recipe, a yoga routine, a history lesson, contract law, IELTS prep, car maintenance, or anything else. There is no exception. Treat every meeting summary like a textbook chapter or SAP/Oracle/Google/Microsoft/AWS internal documentation — never like meeting minutes.

ABSOLUTELY BANNED — NEVER start a sentence with these phrases, no matter the subject:
- "The speaker discussed/explained/demonstrated/showed/emphasized/highlighted/covered/talked about/touched on/mentioned/described/presented/introduced/illustrated/compared..."
- "They discussed/explained/demonstrated/showed/emphasized/..." (when referring back to the speaker)
- "The instructor/trainer/presenter/teacher/lecturer/host discussed/explained/..."
- "In the meeting/session/class/lecture/presentation, the speaker..."
- "During the meeting/session/class/lecture/presentation, ..."
- "The speaker's explanation/discussion/demonstration of..."

REQUIRED — Write in CONTENT-FIRST DECLARATIVE VOICE in every section:
- The TOPIC, CONCEPT, TOOL, PROCESS, or SUBJECT MATTER is the subject of the sentence — NEVER the speaker.
- Explain WHAT something is, HOW it works, WHY it matters, WHEN it applies — not WHO said it.
- Use teaching verbs: "is", "are", "works by", "requires", "produces", "consists of", "depends on", "follows the formula", "is calculated by", "is performed using", "begins with", "indicates", "is measured by".

UNIVERSAL EXAMPLES — these patterns hold for every domain:

1) Software / Programming —
WRONG: "The speaker explained how to create a K-means model and calculate the sum of squared errors."
RIGHT: "A K-means model is created with KMeans(n_clusters=k) and fitted via .fit(). The sum of squared errors is accessed through the .inertia_ attribute and indicates cluster tightness — lower values mean tighter clusters."

2) Enterprise IT (SAP / networking / cloud) —
WRONG: "The speaker demonstrated how to install SAP HANA using the HDBLCM tool."
RIGHT: "SAP HANA is installed using the HDBLCM tool, which provides a guided installation flow covering system parameters, network configuration, and license setup. The installer requires the SID, instance number, and system administrator password before proceeding."

3) Business meeting (sales / project / quarterly review) —
WRONG: "The speaker discussed Q3 revenue performance and explained how the team missed targets."
RIGHT: "Q3 revenue came in 8% below target, driven by delayed enterprise renewals and slower mid-market pipeline conversion. Pipeline coverage for Q4 is currently 2.1x against the standard 3x benchmark, indicating a need for accelerated outbound activity."

4) HR / recruitment / process —
WRONG: "The speaker described the recruitment process and explained that screening happens before interviews."
RIGHT: "The recruitment process begins with resume screening, followed by a phone interview, a technical or role-specific assessment, and a panel interview before final selection. Each stage has a defined rejection criteria sheet."

5) Finance / accounting —
WRONG: "The speaker explained how to calculate ROI by dividing net profit by total investment."
RIGHT: "ROI is calculated by dividing net profit by total investment and multiplying by 100 to express as a percentage. A positive ROI indicates the investment generated a return greater than its cost; a negative ROI indicates a loss."

6) Healthcare / medical —
WRONG: "The speaker described how blood pressure should be measured at the same time each day."
RIGHT: "Blood pressure should be measured at the same time each day, after 5 minutes of seated rest, with the arm supported at heart level. Two readings taken one minute apart are averaged for the final value."

7) Cooking / lifestyle —
WRONG: "The speaker demonstrated how to make pasta dough by mixing flour and eggs."
RIGHT: "Pasta dough is made by combining 100g of flour per egg, kneading for 10 minutes until smooth and elastic, and resting it covered for 30 minutes before rolling. Resting allows the gluten to relax, making the dough easier to roll thin."

8) Fitness / general instructional —
WRONG: "The speaker emphasized that proper squat form requires keeping the knees aligned with the toes."
RIGHT: "Proper squat form requires keeping the knees aligned with the toes, the chest upright, and the weight distributed through the heels. Allowing the knees to collapse inward places stress on the medial ligaments and reduces force transfer."

The pattern is the same in every example: subject of the sentence = the topic, not the speaker.

Speaker attribution is permitted ONLY in two cases:
1. Action Items section — naming who owns each task is the whole point.
2. A specific decision or stated opinion whose weight depends on the speaker (e.g., "The CFO approved the revised budget.", "The team lead decided to defer the migration to next sprint.").

For ALL concept explanations, definitions, procedures, examples, troubleshooting steps, best practices, and key takeaways — across every domain — NO speaker references. Just the content.
---------------------------------------------------------------------

QUALITY STANDARDS (apply WITHIN grounded topics):
- Tone: clear, professional, instructional, like SAP/Oracle/Google/Microsoft/AWS internal documentation.
- Teaching style: step-by-step explanation, beginner-friendly structure, simple language.
- Depth: proportional to how much was actually said about each topic. If 20 minutes with 15 steps, write all 15 steps. Never pad to hit a length target.
- Tools/commands/APIs: preserve exactly as the speaker mentioned them. Do not add tools that were not mentioned.
- Consolidate duplicate or fragmented instructions: if a step is repeated, merge into one. Do not invent missing steps to "complete" a procedure.
- Replace any real values (IP addresses, passwords, hostnames, ports, usernames, emails) with placeholder tags: <ip>, <password>, <hostname>, <port>, <username>, <email>. Do NOT alter sentence structure — just swap the values.
---------------------------------------------------------------------

TRANSCRIPT:
\"\"\"{transcript}\n\n{context}\"\"\"

---------------------------------------------------------------------
OUTPUT FORMAT — Use these EXACT section headers in this EXACT order:
---------------------------------------------------------------------

## Executive Summary
Write a comprehensive overview of the entire meeting/session. Cover: what was the main topic, what was discussed or taught, what tools/technologies were mentioned, what procedures were demonstrated, and what was the outcome or conclusion. Write as many sentences as needed to fully capture the scope — do NOT compress to just 2-3 sentences if the meeting covered substantial content.

## Key Topics

### [Topic Name 1]
The topic must come from the transcript. Describe it clearly with proportional depth — short if briefly mentioned, longer if extensively explained. You may add natural context and explanation that helps the reader understand what the speaker said. Do NOT add unrelated topics or applications the speaker never mentioned.

### [Topic Name 2]
Same — topic must be from the transcript, written with appropriate depth and clarity.

### [Topic Name 3]
Same — topic must be from the transcript, written with appropriate depth and clarity.

(Add as many topics as were covered in the transcript. Typically 3-7 topics.)

## Action Items
CONDITIONAL: Include this section ONLY if the speaker explicitly assigned tasks. If no action items were assigned, DO NOT include this section at all — skip the heading entirely. Do NOT write "No specific action items..." or any placeholder text.
- [Who/Group] Action item or task to complete.
- [Who/Group] Another action item.
## Detailed Discussion

### [Topic Name 1]
- Detailed bullet point covering a key concept, step, or decision discussed. Include all relevant tools, platforms, or interfaces used.
- Another detailed point with specifics — commands, values, tools, configurations, paths, parameters.
- Step-by-step technical or procedural instructions if any were discussed, clearly numbered and logically ordered.
- Tips, warnings, Important Notes, or expected outcomes where necessary.
- Validation: how to confirm success — expected outputs, system checks, visual indicators.
- Continue with ALL important details from the transcript for this topic. Do NOT skip or compress steps.

### [Topic Name 2]
- Same thorough treatment. Include any commands, code, configurations, or procedures mentioned.
- Cover technical specifics, parameter values, tool names.
- Include troubleshooting content if discussed: frequent issues, conditions that cause them, corrective actions, file paths, log tools, error codes.
- Include best practices if discussed: efficiency improvements, security tips, SOPs, pitfalls to avoid.

(Repeat for each topic from Key Topics. Each topic should have as many bullets as needed — 3 minimum, but no upper limit. Write everything that was discussed.)

## Examples & Demonstrations
IMPORTANT CONDITIONAL RULE: Include this section ONLY if the transcript contains ACTUAL demonstrations, live examples, code walkthroughs, calculations, predictions, sample data walkthroughs, or hands-on demos performed by the speaker.
If this is a general meeting (standup, project update, review, planning, discussion) with NO demonstrations or worked examples, DO NOT include this section at all — skip the heading entirely.

### Example 1: [Short Descriptive Title]
Full description of the example that was demonstrated. Include:
- Specific input values, parameters, data used
- Step-by-step walkthrough of what was done
- Output/results obtained
- What the example proved, showed, or taught
- Any observations or comparisons made during the demo

### Example 2: [Short Descriptive Title]
Same thorough treatment for the next example.

(Extract ALL real examples from the transcript. Include specific values, inputs, outputs. Do NOT summarize examples in one line — write the full walkthrough as it was demonstrated.)

## Complete Code Example
DECISION RULE — only two outcomes are allowed. NO middle ground.

OUTCOME A: Generate real working code (include the heading + full code)
Trigger this if ANY of the following is true in the transcript:
1. The speaker wrote or dictated actual code
2. The speaker mentioned ANY programming language (Python, Java, JavaScript, C++, SQL, R, Go, Ruby, etc.)
3. The speaker mentioned ANY programming concept (function, class, variable, loop, library, framework, API, IDE, debugging, package, module, etc.)
4. The speaker discussed any algorithm or technique that is normally implemented in code (K-means, linear regression, sorting, neural networks, decision trees, REST APIs, web scraping, data analysis, machine learning, etc.)

When triggered: ANALYZE what concept the speaker actually taught, then GENERATE a complete runnable code example that demonstrates EXACTLY that concept. Use the language the speaker mentioned (default to Python if no specific language). The code MUST include:
- Full imports / package statements
- Real implementation (not pseudocode, not comments-only)
- Inline comments explaining each step
- Correct indentation and syntax
- Sample data or example call so it runs end-to-end
- Expected output written as a comment

Example: speaker discussed K-means clustering and mentioned Python → generate real working Python code using sklearn.cluster.KMeans with sample data, fit/predict, and the elbow technique. NOT an apology paragraph.

OUTCOME B: Omit the entire section
Trigger ONLY when the transcript has ZERO programming content (no language, no programming concept, no algorithm normally implemented in code). Examples: cooking lessons, history lectures, pure business meetings, philosophy discussions.
When triggered: DO NOT write the heading. DO NOT write any text. The reader should see Detailed Discussion flow directly into Key Takeaways with no Code section between them.

ABSOLUTELY FORBIDDEN — these phrases must NEVER appear:
- "No code was discussed in the transcript"
- "No code was explicitly mentioned"
- "No code example was provided"
- "While no code was shown..."
- "Since no code was demonstrated..."
- "However, the speaker mentioned [language]..." (as a substitute for actual code)
- ANY apology text about absent code

If you find yourself about to write apology text, that is the signal to either (a) GENERATE real code based on what was discussed, or (b) OMIT the section entirely. There is no third option.

## Key Takeaways
Write a thorough summary of the most important points from the session. What should attendees remember? What are the critical insights? Include:
- What was implemented or discussed
- Expected outcomes and readiness indicators
- Best practices highlighted during the session
- Any prerequisites or tools required for follow-up
Write as many sentences as needed — do NOT compress to just 2-3 sentences if the meeting was substantial.

## Concept Map
```dot
digraph ConceptMap {{
    node[shape=box, style="rounded,filled"];
    root[label="Main Topic"];
    topic1[label="Topic 1"];
    topic2[label="Topic 2"];
    sub1[label="Sub-topic 1"];
    sub2[label="Sub-topic 2"];
    root -> topic1;
    root -> topic2;
    topic1 -> sub1;
    topic2 -> sub2;
}}
```
(Build a proper concept map covering ALL key topics and their relationships. Include sub-topics, tools, and concepts as nodes.)

---

Was this summary helpful? Share your feedback to help us improve.

---------------------------------------------------------------------
FINAL RULES:
---------------------------------------------------------------------
1. Use the EXACT section headers shown above: ## Executive Summary, ## Key Topics, ## Action Items, ## Detailed Discussion, ## Examples & Demonstrations, ## Complete Code Example, ## Key Takeaways, ## Concept Map
2. Topic names in Key Topics and Detailed Discussion must match exactly.
3. Examples & Demonstrations: ONLY include if transcript has real demos/examples/walkthroughs. General meetings without demos must NOT have this section.
4. Complete Code Example: ONLY include if coding detected per trigger conditions. Non-coding meetings must NOT have this section.
5. Keep technical terms, tool names, commands, and code exactly as mentioned in transcript.
6. The concept map must be valid Graphviz DOT format enclosed in triple backticks.
7. Replace real values with <ip>, <password>, <hostname>, <port>, <username>, <email>.
8. No emojis or decorative formatting.
9. DO NOT compress content. Write proportionally to how much was discussed. Long discussions = long sections.
10. Always end with: "Was this summary helpful? Share your feedback to help us improve."
11. Cross-check all tools, commands, APIs — correct outdated or wrong ones.
12. Merge duplicate instructions into single complete versions.
13. If a conditional section (Action Items, Examples & Demonstrations, Complete Code Example) has no content from the transcript, OMIT it entirely — no heading, no explanation, no placeholder. Do NOT write text like "No code was discussed in the transcript", "No specific examples were given", or "No specific action items were identified". Just leave nothing in its place.
"""

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You are a technical documentation writer producing a content-first reference summary (like SAP/Oracle/Google internal docs), NOT meeting minutes. The transcript decides WHAT topics to cover — never invent NEW topics, action items, or examples. CRITICAL UNIVERSAL VOICE RULE (applies to every topic — software, enterprise, business, healthcare, cooking, fitness, education, anything): NEVER start sentences with 'The speaker discussed/explained/demonstrated', 'They explained/showed', 'The instructor/trainer/presenter...', 'In the meeting/session/class...', or 'During the meeting/session...'. Write content-first declarative sentences where the SUBJECT is the topic/concept/tool/process, not the speaker. Use 'is/works by/requires/produces/is calculated by/begins with' instead of 'the speaker explained/showed/demonstrated'. Speaker attribution allowed ONLY in Action Items (task owners) and for specific decisions whose weight depends on the speaker. SPECIAL RULE for Complete Code Example: if the transcript mentions any programming language, concept, or algorithm normally implemented in code, GENERATE complete working code demonstrating that exact concept. If zero programming content, omit the Code section entirely. NEVER write apology text like 'No code was discussed but...' — produce real code or produce nothing. Same omit-or-nothing rule for Action Items and Examples sections."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=8000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"[GROQ SUMMARY ERROR] {e}")
        return "Summary generation failed."


# ============================================================
# CHANGE 3: ADD this template + REPLACE _generate_chunked_summary
# ============================================================

CHUNKED_FINAL_PROMPT_TEMPLATE = """
You are a senior documentation and technical writing expert. Your task is to convert the following extracted transcript content into a comprehensive, highly accurate, and professionally structured meeting summary document.

---------------------------------------------------------------------
GROUNDING RULES (HIGHEST PRIORITY — TOPIC-LEVEL FIDELITY):

The extracted content decides WHAT topics to cover. Within those topics, write naturally with clear explanation.

ALLOWED — when a topic IS in the extractions:
- Explain it well with natural depth and helpful context.
- Merge related content from multiple section extractions into one unified topic description.
- Use proportional length based on how much was actually extracted.

NOT ALLOWED:
- DO NOT introduce topics the extractions never mention.
- DO NOT fabricate action items. If none in the extractions, OMIT the entire Action Items section (no heading, no placeholder text). Same rule for Examples & Demonstrations and Complete Code Example.
- DO NOT invent examples for the Examples & Demonstrations section. Omit it if no examples are in the extractions.
- DO NOT invent code. Omit Complete Code Example if no code was in the extractions.
- DO NOT correct or modernize tools/commands/APIs — keep as recorded.
- DO NOT add applications, use cases, or industry context the extractions don't mention.

VOICE RULE — CONTENT-FIRST, NOT NARRATION (CRITICAL — APPLIES TO EVERY TOPIC):
Universal rule. Applies whether the meeting is about software, enterprise IT, sales, HR, finance, healthcare, cooking, fitness, history, law, education, or any other subject. Treat the document like a textbook/reference doc, not meeting minutes.

ABSOLUTELY BANNED — never start a sentence with:
"The speaker discussed/explained/demonstrated/showed/emphasized/highlighted/covered/talked about/touched on/mentioned/described/presented/introduced..."
"They discussed/explained/demonstrated/showed/emphasized..."
"The instructor/trainer/presenter/teacher/host discussed/explained..."
"In the meeting/session/class..." or "During the meeting/session/class..."
"The speaker's explanation/discussion/demonstration of..."

REQUIRED — content-first declarative voice everywhere:
- Subject of each sentence = the topic / concept / tool / process / subject matter, NOT the speaker.
- Use teaching verbs: "is", "works by", "requires", "produces", "is calculated by", "begins with", "follows the formula", "depends on".

UNIVERSAL EXAMPLES (pattern is the same in every domain):
- Coding: WRONG "The speaker explained how to create a K-means model." → RIGHT "A K-means model is created with KMeans(n_clusters=k) and fitted via .fit(); SSE is read from the .inertia_ attribute."
- Enterprise IT: WRONG "The speaker demonstrated how to install SAP HANA using HDBLCM." → RIGHT "SAP HANA is installed using HDBLCM, which provides a guided flow covering system parameters, network config, and license setup."
- Business: WRONG "The speaker discussed Q3 revenue and explained the team missed targets." → RIGHT "Q3 revenue came in 8% below target, driven by delayed enterprise renewals and slower mid-market conversion."
- Process/HR: WRONG "The speaker described the recruitment process." → RIGHT "The recruitment process begins with resume screening, followed by a phone interview, technical assessment, and panel interview before final selection."
- General/lifestyle: WRONG "The speaker demonstrated how to make pasta dough." → RIGHT "Pasta dough is made by combining 100g of flour per egg, kneading 10 minutes until smooth, and resting 30 minutes before rolling."
- Healthcare: WRONG "The speaker described how blood pressure should be measured." → RIGHT "Blood pressure should be measured at the same time each day, after 5 minutes of seated rest, with the arm at heart level."

Speaker attribution allowed ONLY in Action Items (task owners) and for specific decisions whose weight depends on the speaker (e.g., "The CFO approved the revised budget."). Everywhere else: omit speaker references entirely.

QUALITY STANDARDS (apply within grounded topics):
- Tone: clear, professional, instructional, like SAP/Oracle/Google/Microsoft/AWS internal documentation.
- Depth: proportional to how much was extracted.
- Consolidate duplicate or fragmented instructions across sections into one. Do not invent missing steps to "complete" a procedure.
- Replace real values with: <ip>, <password>, <hostname>, <port>, <username>, <email>.
---------------------------------------------------------------------

EXTRACTED CONTENT FROM {num_sections} SECTIONS:
{combined_partials}

---------------------------------------------------------------------
OUTPUT FORMAT — Use these EXACT section headers in this EXACT order:
---------------------------------------------------------------------

## Executive Summary
Comprehensive overview of the entire meeting/session. Write as many sentences as needed — do NOT compress.

## Key Topics

### [Topic Name 1]
The topic must come from the extractions. Describe it clearly with proportional depth. You may add natural context and explanation, but do NOT add topics or applications the extractions don't mention.

### [Topic Name 2]
Same — topic must come from the extractions, written with appropriate depth and clarity.

(Add all topics. Merge related content from different sections into unified topics.)

## Action Items
CONDITIONAL: Include ONLY if the speaker explicitly assigned tasks. If none, DO NOT include this section at all — skip the heading entirely. Do NOT write any placeholder text.
- [Who/Group] Action item or task.

## Detailed Discussion

### [Topic Name 1]
- Detailed bullet points with ALL specifics: commands, tools, configurations, procedures, parameters
- Step-by-step instructions if discussed, clearly numbered
- Validation steps, troubleshooting, best practices where applicable
- As many bullets as needed — no compression

### [Topic Name 2]
- Same thorough treatment

(Merge content from multiple sections for same topic. Each topic: no minimum limit, write everything.)

## Examples & Demonstrations
CONDITIONAL: Include ONLY if extracted content contains ACTUAL demonstrations, live examples, code walkthroughs, calculations, hands-on demos. If general meeting with NO demos, DO NOT include this section.

### Example 1: [Short Title]
Full walkthrough with specific inputs, steps, outputs, observations.

## Complete Code Example
DECISION RULE — only two outcomes. NO middle ground.

OUTCOME A: Generate real working code
Trigger if the extractions mention ANY of:
1. Actual code dictated by the speaker
2. ANY programming language (Python, Java, JavaScript, C++, SQL, R, etc.)
3. ANY programming concept (function, class, variable, library, framework, API, etc.)
4. Any algorithm normally implemented in code (K-means, regression, sorting, neural networks, REST API, etc.)

When triggered: ANALYZE the concept actually taught, then GENERATE complete runnable code demonstrating EXACTLY that concept. Use the language mentioned (default Python). Include real imports, real implementation, inline comments, sample data, and expected output as a comment. Not pseudocode.

OUTCOME B: Omit the entire section
Trigger only when extractions have ZERO programming content. Do NOT write the heading. Do NOT write any text.

ABSOLUTELY FORBIDDEN — never write phrases like:
- "No code was discussed"
- "No code was explicitly mentioned but the speaker discussed..."
- "While no code was shown..."
- ANY apology text about absent code

Either generate real working code, or omit the section completely. No third option.

## Key Takeaways
Thorough summary — what was implemented, expected outcomes, best practices, prerequisites. Write proportionally.

## Concept Map
```dot
digraph ConceptMap {{{{
    node[shape=box, style="rounded,filled"];
    root[label="Main Topic"];
    (build proper concept map covering ALL key topics and relationships)
}}}}
```

---

Was this summary helpful? Share your feedback to help us improve.

---------------------------------------------------------------------
FINAL RULES:
---------------------------------------------------------------------
1. Use EXACT section headers: ## Executive Summary, ## Key Topics, ## Action Items, ## Detailed Discussion, ## Examples & Demonstrations, ## Complete Code Example, ## Key Takeaways, ## Concept Map
2. Topic names must match between Key Topics and Detailed Discussion
3. Merge related topics from different sections — no duplicates
4. Examples & Demonstrations: ONLY if real demos exist. Skip entirely for general meetings.
5. Complete Code Example: ONLY if coding detected. Skip entirely otherwise.
6. DO NOT compress content. Long discussions = long sections.
7. Replace real values with <ip>, <password>, <hostname>, <port>, <username>, <email>
8. Concept map must be valid Graphviz DOT format
9. Cross-check tools/commands — correct outdated ones
10. Always end with feedback line
11. If a conditional section (Action Items, Examples & Demonstrations, Complete Code Example) has no content, OMIT it entirely — no heading, no explanation, no placeholder. Do NOT write text like "No code was discussed", "No specific examples", or "No specific action items were identified". Just leave nothing in its place.
"""


def _generate_chunked_summary(transcript: str, context: str = "") -> str:
    """
    Summarize long transcripts (2-4 hour meetings) using chunked approach.
    Same quality and format as single summary.
    """
    words = transcript.split()
    chunk_size = 8000
    transcript_chunks = []
    for i in range(0, len(words), chunk_size):
        chunk_text = " ".join(words[i:i + chunk_size])
        transcript_chunks.append(chunk_text)

    logger.info(f"Split into {len(transcript_chunks)} chunks for summarization")

    # Step 1: Extract key content from each chunk
    partial_summaries = []
    for idx, chunk_text in enumerate(transcript_chunks):
        try:
            logger.info(f"Extracting content from chunk {idx + 1}/{len(transcript_chunks)}")


            partial_prompt = f"""You are extracting information from a transcript section.

TRANSCRIPT SECTION {idx + 1} of {len(transcript_chunks)}:
\"\"\"{chunk_text}\"\"\"

For each category below, extract what is actually in this section. The TOPICS, ITEMS, EXAMPLES, and TASKS must come from the transcript — do not invent them. Within topics that ARE present, you may explain them clearly and add helpful context. If a category has nothing in this section, write "Not present in this section."

1. MAIN TOPICS: List topics actually discussed. Describe each clearly with proportional depth — you may add natural explanation, but the TOPIC ITSELF must come from the speaker.
2. KEY CONCEPTS: Explain each concept the speaker introduced, with helpful clarity.
3. TECHNICAL TERMS: List tools, technologies, APIs, commands, configurations actually mentioned. Do not add tools that were not mentioned.
4. STEP-BY-STEP PROCEDURES: Include ALL steps the speaker actually stated, in order, with all commands, parameters, and paths. Do not invent missing steps, but do not skip any either.
5. EXAMPLES/DEMOS: Include ALL examples the speaker actually performed, with specific input values, steps, and output results. Do not fabricate illustrations.
6. VALIDATION: Include all validation steps the speaker mentioned — expected outputs, checks, indicators.
7. TROUBLESHOOTING: Include all issues the speaker discussed — problems, error codes, corrective actions.
8. BEST PRACTICES: Include all tips, recommendations, efficiency improvements, security tips, SOPs the speaker stated.
9. ACTION ITEMS: Include only tasks the speaker explicitly assigned.
10. CHALLENGES: Include all problems, misconceptions, or pitfalls the speaker mentioned.

LENGTH RULE — proportional, not fixed:
- Output as much as the actual content supports. If the speaker discussed 15 steps, extract all 15.
- Do NOT compress: extract EVERYTHING substantive that the speaker actually said.
- Do NOT pad: do not invent content to reach a target length. If the speaker only said a little, extract only a little.

Replace real values with tags: <ip>, <password>, <hostname>, <port>, <username>, <email>."""
            response = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": "You extract information from the provided transcript. Topics, examples, action items, and procedures must come from the transcript — do not invent them. Within topics that ARE present, you may explain clearly with helpful context. CRITICAL UNIVERSAL VOICE RULE (applies to every topic — software, business, medical, cooking, fitness, education, anything): when you write the extracted content, do NOT use narration phrases like 'The speaker discussed/explained/demonstrated', 'They explained/showed', 'The instructor said', 'In the meeting...'. Instead, write content-first declarative sentences where the topic/concept/process is the subject. Examples — instead of 'The speaker explained how MinMaxScaler normalizes values', write 'MinMaxScaler normalizes values to a 0-1 range using (x - min) / (max - min)'. Instead of 'The speaker discussed the recruitment process', write 'The recruitment process begins with resume screening followed by an interview round'. Instead of 'The speaker demonstrated how to make pasta dough', write 'Pasta dough is made by combining 100g flour per egg and kneading for 10 minutes'. If a category has nothing in the transcript, mark it as not present."},
                    {"role": "user", "content": partial_prompt}
                ],
                temperature=0.3,
                max_tokens=4000
            )

            partial = response.choices[0].message.content.strip()
            partial_summaries.append(partial)
            logger.info(f"Chunk {idx + 1} extracted ({len(partial)} chars)")

        except Exception as chunk_error:
            logger.error(f"[CHUNK SUMMARY ERROR] Chunk {idx + 1}: {chunk_error}")
            partial_summaries.append(f"[Section {idx + 1} extraction failed]")

    # Step 2: Combine
    combined_partials = "\n\n=====================================\n\n".join(
        f"EXTRACTED CONTENT FROM SECTION {i + 1}:\n{p}" for i, p in enumerate(partial_summaries)
    )

    # Step 3: Generate final summary using structured format with full quality rules
    final_prompt = CHUNKED_FINAL_PROMPT_TEMPLATE.format(
        num_sections=len(partial_summaries),
        combined_partials=combined_partials
    )

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": "You merge transcript extractions into a final content-first reference summary (like SAP/Oracle/Google internal docs), NOT meeting minutes. The TOPICS, action items, and examples must come from the extractions — never invent NEW ones. CRITICAL UNIVERSAL VOICE RULE (applies to every topic — software, enterprise, business, healthcare, cooking, fitness, education, anything): NEVER start sentences with 'The speaker discussed/explained/demonstrated', 'They explained/showed', 'The instructor/trainer/presenter...', 'In the meeting/session/class...', or 'During the meeting/session...'. Write content-first declarative sentences where the SUBJECT is the topic/concept/tool/process, not the speaker. Use 'is/works by/requires/produces/is calculated by/begins with' instead of 'the speaker explained/showed/demonstrated'. Speaker attribution allowed ONLY in Action Items (task owners) and for specific decisions whose weight depends on the speaker. SPECIAL RULE for Complete Code Example: if the extractions mention any programming language, concept, or algorithm normally implemented in code, GENERATE complete working code demonstrating that exact concept. If extractions have zero programming content, omit the Code section entirely. NEVER write apology text like 'No code was discussed but...' — produce real code or produce nothing. Same omit-or-nothing rule for Action Items and Examples."},
                {"role": "user", "content": final_prompt}
            ],
            temperature=0.4,
            max_tokens=8000
        )
        final_summary = response.choices[0].message.content.strip()
        logger.info(f"Final combined summary generated ({len(final_summary)} chars)")
        return final_summary
    except Exception as e:
        logger.error(f"[FINAL SUMMARY ERROR] {e}")
        return "\n\n".join(partial_summaries)
        
def split_text_into_timed_segments(text: str, chunk_start: float, chunk_duration: float) -> list:
    """
    Split a block of text into subtitle-sized segments with estimated timing.
    Used because gpt-4o-transcribe doesn't return per-segment timestamps.
    Distributes time proportionally based on character count.
    """
    import re
    
    # Split into sentences
    sentences = re.split(r'(?<=[.!?।])\s+', text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]
    
    if not sentences:
        return [{
            "start": chunk_start,
            "end": chunk_start + chunk_duration,
            "text": text
        }]
    
    # Break long sentences (>80 chars) into subtitle-friendly pieces
    refined_segments = []
    for sentence in sentences:
        if len(sentence) <= 80:
            refined_segments.append(sentence)
        else:
            parts = re.split(r',\s+', sentence)
            current = ""
            for part in parts:
                if len(current) + len(part) <= 80:
                    current = (current + ", " + part) if current else part
                else:
                    if current:
                        refined_segments.append(current)
                    current = part
            if current:
                refined_segments.append(current)
    
    # Calculate timing proportional to character count
    total_chars = sum(len(s) for s in refined_segments)
    if total_chars == 0:
        return []
    
    timed_segments = []
    current_time = chunk_start
    
    for segment_text in refined_segments:
        segment_duration = (len(segment_text) / total_chars) * chunk_duration
        segment_duration = max(1.0, min(6.0, segment_duration))
        
        timed_segments.append({
            "start": round(current_time, 2),
            "end": round(current_time + segment_duration, 2),
            "text": segment_text
        })
        current_time += segment_duration
    
    # Ensure last segment ends at chunk end
    if timed_segments:
        timed_segments[-1]["end"] = round(chunk_start + chunk_duration, 2)
    
    return timed_segments

def translate_segments_to_target_lang(segments: list, target_lang: str) -> list:
    """
    Translate English segments to target language (Hindi, Telugu, etc.) using Groq Llama.
    Fast, reliable, preserves technical terms. Used for subtitle generation.
    
    CRITICAL: This function translates ENGLISH → TARGET LANGUAGE (Hindi/Telugu/etc.)
    Opposite direction from translate_segments_batch which goes TO English.
    """
    if not segments or target_lang == "en":
        return segments
    
    if groq_client is None:
        logger.warning(f"groq_client is None, cannot translate to {target_lang}")
        return segments
    
    # Language code to full name mapping
    LANG_NAMES = {
        "hi": "Hindi",
        "te": "Telugu",
        "ta": "Tamil",
        "kn": "Kannada",
        "ml": "Malayalam",
        "bn": "Bengali",
        "mr": "Marathi",
        "gu": "Gujarati",
        "pa": "Punjabi",
    }
    
    # Script name for validation
    SCRIPT_NAMES = {
        "hi": "Devanagari",
        "te": "Telugu",
        "ta": "Tamil",
        "kn": "Kannada",
        "ml": "Malayalam",
        "bn": "Bengali",
        "mr": "Devanagari",
        "gu": "Gujarati",
        "pa": "Gurmukhi",
    }
    
    target_lang_name = LANG_NAMES.get(target_lang, target_lang)
    target_script = SCRIPT_NAMES.get(target_lang, target_lang)
    
    BATCH_SIZE = 30
    translated_all = []
    
    for batch_start in range(0, len(segments), BATCH_SIZE):
        batch = segments[batch_start:batch_start + BATCH_SIZE]
        
        try:
            numbered = "\n".join(f"[{idx}] {s['text']}" for idx, s in enumerate(batch))
            
            response = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {
                        "role": "system",
                        "content": (
                            f"You are a professional subtitle translator. "
                            f"You translate English text into {target_lang_name} written in {target_script} script.\n\n"
                            f"STRICT RULES:\n"
                            f"1. Translate every numbered English line into natural, conversational {target_lang_name}.\n"
                            f"2. Output MUST be in {target_script} script (native writing system of {target_lang_name}).\n"
                            f"3. DO NOT output English text. DO NOT output Latin/Roman letters (except for technical terms in rule 5).\n"
                            f"4. Use simple, everyday {target_lang_name} suitable for video subtitles (max ~80 characters per line).\n"
                            f"5. Keep technical English terms UNCHANGED: Python, Java, JavaScript, C++, SQL, React, Node.js, AWS, Azure, GCP, SAP, Oracle, Docker, Kubernetes, Jenkins, Linux, Selenium, API, CPU, IDE, variable, integer, float, string, boolean, dictionary, function, class, loop, array, database.\n"
                            f"6. Keep product names UNCHANGED: Windows, Linux, macOS, Chrome, Firefox, SAP HANA, Salesforce.\n"
                            f"7. Output ONLY numbered lines in EXACT format:\n"
                            f"   [0] {target_lang_name} translation in {target_script} script\n"
                            f"   [1] {target_lang_name} translation in {target_script} script\n"
                            f"   [2] {target_lang_name} translation in {target_script} script\n"
                            f"8. No preamble, no explanation, no markdown, no quotes, no English prefixes.\n\n"
                            f"EXAMPLE (for Hindi):\n"
                            f"Input:  [0] Variable name is case sensitive\n"
                            f"Output: [0] Variable का नाम case sensitive होता है\n\n"
                            f"EXAMPLE (for Telugu):\n"
                            f"Input:  [0] Python is a programming language\n"
                            f"Output: [0] Python ఒక programming language."
                        )
                    },
                    {
                        "role": "user",
                        "content": (
                            f"Translate these English lines into {target_lang_name} ({target_script} script).\n"
                            f"Remember: Output MUST be in {target_script} script, NOT in English letters.\n\n"
                            f"{numbered}"
                        )
                    }
                ],
                temperature=0.3,
                max_tokens=4000
            )
            
            translated_text = response.choices[0].message.content.strip()
            
            # Parse back into map: index -> translated text
            translated_map = {}
            for line in translated_text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                m = re.match(r"\[(\d+)\]\s*(.*)", line)
                if m:
                    translated_map[int(m.group(1))] = m.group(2).strip()
            
            # Apply translations, keep original timing
            for idx, seg in enumerate(batch):
                translated_seg_text = translated_map.get(idx, seg["text"])
                translated_all.append({
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": translated_seg_text if translated_seg_text else seg["text"]
                })
            
            logger.info(f"[SUBTITLE {target_lang}] Batch {batch_start // BATCH_SIZE + 1}: translated {len(translated_map)}/{len(batch)} segments to {target_script} script")
            
        except Exception as batch_error:
            logger.error(f"[SUBTITLE {target_lang}] Batch translation failed: {batch_error}")
            for seg in batch:
                translated_all.append({
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": seg["text"]
                })
    
    logger.info(f"[SUBTITLE {target_lang}] Completed: {len(translated_all)}/{len(segments)} segments translated to {target_lang_name}")
    return translated_all

def clean_garbage_scripts(text: str) -> str:
    """
    Remove ONLY true garbage/control characters.
    ALL real languages (Chinese, Japanese, Korean, Arabic, Russian, etc.)
    are kept and will be translated to English by GPT-4o-mini.
    """
    if not text:
        return text
    
    # Remove ONLY non-language garbage (control chars, private use, symbols)
    garbage_ranges = [
        (0x0000, 0x001F),   # Control characters
        (0x007F, 0x009F),   # More control characters
        (0xFFF0, 0xFFFF),   # Specials
        (0xE000, 0xF8FF),   # Private Use Area (garbage symbols)
        (0xFE00, 0xFE0F),   # Variation Selectors
    ]
    
    def is_garbage(char):
        code = ord(char)
        for start, end in garbage_ranges:
            if start <= code <= end:
                return True
        return False
    
    cleaned = "".join(" " if is_garbage(char) else char for char in text)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def remove_repetitions(text: str) -> str:
    """
    Remove repeated phrases/sentences that Whisper hallucinates on silence
    or low-information audio (e.g., screen-share with no narration).

    Strategy:
    - 5+ identical consecutive sentences = hallucination → DROP ALL copies.
      Keeping even one copy contaminates the transcript with orphan stubs
      that concatenate into nonsense paragraphs across chunks.
    - 2–4 identical consecutive sentences = normal speech emphasis → keep 1 copy.
    """
    if not text or len(text) < 50:
        return text

    # PASS 1: Sentence-level repetition handling
    sentences = re.split(r'(?<=[.!?])\s+', text)
    if not sentences:
        return text

    deduped_sentences = []
    i = 0
    while i < len(sentences):
        current = sentences[i]
        current_norm = re.sub(r'\s+', ' ', current.strip().lower())

        # Count consecutive identical sentences starting at i
        run_length = 1
        j = i + 1
        while j < len(sentences):
            next_norm = re.sub(r'\s+', ' ', sentences[j].strip().lower())
            if next_norm == current_norm and current_norm:
                run_length += 1
                j += 1
            else:
                break

        if run_length >= 5:
            # Hallucination signature — drop ALL copies entirely
            logger.warning(
                f"[REPEAT FILTER] Hallucination — dropping all {run_length} "
                f"copies of: '{current[:60]}...'"
            )
        elif run_length >= 2:
            # Normal speech emphasis — keep one copy
            deduped_sentences.append(current)
            for _ in range(run_length - 1):
                logger.warning(
                    f"[REPEAT FILTER] Removed duplicate sentence: '{current[:60]}...'"
                )
        else:
            deduped_sentences.append(current)

        i = j

    text = " ".join(deduped_sentences)

    # PASS 2: Intra-sentence phrase repetition (unchanged behavior for 2-4 repeats,
    # but 5+ repeats of a phrase are now also dropped entirely)
    for word_count in [8, 7, 6, 5, 4, 3]:
        word_pattern = r'\b(' + r'\S+\s+' * (word_count - 1) + r'\S+)' + r'(?:\s+\1){2,}'
        matches = list(re.finditer(word_pattern, text, re.IGNORECASE))
        for match in matches:
            repeated_phrase = match.group(1)
            full_match = match.group(0)
            repeat_count = len(full_match.split(repeated_phrase))
            if repeat_count >= 5:
                logger.warning(
                    f"[REPEAT FILTER] Phrase hallucination — dropping all "
                    f"{repeat_count} copies of: '{repeated_phrase[:50]}...'"
                )
                text = text.replace(full_match, "")
            else:
                logger.warning(
                    f"[REPEAT FILTER] Removed repeated phrase: '{repeated_phrase[:50]}...' "
                    f"(appeared {repeat_count} times)"
                )
                text = text.replace(full_match, repeated_phrase)

    return text.strip()

def translate_segments_batch(segments: list, source_lang: str) -> list:
    """
    Translate a batch of transcript segments to clean English using LLM.
    Handles Indic languages (Tamil, Hindi, Telugu, Kannada, etc.) and code-switched audio.
    
    IMPROVED v2:
    - Batch size reduced to 15 for better Tamil/Indic reliability
    - 3 retries with individual segment retry on failure
    - Stricter 3% non-ASCII threshold
    - Explicit Tamil script detection and examples in prompt
    - Indic script detection helper for targeted retries
    """
    if not segments:
        return segments
    
    if groq_client is None:
        logger.warning("groq_client is None, cannot translate - returning original")
        return segments
    
    BATCH_SIZE = 15  # Smaller batches = more reliable for Tamil/Indic
    MAX_RETRIES = 3
    translated_all = []
    
    def _has_indic_content(text: str) -> bool:
        """Check if text contains ANY Indic script characters (Tamil, Hindi, Telugu, etc.)."""
        for char in text:
            code = ord(char)
            if (0x0900 <= code <= 0x097F or   # Devanagari (Hindi, Marathi)
                0x0980 <= code <= 0x09FF or   # Bengali
                0x0A00 <= code <= 0x0A7F or   # Gurmukhi (Punjabi)
                0x0A80 <= code <= 0x0AFF or   # Gujarati
                0x0B00 <= code <= 0x0B7F or   # Oriya
                0x0B80 <= code <= 0x0BFF or   # Tamil
                0x0C00 <= code <= 0x0C7F or   # Telugu
                0x0C80 <= code <= 0x0CFF or   # Kannada
                0x0D00 <= code <= 0x0D7F):    # Malayalam
                return True
        return False
    

    def _call_llm(batch_segments, retry_count=0):
        """Single LLM call for a batch. Returns translated_map {idx: text}."""
        numbered = "\n".join(f"[{idx}] {s['text']}" for idx, s in enumerate(batch_segments))
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a translator. Convert ALL input to English using Latin alphabet ONLY.\n\n"
                        "RULES:\n"
                        "1. Output 100% English. ZERO Devanagari/Telugu/Tamil/Kannada/Bengali/Malayalam allowed.\n"
                        "2. Translate MEANING, not sound. 'है' → 'is', NOT 'hai'. 'चेयाली' → 'should do', NOT 'cheyali'.\n"
                        "3. Whisper writes Telugu speech in Devanagari. Common Telugu-in-Devanagari words:\n"
                        "   गद=right?, मनम/मनकी=we/our, चेयाली/चेसको=should do, उन्टे=if there is,\n"
                        "   राइसेयाम=let us write, अलागे=similarly, तरवात=after that, इप्पुडु=now,\n"
                        "   एक्सिकुट=execute, दानिको=for that, वेदुन्दी=need to know, लो=in, नी=of\n"
                        "4. Hindi words: है=is, और=and, में=in, को=to, का/की=of, यह=this, हम=we,\n"
                        "   पहले=first, अब=now, कैसे=how, करना=to do, देखो=see\n"
                        "5. Romanized: hai→is, cheyali→should do, ante→means, manam→we, ippudu→now\n"
                        "6. Keep English technical terms unchanged: install, Linux, command, server, IP, ping, etc.\n"
                        "7. If text is garbled nonsense (not real English or any language), return ONLY the real words.\n"
                        "8. One output line per input line. Format: [0] english text\n\n"
                        "EXAMPLES:\n"
                        "[0] command line bash bash वा ने नुवु नी virtual machine turn on चेसीन\n"
                        "[1] modules उन्ने गद, modules so, इ modules लो, मनकी functions उन्टे\n"
                        "[2] math.seal न राइसेयाम। अलागे math.floor न राइसेयाम।\n"
                        "[3] So NumPy is complete working on arrays\n"
                        "[4] तरवात ये version use चेस्थ नम्मु चेएड़ान की दानिको वेदुन्दी\n"
                        "[5] Patient ko din mein do baar medicine deni hai\n"
                        "[6] इप्पड़ु इदी मनकी 3-Dimension format लो उस्तों\n"
                        "[7] haberdigon install the operating system gesht\n"
                        "[8] பார்க்கலாம் இந்த formula use பண்ணணும்\n"
                        "[9] ఈ test case execute చేయాలి automation framework లో\n"
                        "[10] এই function টা database থেকে data fetch করবে\n"
                        "[11] हा concept मला नीट समजला नाही, please explain\n"
                        "[12] આ table માં primary key add કરવાની છે\n"
                        "[13] vamos a configurar el servidor con docker ahora\n"
                        "[14] 我们 now need to deploy 这个 application\n"
                        "[15] هذا الكود يحتاج إلى debugging قبل النشر\n\n"
                        "OUTPUT:\n"
                        "[0] In command line bash, you need to turn on the virtual machine\n"
                        "[1] There are modules, in these modules we have functions\n"
                        "[2] Let us write math.ceil. Similarly let us write math.floor.\n"
                        "[3] So NumPy is complete working on arrays\n"
                        "[4] After that we use this version for what we need to know\n"
                        "[5] Give the patient medicine twice a day\n"
                        "[6] Now this is in 3-Dimension format for us\n"
                        "[7] Install the operating system\n"
                        "[8] Let us see, we need to use this formula\n"
                        "[9] We need to execute this test case in the automation framework\n"
                        "[10] This function will fetch data from the database\n"
                        "[11] I did not understand this concept properly, please explain\n"
                        "[12] We need to add a primary key in this table\n"
                        "[13] Now we are going to configure the server with docker\n"
                        "[14] We now need to deploy this application\n"
                        "[15] This code needs debugging before deployment"
                    )
                },
                {
                    "role": "user",
                    "content": (
                        f"Translate {len(batch_segments)} lines to English. Remove garbled nonsense. Keep technical terms.\n\n"
                        f"{numbered}"
                    )
                }
            ],
            temperature=0.1,
            max_tokens=4000
        )
        
        translated_text = response.choices[0].message.content.strip()
        
        # Parse numbered output
        t_map = {}
        for line in translated_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            m = re.match(r"\[(\d+)\]\s*(.*)", line)
            if m:
                idx = int(m.group(1))
                text = m.group(2).strip()
                if text:
                    t_map[idx] = text
        return t_map    
    def _is_clean_english(text: str) -> bool:
        """Empty/missing translations are NOT clean — they need retry."""
        if not text or not text.strip():
            return False
        non_ascii = sum(1 for c in text if ord(c) > 127)
        return (non_ascii / len(text)) < 0.03
    
    # Process segments in smaller batches
    for batch_start in range(0, len(segments), BATCH_SIZE):
        batch = segments[batch_start:batch_start + BATCH_SIZE]
        
        try:
            # First attempt — batch translation
            translated_map = _call_llm(batch)
            
            # Retry loop — check and re-translate segments that are still non-English
            for retry_round in range(MAX_RETRIES):
                missing_or_bad = []
                for idx, seg in enumerate(batch):
                    text = translated_map.get(idx, "")
                    if not text or not _is_clean_english(text) or _has_indic_content(text):
                        missing_or_bad.append(idx)
                
                if not missing_or_bad:
                    break  # All segments are clean English
                
                logger.warning(f"[TRANSLATE] Batch {batch_start}, retry {retry_round+1}: {len(missing_or_bad)} segments still non-English")
                
                if retry_round == 0:
                    # First retry: try as a smaller sub-batch
                    retry_batch = [batch[idx] for idx in missing_or_bad]
                    retry_map = _call_llm(retry_batch)
                    for retry_idx, original_idx in enumerate(missing_or_bad):
                        if retry_idx in retry_map and _is_clean_english(retry_map[retry_idx]) and not _has_indic_content(retry_map[retry_idx]):
                            translated_map[original_idx] = retry_map[retry_idx]
                else:
                    # Subsequent retries: translate individually for maximum accuracy
                    for bad_idx in missing_or_bad:
                        try:
                            single_result = _call_llm([batch[bad_idx]])
                            if 0 in single_result and _is_clean_english(single_result[0]) and not _has_indic_content(single_result[0]):
                                translated_map[bad_idx] = single_result[0]
                                logger.info(f"[TRANSLATE] Individual retry SUCCESS for segment {bad_idx}")
                        except Exception as single_err:
                            logger.warning(f"[TRANSLATE] Individual retry FAILED for segment {bad_idx}: {single_err}")
            
            # Build translated segments
            for idx, seg in enumerate(batch):
                translated_text = translated_map.get(idx, "")
                
                # If still no valid translation, force-strip non-ASCII from original
                if not translated_text or not _is_clean_english(translated_text) or _has_indic_content(translated_text):
                    # Try stripping from translated text first, then original
                    source_text = translated_text if translated_text else seg["text"]
                    stripped = "".join(c if ord(c) < 128 else " " for c in source_text)
                    stripped = re.sub(r'\s+', ' ', stripped).strip()
                    
                    if stripped and len(stripped) > 5:
                        translated_text = stripped
                        logger.warning(f"[TRANSLATE] Segment {idx}: LLM failed, used ASCII-stripped fallback")
                    else:
                        # Not enough English content to keep - skip this segment
                        translated_text = ""
                        logger.warning(f"[TRANSLATE] Segment {idx}: Dropped (mostly non-English, no salvageable content)")
                
                # Only keep segments with real content
                if translated_text and len(translated_text.strip()) > 2:
                    translated_all.append({
                        "start": seg["start"],
                        "end": seg["end"],
                        "text": translated_text
                    })
            
            logger.info(f"[TRANSLATE] Batch {batch_start // BATCH_SIZE + 1}: kept {len([s for s in translated_all if s])}/{len(batch)} segments")
            
        except Exception as e:
            logger.error(f"[TRANSLATE ERROR] Batch {batch_start}: {e}")
            # On total failure, try to salvage ASCII content from original
            for seg in batch:
                original = seg["text"]
                stripped = "".join(c if ord(c) < 128 else " " for c in original)
                stripped = re.sub(r'\s+', ' ', stripped).strip()
                if stripped and len(stripped) > 5:
                    translated_all.append({
                        "start": seg["start"],
                        "end": seg["end"],
                        "text": stripped
                    })
    
    logger.info(f"[TRANSLATE] Complete: {len(translated_all)} clean English segments from {len(segments)} input")
    return translated_all

def analyze_trainer_performance(transcript: str) -> dict:
    """
    Analyze trainer's technical content, explanation clarity, friendliness, and communication
    based purely on transcript semantics and delivery style markers.
    Uses GPT-4o for text-based behavioral and technical scoring.
    """
    if not transcript.strip():
        return {
            "technical_content": 0,
            "explanation_clarity": 0,
            "friendliness": 0,
            "communication": 0,
            "overall_feedback": "No speech detected for evaluation."
        }
    
    # ✅ CRITICAL: Check if Groq client is available
    if groq_client is None:
        logger.error("❌ Cannot evaluate trainer - groq_client is None!")
        logger.error(f"❌ GROQ_API_KEY value: '{os.getenv('GROQ_API_KEY', 'NOT_SET')[:10]}...'")
        return {
            "technical_content": 0,
            "explanation_clarity": 0,
            "friendliness": 0,
            "communication": 0,
            "overall_feedback": "Evaluation unavailable - Groq API not configured.",
            "error": "groq_client is None"
        }

    # Truncate transcript if too long (trainer eval doesn't need full text)
    words = transcript.split()
    MAX_WORDS_FOR_EVAL = 8000  # About 30 mins of speech — sufficient for style evaluation
    
    if len(words) > MAX_WORDS_FOR_EVAL:
        # Sample from beginning, middle, and end for representative evaluation
        third = MAX_WORDS_FOR_EVAL // 3
        beginning = " ".join(words[:third])
        middle_start = (len(words) // 2) - (third // 2)
        middle = " ".join(words[middle_start:middle_start + third])
        end = " ".join(words[-third:])
        sampled_transcript = f"{beginning}\n\n[... middle section ...]\n\n{middle}\n\n[... later section ...]\n\n{end}"
        logger.info(f"📊 Trainer eval: sampled {MAX_WORDS_FOR_EVAL} words from {len(words)} total words")
    else:
        sampled_transcript = transcript
    
    prompt = f"""
You are an expert communication and training evaluator.
Evaluate the trainer's communication quality, tone, and content in the transcript below.

TRANSCRIPT:
\"\"\"{sampled_transcript}\"\"\"

Evaluate across the following dimensions (each scored 0–100%):
1. Technical Content — accuracy, depth, and domain clarity.
2. Explanation Clarity — how logically and simply ideas are explained.
3. Friendliness — warmth, politeness, and positive tone.
4. Communication – evaluate using Indian English standards:
   - Focus on fluency, confidence, and comfort with Indian accent.
   - Accept light Indianisms such as “basically”, “ok na”, “ya”, etc.
   - Penalize only unclear speech or excessive filler use (“umm”, “like”, “you know”).
   - Do not reduce marks for accent style — evaluate clarity, not foreign pronunciation.
Output a short, factual JSON report with numeric scores and 1–2 lines of feedback.
Format exactly as:

{{
  "technical_content": <number>,
  "explanation_clarity": <number>,
  "friendliness": <number>,
  "communication": <number>,
  "overall_feedback": "<short summary>"
}}
"""
    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional behavioral and technical communication evaluator."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.2,
            max_tokens=400
        )

        raw = response.choices[0].message.content.strip()

        # Safe JSON extraction
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if not match:
            logger.warning("Trainer evaluation JSON not found")
            return {"error": "Invalid evaluation format"}

        return json.loads(match.group(0))

    except Exception as e:
        logger.error(f"[GROQ TRAINER EVAL ERROR] {e}")
        return {"error": "Trainer evaluation failed"}

# === ENHANCED VIDEO PROCESSING WITH FASTAPI FEATURES ===
# Add these imports at the very top of your file
try:
    import torch
except ImportError:
    torch = None

try:
    from transformers import MarianMTModel, MarianTokenizer
except ImportError:
    MarianMTModel = None
    MarianTokenizer = None
import logging

def get_transformers():
    """Lazy loader for transformers - returns (None, None) if not installed."""
    try:
        from transformers import MarianMTModel, MarianTokenizer
        return MarianMTModel, MarianTokenizer
    except ImportError:
        return None, None

class LocalIndianLanguageTranslator:
    """Fast local translation for Hindi and Telugu using Helsinki-NLP models"""
    
    def __init__(self):
        self.models = {}
        self.tokenizers = {}
        torch = get_torch()
        self.device = "cuda" if torch and torch.cuda.is_available() else "cpu"
        
        # Model names for Indian languages
        self.model_names = {
            "hi": "Helsinki-NLP/opus-mt-en-hi",  # English to Hindi
            # "te": "Helsinki-NLP/opus-mt-en-mul",  # Use Hindi model for Telugu (fallback)
        }
        
        logging.info(f"🔧 LocalTranslator initialized on device: {self.device}")
    
    def load_model(self, lang):
        """Lazy load model only when needed"""
        if lang in self.models:
            return  # Already loaded
        
        try:
            model_name = self.model_names.get(lang)
            if not model_name:
                raise Exception(f"No model available for {lang}")
            
            logging.info(f"📥 Loading {lang} translation model: {model_name}")
            
            MarianMTModel, MarianTokenizer = get_transformers()
            if not MarianMTModel or not MarianTokenizer:
                raise Exception("Transformers library not available")
            
            self.tokenizers[lang] = MarianTokenizer.from_pretrained(model_name)
            self.models[lang] = MarianMTModel.from_pretrained(model_name)
            
            # Move to GPU if available
            if self.device == "cuda":
                self.models[lang] = self.models[lang].cuda()
                logging.info(f"✅ {lang} model loaded on GPU")
            else:
                logging.info(f"✅ {lang} model loaded on CPU")
                
        except Exception as e:
            logging.error(f"❌ Failed to load {lang} model: {e}")
            raise
    
    def translate_batch(self, texts, target_lang):
        """Translate a batch of texts to target language"""
        if target_lang not in self.model_names:
            raise Exception(f"Unsupported language: {target_lang}")
        
        # Load model if not already loaded
        self.load_model(target_lang)
        
        try:
            # Tokenize input texts
            inputs = self.tokenizers[target_lang](
                texts, 
                return_tensors="pt", 
                padding=True, 
                truncation=True, 
                max_length=512
            )
            
            # Move to GPU if available
            if self.device == "cuda":
                inputs = {k: v.cuda() for k, v in inputs.items()}
            
            # Generate translations
            # Generate translations
            torch = get_torch()
            with torch.no_grad():  # Disable gradient calculation for inference
                outputs = self.models[target_lang].generate(
                    **inputs,
                    max_length=512,
                    num_beams=4,  # Better quality
                    early_stopping=True
                )
            
            # Decode translations
            translations = self.tokenizers[target_lang].batch_decode(
                outputs, 
                skip_special_tokens=True
            )
            
            return translations
            
        except Exception as e:
            logging.error(f"Translation error for {target_lang}: {e}")
            # Return fallback translations
            return ["[Translation unavailable]"] * len(texts)
    
    def translate_segments(self, segments, target_lang, batch_size=32):
        """Translate all segments for a language in batches"""
        if target_lang == "en":
            return segments  # No translation needed
        
        logging.info(f"🌐 Translating {len(segments)} segments to {target_lang} using local model...")
        
        translated_segments = []
        total_batches = (len(segments) + batch_size - 1) // batch_size
        
        for batch_idx in range(0, len(segments), batch_size):
            batch_segments = segments[batch_idx:batch_idx + batch_size]
            batch_texts = [seg["text"].strip() for seg in batch_segments]
            
            # Translate batch
            translations = self.translate_batch(batch_texts, target_lang)
            
            # Create translated segments
            for i, seg in enumerate(batch_segments):
                translated_segments.append({
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": translations[i] if i < len(translations) else "[Translation failed]"
                })
            
            # Progress logging
            current_batch = (batch_idx // batch_size) + 1
            if current_batch % 5 == 0 or current_batch == total_batches:
                progress = (len(translated_segments) / len(segments)) * 100
                logging.info(f"   {target_lang}: {progress:.1f}% ({len(translated_segments)}/{len(segments)} segments)")
        
        return translated_segments
    
    def cleanup(self):
        """Free GPU memory"""
        for model in self.models.values():
            del model
        self.models.clear()
        self.tokenizers.clear()
        
        torch = get_torch()
        if torch and torch.cuda.is_available():
            torch.cuda.empty_cache()

def process_video_sync(video_path: str, meeting_id: str, user_id: str):
    """Process video with GPU acceleration - UPDATED WITH GROQ + DOCX + TRAINER EVAL"""
    import logging
    import subprocess
    from tempfile import TemporaryDirectory
    import os
    import json
    import re
    from datetime import datetime

    logging.info(f"🎬 Starting video processing: {video_path}")
    
    with TemporaryDirectory() as workdir:
        try:
            # Disk space check - need at least 3x video size free
            import shutil as _shutil
            disk_stat = _shutil.disk_usage(workdir)
            free_gb = disk_stat.free / (1024**3)
            logging.info(f"💾 Free disk in workdir: {free_gb:.2f} GB")
            if free_gb < 5:
                raise Exception(f"Insufficient disk space: {free_gb:.2f} GB free (need at least 5 GB)")
            
            # Check if video_path is an S3 key (not a local file)
            if not os.path.exists(video_path) and '/' in video_path:
                # Download from S3 to workdir
                local_video = os.path.join(workdir, os.path.basename(video_path))
                logging.info(f"📥 Downloading from S3: {video_path}")
                if not download_from_s3(video_path, local_video):
                    raise Exception(f"Failed to download from S3: {video_path}")
                video_path = local_video
                logging.info(f"✅ Downloaded to: {video_path}")
            
            if not os.path.exists(video_path):
                raise Exception(f"Input video file not found: {video_path}")
            input_size = os.path.getsize(video_path)
            logging.info(f"📁 Input file size: {input_size} bytes")
            
            if input_size == 0:
                raise Exception("Input video file is empty")
            
            input_ext = os.path.splitext(video_path)[1].lower()
            logging.info(f"📹 Processing {input_ext} file: {video_path}")
            
            # Extract session_id early so it's available for subtitle/transcript uploads
            session_id = None
            basename = os.path.basename(video_path).replace('.mp4', '')
            parts = basename.split('_')
            # Format: recording_meetingid_sessionid_timestamp.mp4
            for part in parts:
                if len(part) == 8 and part.isalnum() and not part.isdigit():
                    session_id = part
                    break
            logging.info(f"📋 Session ID: {session_id}")
            
            # ========== PROBE INPUT FILE ==========
            compressed = os.path.join(workdir, "compressed.mp4")
            
            probe_cmd = [
                "ffprobe", "-v", "quiet", "-print_format", "json", 
                "-show_streams", "-show_format", video_path
            ]
            
            try:
                probe_result = subprocess.run(probe_cmd, capture_output=True, text=True, check=True, timeout=30)
                streams_info = json.loads(probe_result.stdout)
                
                has_audio = any(stream.get('codec_type') == 'audio' for stream in streams_info.get('streams', []))
                has_video = any(stream.get('codec_type') == 'video' for stream in streams_info.get('streams', []))
                video_duration = float(streams_info.get('format', {}).get('duration', 0))
                
                video_stream = next((s for s in streams_info.get('streams', []) if s.get('codec_type') == 'video'), {})
                audio_stream = next((s for s in streams_info.get('streams', []) if s.get('codec_type') == 'audio'), {})
                
                logging.info(f"📊 File analysis: has_video={has_video}, has_audio={has_audio}, duration={video_duration:.2f}s")
                
                if video_stream:
                    logging.info(f"📺 Video: {video_stream.get('codec_name', 'unknown')} {video_stream.get('width', 0)}x{video_stream.get('height', 0)}")
                if audio_stream:
                    logging.info(f"🔊 Audio: {audio_stream.get('codec_name', 'unknown')} {audio_stream.get('sample_rate', 0)}Hz")
                
            except Exception as probe_error:
                logging.warning(f"⚠ Failed to probe input file: {probe_error}")
                has_audio = True
                has_video = True
                video_duration = 0
            
            if not has_video:
                raise Exception("Input file does not contain a valid video stream")
            
            if video_duration <= 0:
                logging.warning(f"⚠ Duration detection failed")
                video_duration = 30.0

            # ========== GPU DETECTION ==========
            nvenc_available = False
            try:
                check_nvenc = subprocess.run(
                    ['ffmpeg', '-h', 'encoder=h264_nvenc'],
                    capture_output=True, text=True, timeout=5
                )
                nvenc_available = (check_nvenc.returncode == 0)
                logging.info(f"{'🚀 GPU (NVENC) detected' if nvenc_available else 'ℹ️ GPU not available - using CPU'}")
            except:
                nvenc_available = False
            
            # ========== VIDEO COMPRESSION ==========
            skip_compression = False
            
            if "_final.mp4" in video_path:
                logging.info("✅ Input is already optimized - skipping compression")
                compressed = video_path
                skip_compression = True
            else:
                logging.info(f"🔄 {'Optimizing' if input_ext == '.mp4' else 'Converting'} video...")
                
                if nvenc_available:
                    if has_audio:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-c:v", "h264_nvenc", "-preset", "p1", "-tune", "hq",
                            "-rc", "vbr", "-cq", "23", "-b:v", "5M",
                            "-maxrate", "10M", "-bufsize", "20M",
                            "-c:a", "aac" if input_ext == '.webm' else "copy",
                            "-ar", "44100", "-ac", "2", "-b:a", "192k",
                            "-movflags", "+faststart", "-pix_fmt", "yuv420p",
                            "-avoid_negative_ts", "make_zero", compressed
                        ]
                    else:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-f", "lavfi", "-i", "anullsrc=channel_layout=stereo:sample_rate=44100",
                            "-c:v", "h264_nvenc", "-preset", "p1", "-tune", "hq",
                            "-rc", "vbr", "-cq", "23", "-b:v", "3M",
                            "-maxrate", "5M", "-bufsize", "6M",
                            "-c:a", "aac", "-ar", "44100", "-ac", "2", "-b:a", "64k",
                            "-shortest", "-movflags", "+faststart",
                            "-pix_fmt", "yuv420p", "-avoid_negative_ts", "make_zero",
                            compressed
                        ]
                else:
                    # CPU fallback
                    if has_audio:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-c:v", "libx264", "-preset", "fast", "-crf", "23",
                            "-maxrate", "10M", "-bufsize", "20M",
                            "-c:a", "aac" if input_ext == '.webm' else "copy",
                            "-ar", "44100", "-ac", "2", "-b:a", "192k",
                            "-movflags", "+faststart", "-pix_fmt", "yuv420p",
                            "-avoid_negative_ts", "make_zero", compressed
                        ]
                    else:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-f", "lavfi", "-i", "anullsrc=channel_layout=stereo:sample_rate=44100",
                            "-c:v", "libx264", "-preset", "fast", "-crf", "23",
                            "-maxrate", "3M", "-bufsize", "6M",
                            "-c:a", "aac", "-ar", "44100", "-ac", "2", "-b:a", "64k",
                            "-shortest", "-movflags", "+faststart",
                            "-pix_fmt", "yuv420p", "-avoid_negative_ts", "make_zero",
                            compressed
                        ]

            if not skip_compression:
                try:
                    # Dynamic timeout: 3 sec per MB, min 10 min, max 2 hours
                    input_size_mb = input_size / (1024 * 1024)
                    # Increased ceiling to 6 hours for long training videos
                    compression_timeout = max(600, min(21600, int(input_size_mb * 5)))
                    logging.info(f"🕐 Video compression timeout: {compression_timeout}s (input: {input_size_mb:.1f}MB)")
                    
                    logging.info(f"🔄 Running compression...")
                    subprocess.run(ffmpeg_cmd, check=True, capture_output=True, text=True, timeout=compression_timeout)
                    logging.info(f"✅ Video compressed using {'GPU' if nvenc_available else 'CPU'}")
                except subprocess.TimeoutExpired:
                    raise Exception("Video compression timed out")
                except subprocess.CalledProcessError as e:
                    logging.error(f"❌ Compression failed: {e.stderr}")
                    raise Exception(f"Video compression failed: {e.stderr}")

            if not os.path.exists(compressed) or os.path.getsize(compressed) == 0:
                raise Exception("Compressed video file is empty")

            compressed_size = os.path.getsize(compressed)
            logging.info(f"✅ Compressed file: {compressed_size} bytes")

            # ==================== UPLOAD VIDEO TO S3 FIRST (CRITICAL) ==========
            meeting_type = get_meeting_type(meeting_id)
            logging.info("☁ Uploading video to S3...")

            # Use session-based S3 path so multiple recordings don't overwrite each other
            video_s3_key = build_s3_video_path(meeting_id, user_id, meeting_type, session_id=session_id)
            logging.info(f"📁 Video S3 key: {video_s3_key}")
            video_url = upload_to_aws_s3(compressed, video_s3_key)
            
            if not video_url:
                raise Exception("Failed to upload video to S3")

            logging.info(f"✅ Video uploaded: {video_url}")

            # ========== THUMBNAIL GENERATION ==========
            thumbnail_url = None
            try:
                thumbnail_path = os.path.join(workdir, "thumbnail.jpg")
                seek_time = min(3.0, max(0.5, video_duration * 0.1))

                # Use -ss AFTER -i for accurate (frame-precise) seek instead of fast seek.
                # Fast seek (-ss before -i) often returns no frame on freshly-encoded video.
                thumb_cmd = [
                    "ffmpeg", "-y",
                    "-i", compressed,
                    "-ss", str(seek_time),
                    "-vframes", "1",
                    "-vf", "scale=640:-2",
                    "-q:v", "3",
                    thumbnail_path,
                ]
                logging.info(f"🖼 Generating thumbnail from {compressed} at {seek_time}s")
                proc = subprocess.run(thumb_cmd, capture_output=True, text=True, timeout=60)

                if proc.returncode != 0:
                    logging.error(
                        f"❌ ffmpeg thumbnail (primary) failed code={proc.returncode}: "
                        f"{(proc.stderr or '')[:500]}"
                    )
                    # Fallback: just grab the very first frame, no seek
                    logging.info("🔁 Retrying thumbnail with no seek (frame 0)...")
                    fallback_cmd = [
                        "ffmpeg", "-y",
                        "-i", compressed,
                        "-vframes", "1",
                        "-vf", "scale=640:-2",
                        "-q:v", "3",
                        thumbnail_path,
                    ]
                    proc2 = subprocess.run(fallback_cmd, capture_output=True, text=True, timeout=60)
                    if proc2.returncode != 0:
                        logging.error(
                            f"❌ ffmpeg thumbnail (fallback) ALSO failed code={proc2.returncode}: "
                            f"{(proc2.stderr or '')[:500]}"
                        )

                if os.path.exists(thumbnail_path) and os.path.getsize(thumbnail_path) > 0:
                    thumb_size = os.path.getsize(thumbnail_path)
                    if session_id:
                        thumb_s3_key = f"thumbnails/{meeting_id}/{session_id}/thumbnail.jpg"
                    else:
                        thumb_s3_key = f"thumbnails/{meeting_id}_{user_id}_thumbnail.jpg"

                    logging.info(f"☁ Uploading thumbnail ({thumb_size} bytes) to s3://{AWS_S3_BUCKET}/{thumb_s3_key}")
                    thumbnail_url = upload_to_aws_s3(thumbnail_path, thumb_s3_key)
                    if thumbnail_url:
                        logging.info(f"✅ Thumbnail uploaded: {thumbnail_url}")
                    else:
                        logging.error(f"❌ Thumbnail S3 upload returned None for key: {thumb_s3_key}")
                else:
                    logging.warning(f"⚠ Thumbnail file not created or empty at: {thumbnail_path}")
            except Exception as thumb_error:
                import traceback
                logging.error(f"❌ Thumbnail generation exception: {thumb_error}")
                logging.error(f"❌ Traceback: {traceback.format_exc()}")
                thumbnail_url = None

            # ========== AUDIO EXTRACTION FOR TRANSCRIPTION (RAW & ACCURATE) ==========
            audio = os.path.join(workdir, "audio.wav")

            audio_extract_cmd = [
                "ffmpeg", "-y",
                "-i", video_path,
                "-vn",
                "-ac", "1",
                "-ar", "16000",
                "-c:a", "pcm_s16le",
                audio
            ]

            try:
                # Dynamic timeout: 1 sec per MB, min 2 min, max 30 min
                video_size_mb = os.path.getsize(video_path) / (1024 * 1024)
                # Increased ceiling to 2 hours for long videos
                audio_timeout = max(120, min(7200, int(video_size_mb * 2)))
                logging.info(f"🕐 Audio extraction timeout: {audio_timeout}s (video: {video_size_mb:.1f}MB)")
                
                subprocess.run(
                    audio_extract_cmd,
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=audio_timeout
                )
                logging.info(f"✅ Raw audio extracted for transcription: {audio}")
                # ========== CACHE AUDIO TO S3 (FOR API-FAILURE RECOVERY) ==========
                audio_s3_key = None
                try:
                    if session_id:
                        audio_s3_key = f"intermediate/{meeting_id}/{session_id}/audio.wav"
                    else:
                        audio_s3_key = f"intermediate/{meeting_id}/{user_id}/audio.wav"

                    audio_size_mb = os.path.getsize(audio) / (1024 * 1024)
                    logging.info(f"☁ Caching audio to S3 for retry recovery ({audio_size_mb:.1f} MB): {audio_s3_key}")

                    audio_cache_url = upload_to_aws_s3(audio, audio_s3_key)
                    if audio_cache_url:
                        logging.info(f"✅ Audio cached for retry: {audio_s3_key}")
                    else:
                        logging.warning(f"⚠️ Audio cache upload returned None — retry recovery may not work")
                        audio_s3_key = None
                except Exception as cache_err:
                    logging.warning(f"⚠️ Audio caching failed (non-fatal): {cache_err}")
                    audio_s3_key = None
            except Exception as e:
                raise Exception(f"Audio extraction for transcription failed: {e}")

            # ========== TRANSCRIPTION (OpenAI gpt-4o-transcribe) ==========
            transcript_text = ""
            segments = []
            api_failure_types = []  # Track what kind of API failures occurred
            
            if client is None:
                logging.error("❌ OPENAI CLIENT IS NONE - Cannot transcribe!")
                logging.error(f"❌ OPENAI_API_KEY in environment: '{os.getenv('OPENAI_API_KEY', 'NOT_SET')[:10]}...'")
                transcript_text = "Transcription unavailable - OpenAI API not configured. Please check OPENAI_API_KEY environment variable."
                segments = [{
                    "start": 0.0,
                    "end": max(5.0, video_duration),
                    "text": "Transcription unavailable - OpenAI API not configured."
                }]
            else:
                logging.info("✅ OpenAI client is available, starting transcription with gpt-4o-transcribe...")

                try:
                    from pydub import AudioSegment
                    import gc

                    # Get duration first without loading whole file into memory
                    probe_result = subprocess.run(
                        ["ffprobe", "-v", "quiet", "-show_entries", "format=duration",
                         "-of", "default=noprint_wrappers=1:nokey=1", audio],
                        capture_output=True, text=True, timeout=30
                    )
                    total_audio_duration = float(probe_result.stdout.strip())
                    chunk_duration_sec_total = 5 * 60  # 5 minute chunks
                    chunk_ms = chunk_duration_sec_total * 1000
                    offset = 0.0
                    
                    # For very long audio (>2 hours), extract chunks via ffmpeg directly (memory-safe)
                    use_ffmpeg_chunking = total_audio_duration > 7200
                    
                    if not use_ffmpeg_chunking:
                        audio_file = AudioSegment.from_file(audio)
                    else:
                        audio_file = None
                        logging.info(f"📏 Long audio ({total_audio_duration:.0f}s) - using ffmpeg chunking to save memory")

                    HALLUCINATION_PATTERNS = [
                        "thanks for watching", "thank you for watching",
                        "please subscribe", "subscribe to my channel",
                        "like and subscribe", "don't forget to subscribe",
                        "thanks for listening", "see you next time",
                        "www.", ".com", "udemy", "coursera",
                        "spotlight shine", "music playing",
                        "[music]", "[applause]", "[silence]",
                    ]

                    def is_hallucination(text: str) -> bool:
                        if not text or len(text.strip()) < 3:
                            return True
                        text_lower = text.lower().strip()
                        for pattern in HALLUCINATION_PATTERNS:
                            if pattern in text_lower:
                                return True
                        words = text_lower.split()
                        if len(words) >= 4:
                            unique_words = set(words)
                            if len(unique_words) / len(words) < 0.3:
                                return True
                        return False

                    if use_ffmpeg_chunking:
                        total_chunks = int(total_audio_duration / chunk_duration_sec_total) + 1
                        chunk_iter = range(0, total_chunks)
                    else:
                        chunk_iter = range(0, len(audio_file), chunk_ms)
                    
                    for chunk_idx in chunk_iter:
                        if use_ffmpeg_chunking:
                            i = chunk_idx
                            start_sec = chunk_idx * chunk_duration_sec_total
                            if start_sec >= total_audio_duration:
                                break
                            chunk_path = os.path.join(workdir, f"chunk_{i}.wav")
                            chunk_extract_cmd = [
                                "ffmpeg", "-y", "-ss", str(start_sec),
                                "-i", audio, "-t", str(chunk_duration_sec_total),
                                "-c", "copy", chunk_path
                            ]
                            try:
                                subprocess.run(chunk_extract_cmd, check=True, capture_output=True, timeout=120)
                                chunk_duration_sec = min(chunk_duration_sec_total, total_audio_duration - start_sec)
                            except Exception as ext_err:
                                logging.error(f"[CHUNK {i}] ffmpeg extract failed: {ext_err}")
                                continue
                        else:
                            i = chunk_idx
                            chunk = audio_file[i:i + chunk_ms]
                            chunk_path = os.path.join(workdir, f"chunk_{i}.wav")
                            chunk.export(chunk_path, format="wav")
                            chunk_duration_sec = len(chunk) / 1000.0
                            del chunk
                            gc.collect()

                        # Retry up to 3 times on API failures
                        result = None
                        last_error = None
                        for attempt in range(3):
                            try:
                                with open(chunk_path, "rb") as f:
                                    result = client.audio.transcriptions.create(
                                        model="gpt-4o-transcribe",
                                        file=f,
                                        response_format="json",
                                        temperature=0.0,
                                    prompt=(
                                        "You are a professional transcription and translation system.\n\n"
                                        
                                        "TASK: Listen to the audio and output ONLY clean English text.\n\n"
                                        
                                        "ABSOLUTE RULE - READ CAREFULLY:\n"
                                        "The output MUST be 100% English written in Latin alphabet (A-Z, a-z, 0-9, punctuation).\n"
                                        "ZERO tolerance for non-Latin scripts. This is non-negotiable.\n\n"
                                        
                                        "WHAT THE SPEAKER MIGHT USE:\n"
                                        "- Pure English → transcribe as spoken\n"
                                        "- Hindi (Devanagari script or spoken) → TRANSLATE to English\n"
                                        "- Telugu (Telugu script or spoken) → TRANSLATE to English\n"
                                        "- Tamil, Kannada, Malayalam, Bengali, Marathi, Gujarati, Punjabi, Urdu → TRANSLATE to English\n"
                                        "- Mixed Hindi/Telugu + English (code-switching) → TRANSLATE non-English parts, keep English\n"
                                        "- Romanized Hindi/Telugu (e.g., 'yela chayali', 'manam cheddam') → TRANSLATE to English meaning\n\n"
                                        
                                        "FORBIDDEN OUTPUT (these will cause system failure):\n"
                                        "- Devanagari characters: अ आ इ ई क ख ग घ etc.\n"
                                        "- Telugu characters: అ ఆ ఇ ఈ క ఖ గ ఘ etc.\n"
                                        "- Tamil, Chinese, Japanese, Korean, Arabic, Hebrew, Cyrillic, or ANY non-Latin script\n"
                                        "- Transliteration like 'ek', 'do', 'kya', 'hai' — always translate to meaning: 'one', 'two', 'what', 'is'\n\n"
                                        
                                        "CONTENT DOMAINS — the speaker could be teaching ANY subject:\n"
                                        "- Software & IT: Python, Java, C++, JavaScript, SQL, DevOps, cloud, databases, cybersecurity\n"
                                        "- Enterprise software: SAP, Oracle, Salesforce, ServiceNow, Workday, Microsoft Dynamics\n"
                                        "- Data & AI: data science, machine learning, deep learning, NLP, statistics, analytics\n"
                                        "- Engineering: mechanical, civil, electrical, electronics, chemical, aerospace\n"
                                        "- Sciences: physics, chemistry, biology, mathematics, astronomy, geology\n"
                                        "- Medical & health: anatomy, pharmacology, nursing, diagnostics, clinical procedures\n"
                                        "- Business: finance, accounting, marketing, sales, HR, operations, strategy, management\n"
                                        "- Legal: contracts, compliance, regulations, corporate law, criminal law\n"
                                        "- Academic: history, geography, economics, political science, literature, philosophy\n"
                                        "- Languages: English grammar, spoken English, IELTS, TOEFL, foreign languages\n"
                                        "- Skills & trades: cooking, carpentry, electrical work, automotive, beauty, fashion\n"
                                        "- Arts & creative: music, painting, photography, writing, design, video editing\n"
                                        "- Fitness & lifestyle: yoga, gym, nutrition, meditation, sports coaching\n"
                                        "- Exam prep: UPSC, SSC, banking, GATE, NEET, JEE, CAT, school/college subjects\n"
                                        "- Religious & spiritual: scripture study, philosophy, meditation practices\n"
                                        "- General knowledge: current affairs, news, documentaries, tutorials\n\n"
                                        
                                        "TRANSLATION EXAMPLES (showing different subjects):\n\n"
                                        
                                        "Speaker (Telugu, tech): 'మనం ఇప్పుడు database లో data insert చేద్దాం'\n"
                                        "Output: 'Now let us insert data into the database'\n\n"
                                        
                                        "Speaker (Hindi, cooking): 'पहले हम pan में oil डालेंगे फिर onion add करेंगे'\n"
                                        "Output: 'First we will put oil in the pan then add onion'\n\n"
                                        
                                        "Speaker (Hindi, medical): 'यह दवा दिन में दो बार लेनी है खाने के बाद'\n"
                                        "Output: 'This medicine should be taken twice a day after meals'\n\n"
                                        
                                        "Speaker (Romanized Telugu, business): 'company profit margin ni improve cheyali ante cost reduce cheyali'\n"
                                        "Output: 'To improve the company profit margin we need to reduce cost'\n\n"
                                        
                                        "Speaker (Mixed, education): 'Indus Valley civilization gurinchi telusukundam ippudu'\n"
                                        "Output: 'Let us now learn about the Indus Valley civilization'\n\n"
                                        
                                        "Speaker (Telugu, fitness): 'ఈ exercise 15 reps చేయాలి three sets లో'\n"
                                        "Output: 'This exercise should be done 15 reps in three sets'\n\n"
                                        
                                        "PRESERVE unchanged (do NOT translate these):\n"
                                        "- Proper nouns: people's names, place names, company names, product names, brand names\n"
                                        "- Technical terms, tool names, software names already in English\n"
                                        "- Acronyms: API, SQL, CPU, CEO, GDP, DNA, UPSC, SSC, IELTS, etc.\n"
                                        "- Numbers, measurements, units (kg, km, ml, °C, $, %)\n"
                                        "- Scientific terms, medical terms, legal terms already in English\n\n"
                                        
                                        "QUALITY STANDARDS:\n"
                                        "- Produce complete, grammatical English sentences with proper punctuation.\n"
                                        "- Keep speaker's depth and terminology intact — do not oversimplify.\n"
                                        "- Do NOT invent content. Do NOT add advertisements, promotions, or URLs.\n"
                                        "- If audio is silent or unclear, skip it — do not fabricate.\n\n"
                                        
                                        "OUTPUT: Only the clean English translation. No metadata, no labels, no headers."
                                    )
                                )
                                break  # success, exit retry loop
                            except Exception as api_error:
                                last_error = api_error
                                failure_type = _detect_api_failure_type(api_error)
                                api_failure_types.append(failure_type)
                                logging.warning(f"[CHUNK {i}] API attempt {attempt+1}/3 failed ({failure_type}): {api_error}")
                                if attempt < 2:
                                    time.sleep(2 ** attempt)  # 1s, then 2s backoff

                        if result is None:
                            logging.error(f"[CHUNK {i}] All 3 attempts failed: {last_error}")
                            offset += chunk_duration_sec
                            try:
                                os.remove(chunk_path)
                            except:
                                pass
                            continue

                        # Extract full transcript from this chunk
                        chunk_text = (getattr(result, "text", "") or "").strip()

                        # STEP 1: Remove garbage scripts (Korean, Chinese, Japanese, Arabic, Hebrew, Cyrillic)
                        original_length = len(chunk_text)
                        chunk_text = clean_garbage_scripts(chunk_text)
                        if len(chunk_text) < original_length:
                            logging.warning(f"[CHUNK {i}] Removed {original_length - len(chunk_text)} garbage chars")

                        # STEP 2: Remove repeated phrases/sentences (Whisper hallucination)
                        before_repeat_filter = len(chunk_text)
                        chunk_text = remove_repetitions(chunk_text)
                        removed_chars = before_repeat_filter - len(chunk_text)

                        if before_repeat_filter > 200 and (removed_chars / before_repeat_filter) > 0.70:
                            logging.warning(
                                f"[CHUNK {i}] HALLUCINATION DETECTED — "
                                f"{removed_chars}/{before_repeat_filter} chars "
                                f"({removed_chars / before_repeat_filter * 100:.0f}%) were repetition. "
                                f"Dropping entire chunk to prevent orphan stubs."
                            )
                            offset += chunk_duration_sec
                            try:
                                os.remove(chunk_path)
                            except:
                                pass
                            continue

                        if removed_chars > 0:
                            logging.warning(f"[CHUNK {i}] Removed repeated content ({removed_chars} chars)")

                        logging.info(f"[CHUNK {i}] Cleaned transcript: {len(chunk_text)} chars in {chunk_duration_sec:.1f}s audio")

                        # STEP 3: Skip ONLY if truly empty (defense in depth)
                        if not chunk_text or len(chunk_text) < 10:
                            logging.info(f"[CHUNK {i}] Skipped (truly empty)")
                            offset += chunk_duration_sec
                            try:
                                os.remove(chunk_path)
                            except:
                                pass
                            continue

                        # STEP 4: Split into subtitle-sized segments with timing
                        chunk_segments = split_text_into_timed_segments(
                            chunk_text,
                            chunk_start=offset,
                            chunk_duration=chunk_duration_sec
                        )
                        logging.info(f"[CHUNK {i}] Split into {len(chunk_segments)} subtitle segments")

                        # Filter hallucinated segments individually (don't drop whole chunk)
                        before_hall = len(chunk_segments)
                        chunk_segments = [
                            seg for seg in chunk_segments
                            if not is_hallucination(seg["text"]) or len(seg["text"]) > 50
                        ]
                        if len(chunk_segments) < before_hall:
                            logging.info(f"[CHUNK {i}] Filtered {before_hall - len(chunk_segments)} hallucinated segments")

                        # STEP 4.5: Further split long mixed-language segments for better translation
                        # Long segments with mixed Indic+English confuse the LLM translator
                        refined_segments = []
                        for seg in chunk_segments:
                            seg_text = seg["text"].strip()
                            if len(seg_text) > 60:
                                # Split at sentence boundaries
                                parts = re.split(r'(?<=[.!?।])\s+', seg_text)
                                if len(parts) > 1:
                                    seg_duration = seg["end"] - seg["start"]
                                    part_duration = seg_duration / len(parts)
                                    for p_idx, part in enumerate(parts):
                                        part = part.strip()
                                        if part:
                                            refined_segments.append({
                                                "start": round(seg["start"] + (p_idx * part_duration), 2),
                                                "end": round(seg["start"] + ((p_idx + 1) * part_duration), 2),
                                                "text": part
                                            })
                                else:
                                    # No sentence boundary found, try splitting at commas
                                    comma_parts = re.split(r',\s+', seg_text)
                                    if len(comma_parts) > 1 and len(seg_text) > 100:
                                        seg_duration = seg["end"] - seg["start"]
                                        part_duration = seg_duration / len(comma_parts)
                                        for p_idx, part in enumerate(comma_parts):
                                            part = part.strip()
                                            if part:
                                                refined_segments.append({
                                                    "start": round(seg["start"] + (p_idx * part_duration), 2),
                                                    "end": round(seg["start"] + ((p_idx + 1) * part_duration), 2),
                                                    "text": part
                                                })
                                    else:
                                        refined_segments.append(seg)
                            else:
                                refined_segments.append(seg)
                        
                        if len(refined_segments) != len(chunk_segments):
                            logging.info(f"[CHUNK {i}] Refined {len(chunk_segments)} → {len(refined_segments)} segments for better translation")
                        chunk_segments = refined_segments

                        if chunk_segments:
                            logging.info(f"[CHUNK {i}] Running LLM translator pass (covers all languages)...")
                            chunk_segments = translate_segments_batch(chunk_segments, "mixed")

                        # STEP 6: Safety net — retry any segment that's still non-English
                        still_non_english = []
                        for idx, seg in enumerate(chunk_segments):
                            text = seg.get("text", "")
                            non_ascii = sum(1 for c in text if ord(c) > 127)
                            if len(text) > 0 and (non_ascii / len(text)) > 0.05:
                                still_non_english.append(idx)

                        if still_non_english:
                            logging.warning(f"[CHUNK {i}] {len(still_non_english)} segments still non-English — retrying")
                            retry_segments = [chunk_segments[idx] for idx in still_non_english]
                            retry_translated = translate_segments_batch(retry_segments, "unknown")
                            for j, idx in enumerate(still_non_english):
                                if j < len(retry_translated):
                                    chunk_segments[idx] = retry_translated[j]

                        for seg in chunk_segments:
                            text = seg.get("text", "")
                            if not text:
                                continue
                            non_ascii = sum(1 for c in text if ord(c) > 127)
                            ratio = non_ascii / len(text)
                            if ratio < 0.05:
                                continue  # already clean English
                            elif ratio < 0.5:
                                cleaned_text = "".join(c if ord(c) < 128 else " " for c in text)
                                cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
                                if cleaned_text and len(cleaned_text) > 5:
                                    seg["text"] = cleaned_text
                                # else keep original — too little English to salvage
                            else:
                                # Translation failed for this segment — preserve original
                                # so the user at least sees what was said in the source language.
                                logger.warning(
                                    f"[CHUNK {i}] Translation failed, preserving original: {text[:60]}..."
                                )
                                # text stays as-is
                        
                        # STEP 7.5: Drop only TRUE garbled nonsense (relaxed filter)
                        clean_segments = []
                        for seg in chunk_segments:
                            text = seg.get("text", "").strip()
                            if not text:
                                continue
                            words = text.split()
                            # Drop only if very short AND no meaningful word
                            if len(words) < 3 and len(text) < 10:
                                logging.debug(f"[GARBLE FILTER] Dropped too-short: '{text}'")
                                continue
                            # Drop only if NO word is 4+ chars AND segment is short
                            has_meaningful_word = any(len(w) >= 4 for w in words)
                            if not has_meaningful_word and len(words) < 5:
                                logging.warning(f"[GARBLE FILTER] Dropped fragment: '{text[:60]}'")
                                continue
                            clean_segments.append(seg)
                        if len(clean_segments) != len(chunk_segments):
                            logging.info(f"[CHUNK {i}] Garble filter: {len(chunk_segments)} -> {len(clean_segments)}")
                        chunk_segments = clean_segments

                        # Only keep non-empty segments
                        chunk_segments = [s for s in chunk_segments if s.get("text", "").strip()]

                        segments.extend(chunk_segments)
                        offset += chunk_duration_sec
                        try:
                            os.remove(chunk_path)
                        except:
                            pass

                except Exception as transcription_error:
                    logging.error(f"❌ Transcription failed: {transcription_error}")
                    import traceback
                    logging.error(f"❌ Traceback: {traceback.format_exc()}")
                    segments = []
                    transcript_text = "Transcription failed."

            needs_retry = False
            retry_reason = None

            if not segments and api_failure_types:
                # Determine the dominant failure type
                from collections import Counter
                failure_counts = Counter(api_failure_types)
                dominant_failure = failure_counts.most_common(1)[0][0]

                if _is_recoverable_failure(dominant_failure):
                    needs_retry = True
                    retry_reason = dominant_failure
                    logging.error(
                        f"⚠️ Transcription FAILED for entire recording — "
                        f"dominant failure: {dominant_failure} "
                        f"(failure breakdown: {dict(failure_counts)})"
                    )
                    logging.error(
                        f"⚠️ Marking recording for auto-retry. Audio cached at: {audio_s3_key}"
                    )

            if not segments:
                segments = [{
                    "start": 0.0,
                    "end": max(5.0, video_duration),
                    "text": "No speech detected."
                }]

            # Group segments into readable paragraphs (~80 words each, broken at sentence ends)
            # so the transcript reads like prose, not a list of one-line fragments.
            def _group_segments_into_paragraphs(segs, target_words=80):
                paragraphs = []
                current = []
                current_words = 0
                for seg in segs:
                    text = seg.get("text", "").strip()
                    if not text:
                        continue
                    current.append(text)
                    current_words += len(text.split())
                    # Break paragraph only when we've reached target length AND
                    # the last sentence ended cleanly with a terminal punctuation mark.
                    if current_words >= target_words and text[-1:] in ".!?":
                        paragraphs.append(" ".join(current))
                        current = []
                        current_words = 0
                if current:
                    paragraphs.append(" ".join(current))
                return paragraphs

            transcript_paragraphs = _group_segments_into_paragraphs(segments, target_words=80)
            transcript_text = "\n\n".join(transcript_paragraphs)

            # ========== TRAINER PERFORMANCE EVALUATION ==========
            trainer_evaluation = {}
            if needs_retry:
                logging.info("⏭️ Skipping trainer evaluation (recording marked for retry)")
                trainer_evaluation = {
                    "technical_content": 0,
                    "explanation_clarity": 0,
                    "friendliness": 0,
                    "communication": 0,
                    "overall_feedback": "Pending API recovery"
                }
            try:
                if not needs_retry and transcript_text and len(transcript_text.strip()) > 50:
                    logging.info("📊 Analyzing trainer performance...")
                    trainer_evaluation = analyze_trainer_performance(transcript_text)
                    
                    if trainer_evaluation and not trainer_evaluation.get("error"):
                        logging.info(f"✅ Trainer evaluation completed:")
                        logging.info(f"   Technical Content: {trainer_evaluation.get('technical_content', 0)}%")
                        logging.info(f"   Explanation Clarity: {trainer_evaluation.get('explanation_clarity', 0)}%")
                        logging.info(f"   Friendliness: {trainer_evaluation.get('friendliness', 0)}%")
                        logging.info(f"   Communication: {trainer_evaluation.get('communication', 0)}%")
                    else:
                        logging.warning(f"⚠ Trainer evaluation returned error: {trainer_evaluation.get('error')}")
                else:
                    logging.info("ℹ️ Insufficient content for trainer evaluation")
                    trainer_evaluation = {
                        "technical_content": 0,
                        "explanation_clarity": 0,
                        "friendliness": 0,
                        "communication": 0,
                        "overall_feedback": "Insufficient content for evaluation"
                    }
            except Exception as eval_error:
                logging.error(f"❌ Trainer evaluation failed: {eval_error}")
                trainer_evaluation = {
                    "error": str(eval_error),
                    "technical_content": 0,
                    "explanation_clarity": 0,
                    "friendliness": 0,
                    "communication": 0,
                    "overall_feedback": "Evaluation failed"
                }

            # ========== TRAINER EVALUATION DOCX (HOST-ONLY) ==========
            trainer_eval_url = None
            try:
                if (
                    trainer_evaluation
                    and not trainer_evaluation.get("error")
                    and any(trainer_evaluation.get(k, 0) > 0 for k in
                            ["technical_content", "explanation_clarity",
                             "friendliness", "communication"])
                ):
                    eval_path = os.path.join(workdir, "trainer_evaluation.docx")
                    save_trainer_evaluation_docx(
                        trainer_evaluation,
                        eval_path,
                        meeting_title=f"Meeting {meeting_id}"
                    )
                    if os.path.exists(eval_path) and os.path.getsize(eval_path) > 0:
                        eval_s3_key = build_s3_document_path(
                            meeting_id, user_id, meeting_type, "trainer_evaluation",
                            session_id=session_id
                        )
                        trainer_eval_url = upload_to_aws_s3(eval_path, eval_s3_key)
                        logging.info(f"✅ Trainer evaluation DOCX uploaded: {trainer_eval_url}")
                    else:
                        logging.warning("⚠ Trainer evaluation DOCX file empty or missing")
                else:
                    logging.info("ℹ️ Skipping trainer evaluation DOCX (no valid scores)")
            except Exception as eval_doc_err:
                logging.warning(f"⚠ Trainer evaluation DOCX generation failed: {eval_doc_err}")

            # ========== SUBTITLES GENERATION (Groq → MarianMT → GoogleTranslator fallback) ==========
            subtitle_urls = {}

            meeting_type = get_meeting_type(meeting_id)
            logging.info(f"Meeting type for subtitles: {meeting_type}")

            if segments and len(segments) > 0:
                logging.info("🎬 Generating subtitles with 3-tier translation fallback...")
                
                for lang in ["en", "hi", "te"]:
                    try:
                        translated_segments = []
                        translation_method = "none"
                        
                        if lang == "en":
                            # English SRT — verify segments are actually English, force-clean if not
                            translated_segments = segments.copy()
                            
                            # Check if segments are actually in English (no non-ASCII chars)
                            sample_text = " ".join(s.get("text", "") for s in translated_segments[:5])
                            non_ascii = sum(1 for c in sample_text if ord(c) > 127)
                            
                            if len(sample_text) > 0 and (non_ascii / len(sample_text)) > 0.1:
                                # Segments still have non-English content — force translate to English
                                logging.warning(f"⚠️ English SRT has non-ASCII content ({non_ascii} chars), forcing translation to English")
                                translated_segments = translate_segments_batch(translated_segments, "mixed")
                                translation_method = "groq_forced_english"
                            else:
                                translation_method = "original"
                            
                            logging.info("✅ English SRT: clean English verified")
                            
                        else:
                            # ===== TIER 1: Try Groq Llama (primary, fastest & most accurate) =====
                            groq_success = False
                            try:
                                logging.info(f"🔄 [{lang}] Trying Tier 1: Groq Llama...")
                                translated_segments = translate_segments_to_target_lang(segments, lang)
                                
                                # Verify translation actually worked (check if output has target script)
                                if translated_segments and len(translated_segments) > 0:
                                    sample_text = " ".join(s.get("text", "") for s in translated_segments[:5])
                                    non_ascii = sum(1 for c in sample_text if ord(c) > 127)
                                    if len(sample_text) > 0 and (non_ascii / len(sample_text)) > 0.3:
                                        groq_success = True
                                        translation_method = "groq_llama"
                                        logging.info(f"✅ [{lang}] Tier 1 (Groq Llama) SUCCESS")
                                    else:
                                        logging.warning(f"⚠️ [{lang}] Tier 1 returned English text, trying Tier 2")
                                        
                            except Exception as groq_error:
                                logging.warning(f"⚠️ [{lang}] Tier 1 (Groq Llama) failed: {groq_error}")
                            
                            # ===== TIER 2: Try MarianMT (local model fallback) =====
                            if not groq_success:
                                marian_success = False
                                try:
                                    logging.info(f"🔄 [{lang}] Trying Tier 2: MarianMT local model...")
                                    local_translator = LocalIndianLanguageTranslator()
                                    translated_segments = local_translator.translate_segments(
                                        segments,
                                        lang,
                                        batch_size=32
                                    )
                                    
                                    # Verify translation worked
                                    if translated_segments and len(translated_segments) > 0:
                                        sample_text = " ".join(s.get("text", "") for s in translated_segments[:5])
                                        non_ascii = sum(1 for c in sample_text if ord(c) > 127)
                                        if len(sample_text) > 0 and (non_ascii / len(sample_text)) > 0.3:
                                            marian_success = True
                                            translation_method = "marian_mt"
                                            logging.info(f"✅ [{lang}] Tier 2 (MarianMT) SUCCESS")
                                        else:
                                            logging.warning(f"⚠️ [{lang}] Tier 2 returned English text, trying Tier 3")
                                    
                                    try:
                                        local_translator.cleanup()
                                    except:
                                        pass
                                        
                                except Exception as marian_error:
                                    logging.warning(f"⚠️ [{lang}] Tier 2 (MarianMT) failed: {marian_error}")
                                
                                # ===== TIER 3: Try GoogleTranslator (final fallback) =====
                                if not marian_success:
                                    google_success = False
                                    try:
                                        logging.info(f"🔄 [{lang}] Trying Tier 3: GoogleTranslator...")
                                        import requests as req_lib
                                        
                                        translated_segments = []
                                        for seg in segments:
                                            text = seg["text"].strip()
                                            
                                            if text and len(text) > 0:
                                                try:
                                                    translator = GoogleTranslator(source="en", target=lang)
                                                    
                                                    original_get = req_lib.get
                                                    
                                                    def patched_get(*args, **kwargs):
                                                        kwargs['verify'] = False
                                                        kwargs['timeout'] = 10
                                                        return original_get(*args, **kwargs)
                                                    
                                                    req_lib.get = patched_get
                                                    
                                                    try:
                                                        translated_text = translator.translate(text)
                                                    finally:
                                                        req_lib.get = original_get
                                                    
                                                    if not translated_text:
                                                        translated_text = text
                                                        
                                                except Exception as translate_error:
                                                    logging.warning(f"[{lang}] GoogleTranslator segment error: {translate_error}")
                                                    translated_text = text
                                            else:
                                                translated_text = text
                                            
                                            translated_segments.append({
                                                "start": seg["start"],
                                                "end": seg["end"],
                                                "text": translated_text
                                            })
                                        
                                        # Verify translation worked
                                        if translated_segments and len(translated_segments) > 0:
                                            sample_text = " ".join(s.get("text", "") for s in translated_segments[:5])
                                            non_ascii = sum(1 for c in sample_text if ord(c) > 127)
                                            if len(sample_text) > 0 and (non_ascii / len(sample_text)) > 0.3:
                                                google_success = True
                                                translation_method = "google_translator"
                                                logging.info(f"✅ [{lang}] Tier 3 (GoogleTranslator) SUCCESS")
                                            else:
                                                logging.warning(f"⚠️ [{lang}] Tier 3 returned English text")
                                                
                                    except Exception as google_error:
                                        logging.error(f"❌ [{lang}] Tier 3 (GoogleTranslator) failed: {google_error}")
                                    
                                    # ===== TIER 4: Final fallback — English (better than nothing) =====
                                    if not google_success:
                                        translated_segments = segments.copy()
                                        translation_method = "fallback_english"
                                        logging.error(f"⚠️ [{lang}] ALL tiers failed, using English (all translators failed)")
                        
                        # Create SRT file and upload
                        if translated_segments:
                            srt_path = os.path.join(workdir, f"subs_{lang}.srt")
                            create_srt_from_segments(translated_segments, srt_path)
                            
                            if os.path.exists(srt_path) and os.path.getsize(srt_path) > 0:
                                logging.info(f"✅ [{lang}] SRT created ({os.path.getsize(srt_path)} bytes) via {translation_method}")
                                
                                subtitles_folder = build_s3_document_path(meeting_id, user_id, meeting_type, "subtitles", session_id=session_id)
                                if session_id:
                                    s3_key = f"{subtitles_folder}/{meeting_id}_{session_id}_{lang}.srt"
                                else:
                                    s3_key = f"{subtitles_folder}/{meeting_id}_{user_id}_{lang}.srt"
                                subtitle_url = upload_to_aws_s3(srt_path, s3_key)
                                
                                if subtitle_url:
                                    subtitle_urls[lang] = subtitle_url
                                    logging.info(f"✅ [{lang}] subtitles uploaded to S3 (method: {translation_method})")
                                else:
                                    logging.error(f"❌ [{lang}] Failed to upload subtitles to S3")
                            else:
                                logging.warning(f"⚠️ [{lang}] SRT file is empty or not created")
                                
                    except Exception as lang_error:
                        logging.error(f"❌ Complete failure for {lang} subtitles: {lang_error}")
                        import traceback
                        logging.error(f"Traceback: {traceback.format_exc()}")
                
                logging.info(f"✅ Generated subtitles for languages: {list(subtitle_urls.keys())}")
            else:
                logging.warning("⚠ No segments available for subtitle generation")

            # ========== SUMMARY GENERATION ==========
            summary = "Processing summary..."
            try:
                if transcript_text and len(transcript_text.strip()) > 10:
                    summary = summarize_segment(transcript_text)
                    logging.info(f"✅ Summary generated ({len(summary)} chars)")
                else:
                    summary = "No sufficient content available."
            except Exception as summary_error:
                logging.warning(f"⚠ Summary generation failed: {summary_error}")
                summary = "Summary generation failed."

            # ========== SUMMARY + MIND MAP (END OF DOCUMENT) ==========
            import re
            from docx import Document
            from docx.shared import Inches
            summary_doc_path = os.path.join(workdir, "summary.docx")

            # ---- Mind Map extraction & rendering ----
            image_url = None
            mindmap_png_path = None
            try:
                dot_code = None

                # Pattern 1: ```dot/graphviz``` block
                dot_match = re.search(r"```(?:dot|graphviz)?\s*(.*?)```", summary, re.DOTALL | re.IGNORECASE)
                if dot_match:
                    raw_dot = dot_match.group(1).strip()
                    if raw_dot.lower().startswith(("digraph", "graph")) or "digraph" in raw_dot.lower():
                        dot_code = raw_dot
                        logging.info("Found DOT code in ``` block (Pattern 1)")

                # Pattern 2: Raw digraph/graph without backticks
                if not dot_code:
                    dot_match = re.search(r"((?:di)?graph\s+[\"']?\w+[\"']?\s*\{[^}]+\})", summary, re.DOTALL | re.IGNORECASE)
                    if dot_match:
                        dot_code = dot_match.group(1)
                        logging.info("Found raw DOT code (Pattern 2)")

                # Pattern 3: Flexible multiline
                if not dot_code:
                    dot_match = re.search(r"(digraph\s+\w+\s*\{[\s\S]*?\n\})", summary, re.IGNORECASE)
                    if dot_match:
                        dot_code = dot_match.group(1)
                        logging.info("Found DOT code (Pattern 3)")

                # Pattern 4: Single line fallback
                if not dot_code:
                    all_text = summary.replace('\n', ' ')
                    dot_match = re.search(r'(digraph\s+\w+\s*\{[^}]+\})', all_text, re.IGNORECASE)
                    if dot_match:
                        dot_code = dot_match.group(1)
                        logging.info("Found DOT code (Pattern 4)")

                if dot_code:
                    dot_code = sanitize_dot_code(dot_code)
                    dot_code = enhance_dot_styling(dot_code)

                    mindmap_png_path = os.path.join(workdir, "mindmap.png")
                    generate_graph(dot_code, mindmap_png_path[:-4])

                    if session_id:
                        mindmap_s3_key = f"{meeting_id}/{session_id}/mindmap.png"
                    else:
                        mindmap_s3_key = f"{meeting_id}/{user_id}/mindmap.png"
                    image_url = upload_to_aws_s3(mindmap_png_path, mindmap_s3_key)

                    logging.info(f"Concept map generated and uploaded successfully")
                else:
                    logging.warning("No valid DOT code found in summary - concept map will not be generated")

            except Exception as mindmap_error:
                logging.warning(f"Concept map generation failed: {mindmap_error}")
                import traceback
                logging.error(f"Concept map traceback: {traceback.format_exc()}")
                image_url = None
                mindmap_png_path = None

            # ---- Build Summary DOCX using save_docx ----
            save_docx(
                content=summary,
                path=summary_doc_path,
                image_path=mindmap_png_path if mindmap_png_path and os.path.exists(mindmap_png_path) else None,
                title="Meeting Summary"
            )
            logging.info(f"Summary DOCX created at {summary_doc_path}")

            transcript_url = None
            summary_url = None

            # --- Transcript DOCX ---
            transcript_path = os.path.join(workdir, "transcript.docx")
            save_docx(transcript_text or "No speech detected.", transcript_path, title="Meeting Transcript")

            summary_path = summary_doc_path

            if os.path.exists(transcript_path) and os.path.getsize(transcript_path) > 0:
                transcript_s3_key = build_s3_document_path(meeting_id, user_id, meeting_type, "transcript", session_id=session_id)
                transcript_url = upload_to_aws_s3(transcript_path, transcript_s3_key)

            if os.path.exists(summary_path) and os.path.getsize(summary_path) > 0:
                summary_s3_key = build_s3_document_path(meeting_id, user_id, meeting_type, "summary", session_id=session_id)
                summary_url = upload_to_aws_s3(summary_path, summary_s3_key)

            logging.info("✅ S3 uploads completed")

            # ========== GET PARTICIPANT EMAILS ==========
            visible_to_emails = get_meeting_participants_emails(meeting_id)
            logging.info(f"✅ Recording visible to {len(visible_to_emails)} users")

            user_details = get_user_details(user_id)
            logging.info(f"👤 User details for recording: {user_details}")

            # ========== CHECK CUSTOM NAME ==========
            custom_recording_name = None
            try:
                custom_name_doc = None
                if session_id:
                    # First: check if a name was already claimed by this session
                    custom_name_doc = collection.find_one({
                        "meeting_id": meeting_id,
                        "claimed_by_session": session_id,
                        "custom_recording_name": {"$exists": True, "$ne": None}
                    })
                    if custom_name_doc:
                        logging.info(f"Found custom name already claimed by session {session_id}: {custom_name_doc.get('custom_recording_name')}")
                    
                    # Second: check pending names matching this session
                    if not custom_name_doc:
                        custom_name_doc = collection.find_one({
                            "meeting_id": meeting_id,
                            "pending_name_update": True,
                            "custom_name_session_id": session_id
                        })
                    if custom_name_doc:
                        logging.info(f"Found custom name by session_id {session_id}: {custom_name_doc.get('custom_recording_name')}")
                
                # Fallback: match by oldest unclaimed name for this meeting (legacy recordings)
                if not custom_name_doc:
                    custom_name_doc = collection.find_one(
                        {
                            "meeting_id": meeting_id,
                            "pending_name_update": True,
                            "claimed_by_session": {"$exists": False}
                        },
                        sort=[("name_stored_at", 1)]  # ✅ Oldest first (FIFO), not newest
                    )
                    if custom_name_doc:
                        logging.info(f"Found custom name by FIFO fallback: {custom_name_doc.get('custom_recording_name')}")
                
                if custom_name_doc:
                    custom_recording_name = custom_name_doc.get("custom_recording_name")
                    # ✅ Mark this name as claimed so other recordings don't grab it
                    collection.update_one(
                        {"_id": custom_name_doc["_id"]},
                        {"$set": {
                            "pending_name_update": False,
                            "claimed_by_session": session_id or "legacy",
                            "claimed_at": datetime.now()
                        }}
                    )
                    logging.info(f"Claimed custom name: {custom_recording_name}")
            except Exception as e:
                logging.warning(f"Failed to check custom name: {e}")

            # ========== SCHEDULE METADATA ==========
            schedule_meta = {}
            if meeting_type == "ScheduleMeeting":
                schedule_meta = get_schedule_meeting_metadata(meeting_id)

            if custom_recording_name:
                display_filename = f"{custom_recording_name}.mp4"
                original_filename = f"{custom_recording_name}.mp4"
            else:
                display_filename = os.path.basename(video_path)
                original_filename = os.path.basename(video_path)

            # ========== SAVE TO MONGODB ==========
            recording_sequence = int(time.time())  # Default timestamp

            # session_id was already extracted from filename at the top of this function
            # Only try S3 key extraction if we don't already have one
            if not session_id and video_s3_key and "/" in video_s3_key:
                path_parts = video_s3_key.split("/")
                if len(path_parts) >= 3:
                    potential_session = path_parts[-2]
                    if len(potential_session) == 8 and "-" not in potential_session:
                        session_id = potential_session
                        logging.info(f"✅ Extracted session_id from S3 key: {session_id}")
            logging.info(f"📋 Final session_id for MongoDB save: {session_id}")

            # For very long summaries, truncate MongoDB copy (full version stays in S3)
            summary_for_mongo = summary
            if len(summary) > 500000:  # 500KB safety threshold
                summary_for_mongo = summary[:500000] + "\n\n[... truncated - see full summary in S3 ...]"
                logging.warning(f"⚠️ Summary too large ({len(summary)} chars), truncating for MongoDB storage")
            
            video_document = {
                "meeting_id": meeting_id,
                "session_id": session_id,  # ✅ NEW: Track recording session
                "recording_sequence": recording_sequence,  # ✅ NEW: Sort multiple recordings
                "user_id": user_id,
                "user_name": user_details.get("user_name", f"User {user_id}"),
                "user_email": user_details.get("user_email", ""),
                "meeting_type": meeting_type,
                "schedule_id": schedule_meta.get("schedule_id"),
                "schedule_title": schedule_meta.get("schedule_title"),
                "schedule_folder": schedule_meta.get("folder_path"),
                "filename": display_filename,
                "original_filename": original_filename,
                "custom_recording_name": custom_recording_name,
                "video_url": video_url,
                "thumbnail_url": thumbnail_url,   
                "transcript_url": transcript_url,
                "summary_url": summary_url,
                "summary_text": summary_for_mongo,
                "image_url": image_url,
                "subtitles": subtitle_urls,
                "timestamp": datetime.now(),
                "visible_to": visible_to_emails,
                "file_size": compressed_size,
                "duration": video_duration,
                "transcription_available": bool(transcript_url),
                "summary_available": bool(summary_url),
                "processing_status": "needs_retry" if needs_retry else "completed",
                "retry_reason": retry_reason if needs_retry else None,
                "audio_s3_key": audio_s3_key if needs_retry else None,
                "retry_count": 0 if needs_retry else None,
                "last_retry_at": None,
                "max_retries": 24,
                "subtitle_format": "srt_multi_language",
                "subtitle_languages": list(subtitle_urls.keys()),
                "embedded_subtitles": False,
                "audio_processing_status": "success" if os.path.exists(audio) else "fallback",
                "audio_preserved": has_audio,
                "source_format": "pyav_mp4",
                "smooth_playback": True,
                "file_type": "video/mp4",
                "is_final_video": True,
                "encoder_used": "GPU (NVENC)" if nvenc_available else "CPU (libx264)",
                "gpu_accelerated": nvenc_available,
                "document_format": "docx",
                "transcription_engine": "openai_gpt_4o_transcribe_with_llm_translation",
                "trainer_evaluation": trainer_evaluation,
                "trainer_evaluation_url": trainer_eval_url
            }
            
            logging.info(f"💾 Saving video data to MongoDB...")

            # ✅ Find existing document for this session and update it
            if session_id:
                # Query by session_id, prefer the most recent document
                existing_doc = collection.find_one(
                    {"meeting_id": meeting_id, "session_id": session_id},
                    sort=[("_id", -1)]  # Newest first
                )
                
                if existing_doc:
                    collection.replace_one({"_id": existing_doc["_id"]}, video_document)
                    logging.info(f"✅ Updated existing document for session: {session_id}")
                    
                    # Clean up any duplicate docs for this session (from stale start-recording entries)
                    try:
                        duplicates = collection.find(
                            {"meeting_id": meeting_id, "session_id": session_id, "_id": {"$ne": existing_doc["_id"]}}
                        )
                        dup_count = 0
                        for dup in duplicates:
                            collection.delete_one({"_id": dup["_id"]})
                            dup_count += 1
                        if dup_count > 0:
                            logging.info(f"🧹 Cleaned up {dup_count} duplicate docs for session {session_id}")
                    except Exception as cleanup_err:
                        logging.warning(f"⚠️ Duplicate cleanup failed: {cleanup_err}")
                else:
                    collection.insert_one(video_document)
                    logging.info(f"✅ Created new document for session: {session_id}")
                
                # Also clean up stale start-recording docs without session_id for this meeting
                try:
                    stale_result = collection.delete_many({
                        "meeting_id": meeting_id,
                        "recording_status": {"$in": ["active", "stopped", "starting"]},
                        "session_id": {"$exists": False}
                    })
                    if stale_result.deleted_count > 0:
                        logging.info(f"🧹 Cleaned up {stale_result.deleted_count} stale start-recording docs")
                except Exception:
                    pass
            else:
                # Legacy recording without session_id - fall back to old behavior
                query_filter = {
                    "meeting_id": meeting_id,
                    "user_id": user_id,
                    "is_final_video": True,
                    "session_id": None  # ✅ Only match old recordings without session_id
                }
                
                existing_doc = collection.find_one(query_filter)
                
                if existing_doc:
                    collection.replace_one({"_id": existing_doc["_id"]}, video_document)
                    logging.info(f"✅ Updated legacy document (no session_id)")
                else:
                    collection.insert_one(video_document)
                    logging.info(f"✅ Created new legacy document (no session_id)")
                    
            # ✅ CRITICAL: Don't delete custom name here - it was already marked as used in recording service
            # The recording service marks names as used in _async_finalize_fast_recording
            # If we delete here, we might delete names for OTHER concurrent recordings
            logging.info(f"ℹ️ Skipping custom name cleanup (already handled by recording service)")

            # ========== SEND NOTIFICATIONS ==========
            if needs_retry:
                logging.info(f"⏭️ Skipping notifications — recording in needs_retry state. Will notify on successful retry.")
                notification_count = 0
            else:
                logging.info(f"📧 Sending notifications...")
                try:
                    notification_count = send_recording_completion_notifications(
                        meeting_id=meeting_id,
                        video_url=video_url,
                        transcript_url=transcript_url,
                        summary_url=summary_url
                    )
                    logging.info(f"✅ Sent {notification_count} notifications")
                except Exception as notif_error:
                    logging.error(f"⚠ Notifications failed: {notif_error}")

            # ========== RETURN SUCCESS ==========
            result_dict = {
                "status": "success",
                "video_url": video_url,
                "thumbnail_url": thumbnail_url,   # 👈 ADD THIS LINE
                "transcript_url": transcript_url,
                "summary_url": summary_url,
                "summary_image_url": image_url,
                "subtitle_urls": subtitle_urls,
                "subtitle_languages": list(subtitle_urls.keys()),
                "file_size": compressed_size,
                "meeting_id": meeting_id,
                "user_id": user_id,
                "subtitle_format": "srt_multi_language",
                "authorized_users_count": len(visible_to_emails),
                "encoder_used": "GPU (NVENC)" if nvenc_available else "CPU (libx264)",
                "gpu_accelerated": nvenc_available,
                "document_format": "docx",
                "transcription_engine": "openai_gpt_4o_transcribe",
                "trainer_evaluation": trainer_evaluation,
                "processing_notes": {
                    "audio_extracted": os.path.exists(audio),
                    "transcription_successful": bool(segments),
                    "subtitles_generated": len(subtitle_urls),
                    "summary_generated": len(summary) > 50,
                    "original_had_audio": has_audio,
                    "audio_preserved": has_audio,
                    "source_format": "pyav_mp4",
                    "smooth_playback_enabled": True,
                    "duration_preserved": video_duration > 0,
                    "only_final_mp4_saved": True,
                    "gpu_acceleration_used": nvenc_available,
                    "trainer_evaluated": bool(trainer_evaluation and not trainer_evaluation.get("error"))
                }
            }
            
            logging.info(f"✅ Processing completed using {'GPU' if nvenc_available else 'CPU'}")
            return result_dict
            
        except Exception as e:
            logging.error(f"❌ Video processing failed: {e}")
            import traceback
            logging.error(f"❌ Traceback: {traceback.format_exc()}")
            
            return {
                "status": "error",
                "error": str(e),
                "error_type": type(e).__name__,
                "video_url": None,
                "transcript_url": None,
                "summary_url": None,
                "subtitle_urls": {},
                "file_size": 0,
                "meeting_id": meeting_id,
                "user_id": user_id,
                "processing_notes": {
                    "failed_at": "video_processing",
                    "input_file_size": input_size if 'input_size' in locals() else 0,
                    "input_extension": input_ext if 'input_ext' in locals() else "unknown"
                }
            }

# ============================================================
# API FAILURE DETECTION (for auto-retry feature)
# ============================================================
def _detect_api_failure_type(exception) -> str:
    """
    Classify API exception so retry task knows what to do.

    Returns:
    - 'auth_failed'      → 401 / bad/expired API key (user/admin must update key)
    - 'payment_required' → 402 / billing issue (user must pay)
    - 'quota_exceeded'   → 429 + quota text (subscription out of credits)
    - 'rate_limit'       → 429 without quota text (temporary, retries work)
    - 'network_error'    → timeout/connection (transient)
    - 'unknown_error'    → anything else
    """
    error_str = str(exception).lower()
    error_type = type(exception).__name__.lower()

    if '401' in error_str or 'unauthorized' in error_str or 'invalid_api_key' in error_str or 'invalid api key' in error_str:
        return 'auth_failed'

    if '402' in error_str or 'payment_required' in error_str or 'payment required' in error_str:
        return 'payment_required'

    if '429' in error_str and ('quota' in error_str or 'insufficient_quota' in error_str or 'billing' in error_str or 'exceeded' in error_str):
        return 'quota_exceeded'

    if '429' in error_str or 'rate_limit' in error_str or 'rate limit' in error_str:
        return 'rate_limit'

    if any(word in error_str for word in ['timeout', 'connection', 'network', 'unreachable', 'dns', 'resolve']):
        return 'network_error'

    if any(word in error_type for word in ['timeout', 'connection']):
        return 'network_error'

    return 'unknown_error'


def _is_recoverable_failure(failure_type: str) -> bool:
    """All failures are technically recoverable, but these benefit most from retry."""
    return failure_type in ('auth_failed', 'payment_required', 'quota_exceeded', 'rate_limit', 'network_error', 'unknown_error')

# === 1. GET ALL VIDEOS ===
@require_http_methods(["GET"])
def get_all_videos(request):
    """Get all video documents with pagination and filtering - STRICT ACCESS CONTROL + MEETING TYPE FILTER"""
    try:
        # Query parameters for pagination and filtering
        page = int(request.GET.get('page', 1))
        limit = int(request.GET.get('limit', 10))
        user_id = request.GET.get('user_id')
        email = request.GET.get('email', '')
        meeting_id = request.GET.get('meeting_id')
        meeting_type = request.GET.get('meeting_type')  # ✅ NEW PARAMETER

        # Build query filter - ONLY SHOW FINAL VIDEOS
        query_filter = {"is_final_video": True}
        
        if meeting_id:
            query_filter['meeting_id'] = meeting_id
        
        # ✅ NEW: Filter by meeting type
        if meeting_type:
            if meeting_type in ['CalendarMeeting', 'ScheduleMeeting', 'InstantMeeting']:
                query_filter['meeting_type'] = meeting_type
                logger.info(f"Filtering videos by meeting_type: {meeting_type}")
            else:
                logger.warning(f"Invalid meeting_type parameter: {meeting_type}")

        # Calculate skip value for pagination
        skip = (page - 1) * limit

        # Debug logging
        logger.info(f"📋 Query params: email={email}, user_id={user_id}, meeting_id={meeting_id}, meeting_type={meeting_type}")
        
        # Fetch raw results
        raw_videos = list(collection.find(query_filter).sort("timestamp", -1).skip(skip).limit(limit))
        logger.info(f"📹 Raw videos found in MongoDB: {len(raw_videos)}")

        allowed_videos = []
        for video in raw_videos:
            video_user_id = video.get("user_id")
            visible_to_emails = video.get("visible_to", [])
            video_meeting_id = video.get("meeting_id")
            
            # STRICT ACCESS CONTROL - ONLY 2 WAYS TO ACCESS:
            
            # Way 1: User uploaded/created this video (is the original uploader)
            user_uploaded = False
            if user_id and video_user_id:
                user_uploaded = str(video_user_id) == str(user_id)
            
            # Way 2: User's email is in the visible_to list (was authorized during processing)
            email_authorized = False
            if email and visible_to_emails:
                # Case-insensitive email comparison
                user_email_lower = email.strip().lower()
                visible_emails_lower = [e.strip().lower() for e in visible_to_emails if e]
                email_authorized = user_email_lower in visible_emails_lower
            
            # ALLOW ACCESS ONLY if user uploaded OR email is in visible_to list
            has_access = user_uploaded or email_authorized
            
            # Debug logging for each video (detailed)
            logger.debug(f"🎬 Video {video.get('_id')} (Meeting: {video_meeting_id}, Type: {video.get('meeting_type')}):")
            logger.debug(f"   - Video uploader: {video_user_id}, Current user: {user_id}, Match: {user_uploaded}")
            logger.debug(f"   - Visible to: {visible_to_emails}")
            logger.debug(f"   - User email: {email}, Match: {email_authorized}")
            logger.debug(f"   - Final access decision: {has_access}")
            
            if has_access:
                # Process video data
                video['_id'] = str(video['_id'])
                video['timestamp'] = video['timestamp'].isoformat() if video.get('timestamp') else None
                
                # ✅ Ensure meeting_type is present (backfill if missing)
                if not video.get('meeting_type'):
                    video['meeting_type'] = get_meeting_type(video_meeting_id)
                    # Update MongoDB with the meeting type
                    try:
                        collection.update_one(
                            {"_id": ObjectId(video['_id'])},
                            {"$set": {"meeting_type": video['meeting_type']}}
                        )
                        logger.info(f"Backfilled meeting_type for video {video['_id']}: {video['meeting_type']}")
                    except Exception as backfill_error:
                        logger.warning(f"Failed to backfill meeting_type: {backfill_error}")
                
                # ✅ Set display name (prioritize custom name over filename)
                if video.get('custom_recording_name'):
                    video['display_name'] = video['custom_recording_name']
                elif video.get('display_name'):
                    video['display_name'] = video['display_name']
                else:
                    # Fallback: clean up filename for display
                    filename = video.get('filename', 'Unnamed Recording')
                    # Remove technical prefixes like "raw_video_" and file extensions
                    if filename.startswith('raw_video_'):
                        filename = filename.replace('raw_video_', '').replace('_final.mp4', '').replace('.mp4', '')
                    video['display_name'] = filename if filename else 'Unnamed Recording'
                
                # Add access reason for debugging
                if user_uploaded:
                    video['access_reason'] = 'uploader'
                    logger.info(f"✅ Video {video.get('_id')}: Access granted (uploader)")
                elif email_authorized:
                    video['access_reason'] = 'authorized_email'
                    logger.info(f"✅ Video {video.get('_id')}: Access granted (authorized email)")
                
                # ✅ CRITICAL FIX: Compute subtitles_available field
                subtitles_dict = video.get('subtitles', {})
                video['subtitles_available'] = bool(subtitles_dict and len(subtitles_dict) > 0)
                
                allowed_videos.append(video)
            else:
                logger.debug(f"❌ Video {video.get('_id')}: Access denied for user {user_id}/{email}")

        logger.info(f"✅ Videos after access filtering: {len(allowed_videos)}/{len(raw_videos)}")

        return JsonResponse({
            "status": "success",
            "data": allowed_videos,
            "videos": allowed_videos,  # Alternative key for compatibility
            "pagination": {
                "page": page,
                "limit": limit,
                "total": len(allowed_videos),
                "total_pages": (len(allowed_videos) + limit - 1) // limit
            },
            "debug": {
                "raw_count": len(raw_videos),
                "filtered_count": len(allowed_videos),
                "query_params": {
                    "email": email,
                    "user_id": user_id,
                    "meeting_id": meeting_id,
                    "meeting_type": meeting_type  # ✅ NEW
                }
            }
        })

    except Exception as e:
        logger.error(f"[ERROR] Failed to get videos: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 2. GET VIDEO BY ID ===
@require_http_methods(["GET"])
def get_video_by_id(request, id):
    """Get a specific video document by ID with access control."""
    try:
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')

        video = collection.find_one({"_id": ObjectId(id)})
        if not video:
            logger.error(f"Video ID {id} not found")
            return JsonResponse({"Error": "Video not found"}, status=404)

        meeting_id = video.get("meeting_id", "")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "You are not authorized to view this video"}, status=403)

        video['_id'] = str(video['_id'])
        video['timestamp'] = video['timestamp'].isoformat() if video.get('timestamp') else None
        
        # ✅ CRITICAL FIX: Compute subtitles_available field
        subtitles_dict = video.get('subtitles', {})
        video['subtitles_available'] = bool(subtitles_dict and len(subtitles_dict) > 0)

        return JsonResponse({
            "status": "success",
            "data": video
        })

    except Exception as e:
        logger.error(f"[ERROR] Failed to get video by ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 3. UPDATE VIDEO (PUT) ===
@require_http_methods(["PUT"])
@csrf_exempt
def update_video(request, id):
    """Update a video document."""
    try:
        data = json.loads(request.body)
        
        update_data = {}
        allowed_fields = ['meeting_id', 'user_id', 'video_url', 'transcript_url','summary_text', 'summary_url', 'image_url', 'subtitles']
        
        for field in allowed_fields:
            if field in data:
                update_data[field] = data[field]
        
        if not update_data:
            return JsonResponse({"Error": "No valid fields to update"}, status=400)
        
        result = collection.update_one(
            {"_id": ObjectId(id)},
            {"$set": update_data}
        )
        
        if result.matched_count == 0:
            logger.error(f"Video ID {id} not found for update")
            return JsonResponse({"Error": "Video not found"}, status=404)
        
        updated_video = collection.find_one({"_id": ObjectId(id)})
        updated_video['_id'] = str(updated_video['_id'])
        updated_video['timestamp'] = updated_video['timestamp'].isoformat() if updated_video.get('timestamp') else None
        
        logger.info(f"Updated video ID {id}")
        return JsonResponse({
            "Message": "Video updated successfully",
            "data": updated_video
        })
    except Exception as e:
        logger.error(f"[ERROR] Failed to update video ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 4. DELETE VIDEO ===
@require_http_methods(["DELETE"])
@csrf_exempt
def delete_video(request, id):
    """Delete a video document and associated S3 files (host-only permission)."""
    try:
        # Validate ObjectId format first
        try:
            if len(id) != 24:  # MongoDB ObjectId is always 24 characters
                return JsonResponse({"Error": "Invalid video ID format"}, status=400)
            video_id = ObjectId(id)
        except Exception as e:
            logger.error(f"Invalid video ID format: {id}")
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)

        video = collection.find_one({"_id": video_id})
        if not video:
            logger.error(f"Video ID {id} not found")
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Check if user is host
        user_id = request.GET.get("user_id")
        if not user_id:
            return JsonResponse({"Error": "Missing user_id for authorization"}, status=400)

        meeting_id = video.get("meeting_id")
        if not meeting_id:
            return JsonResponse({"Error": "Invalid video metadata (missing meeting_id)"}, status=400)

        # FIXED: More flexible meeting host check
        is_host = False
        
        # First check if user uploaded this video (direct ownership)
        if str(video.get("user_id", "")) == str(user_id):
            is_host = True
            logger.info(f"User {user_id} is the uploader of video {id}")
        else:
            # Then check if user is host in tbl_Meetings
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                    row = cursor.fetchone()
                    if row:
                        host_id = str(row[0])
                        if host_id == str(user_id):
                            is_host = True
                            logger.info(f"User {user_id} is the host of meeting {meeting_id}")
                    else:
                        # Meeting not found in tbl_Meetings, check if user uploaded the video
                        logger.warning(f"Meeting {meeting_id} not found in tbl_Meetings table")
                        if str(video.get("user_id", "")) == str(user_id):
                            is_host = True
                            logger.info(f"Meeting not in DB but user {user_id} is video uploader - allowing delete")
            except Exception as db_error:
                logger.error(f"Database error checking host: {db_error}")
                # Fallback: allow if user uploaded the video
                if str(video.get("user_id", "")) == str(user_id):
                    is_host = True
                    logger.info(f"DB error but user {user_id} is video uploader - allowing delete")

        if not is_host:
            logger.warning(f"User {user_id} is not authorized to delete video {id}")
            return JsonResponse({"Error": "Only the meeting host or video uploader can delete this recording"}, status=403)

        # Delete associated S3 files
        s3_keys_to_delete = []
        
        # Video, transcript, summary, image, thumbnail URLs
        for url_field in ['video_url', 'transcript_url', 'summary_url', 'image_url', 'thumbnail_url']:
            url = video.get(url_field)
            if url:
                try:
                    # Extract S3 key from URL more reliably
                    if AWS_S3_BUCKET in url:
                        # Format: https://bucket.s3.region.amazonaws.com/folder/file
                        s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    else:
                        # Fallback: take last two parts
                        url_parts = url.split('/')
                        if len(url_parts) >= 2:
                            s3_key = '/'.join(url_parts[-2:])
                        else:
                            continue
                    s3_keys_to_delete.append(s3_key)
                except Exception as e:
                    logger.warning(f"Failed to extract S3 key from {url}: {e}")
        
        # Subtitle URLs
        subtitles = video.get('subtitles', {})
        for lang, url in subtitles.items():
            if url:
                try:
                    if AWS_S3_BUCKET in url:
                        s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    else:
                        url_parts = url.split('/')
                        if len(url_parts) >= 2:
                            s3_key = '/'.join(url_parts[-2:])
                        else:
                            continue
                    s3_keys_to_delete.append(s3_key)
                except Exception as e:
                    logger.warning(f"Failed to extract S3 key from subtitle {url}: {e}")

        # Delete files from S3 (track success/failure)
        deleted_count = 0
        for s3_key in s3_keys_to_delete:
            try:
                if delete_from_s3(s3_key):
                    deleted_count += 1
            except Exception as e:
                logger.warning(f"Failed to delete S3 file {s3_key}: {e}")

        # Delete video document from MongoDB
        delete_result = collection.delete_one({"_id": video_id})
        if delete_result.deleted_count == 0:
            return JsonResponse({"Error": "Failed to delete video document"}, status=500)
            
        logger.info(f"Successfully deleted video document ID {id}, removed {deleted_count}/{len(s3_keys_to_delete)} S3 files")
        return JsonResponse({
            "Message": "Video deleted successfully",
            "deleted_s3_files": deleted_count,
            "total_s3_files": len(s3_keys_to_delete)
        }, status=200)

    except Exception as e:
        logger.error(f"[ERROR] Failed to delete video ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 5. STREAM VIDEO ===
@require_http_methods(["GET", "HEAD", "OPTIONS"])
@csrf_exempt
def stream_video(request, id):
    """Optimized video streaming with automatic URL repair and better error handling."""
    
    # Handle CORS preflight
    if request.method == 'OPTIONS':
        response = HttpResponse()
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Range, Content-Type, Accept, Authorization'
        response['Access-Control-Expose-Headers'] = 'Content-Range, Accept-Ranges, Content-Length'
        response['Access-Control-Max-Age'] = '86400'
        return response
    
    try:
        # Find video document
        try:
            video = collection.find_one({"_id": ObjectId(id)})
        except Exception as id_error:
            logger.error(f"Invalid video ID format: {id}, Error: {id_error}")
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)
            
        if not video:
            logger.warning(f"Video ID {id} not found in MongoDB")
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Access control check
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        meeting_id = video.get("meeting_id", id)
        
        try:
            access_allowed = is_user_allowed(meeting_id, email=email, user_id=user_id)
            if not access_allowed:
                access_allowed = is_user_allowed_debug(meeting_id, email=email, user_id=user_id)
                
            if not access_allowed:
                logger.warning(f"Access denied for user {user_id}/{email} to video {id}")
                return JsonResponse({"Error": "Access denied"}, status=403)
        except Exception as access_error:
            logger.warning(f"Access control check failed: {access_error}")

        # Get video URL and verify/repair if needed
        video_url = video.get("video_url")
        if not video_url:
            logger.error(f"Video URL not found in MongoDB for ID {id}")
            return JsonResponse({"Error": "Video URL not found"}, status=404)

        # Verify URL and attempt repair if broken
        logger.info(f"Verifying video URL for streaming: {id}")
        video = verify_and_repair_video_url(video)
        
        video_url = video.get("video_url")
        if not video_url:
            logger.error(f"Video URL still missing after repair attempt for ID {id}")
            return JsonResponse({"Error": "Video not accessible"}, status=404)

        # SAFETY: Prefer file_path (merged with audio) over rebuilt video_url (may be video-only)
        file_path_key = video.get("file_path")
        if file_path_key:
            fp_size = get_s3_object_size(file_path_key)
            if fp_size > 0:
                current_key = extract_s3_key_from_url(video_url, AWS_S3_BUCKET)
                current_size = get_s3_object_size(current_key) if current_key else 0
                if fp_size > current_size:
                    logger.info(f"📹 Streaming merged file: {file_path_key} ({fp_size} bytes) instead of {current_key} ({current_size} bytes)")
                    video_url = f"https://{AWS_S3_BUCKET}.s3.{AWS_REGION}.amazonaws.com/{file_path_key}"

        # Extract S3 key with improved method
        s3_key = extract_s3_key_from_url(video_url, AWS_S3_BUCKET)
        
        if not s3_key:
            logger.error(f"Failed to extract S3 key from URL: {video_url}")
            return JsonResponse({"Error": "Invalid video URL format"}, status=400)

        logger.info(f"S3 Key extracted: {s3_key}")

        # Check file size
        file_size = get_s3_object_size(s3_key)
        if file_size <= 0:
            logger.error(f"Video file not found or empty in S3: {s3_key}")
            return JsonResponse({"Error": "Video file not accessible in S3"}, status=404)

        logger.info(f"Video file found in S3: {s3_key} ({file_size} bytes)")

        # Determine content type
        file_ext = os.path.splitext(s3_key)[1].lower()
        content_type_map = {
            '.mp4': 'video/mp4',
            '.avi': 'video/x-msvideo',
            '.mov': 'video/quicktime',
            '.wmv': 'video/x-ms-wmv',
            '.webm': 'video/webm',
            '.mkv': 'video/x-matroska'
        }
        content_type = content_type_map.get(file_ext, 'video/mp4')
        logger.info(f"Content type: {content_type}")

        # Handle HEAD requests
        if request.method == 'HEAD':
            response = HttpResponse(content_type=content_type)
            response['Accept-Ranges'] = 'bytes'
            response['Content-Length'] = str(file_size)
            response['Cache-Control'] = 'public, max-age=3600'
            response['Access-Control-Allow-Origin'] = '*'
            logger.info(f"HEAD request served for video {id}")
            return response

        # Handle range requests
        range_header = request.META.get('HTTP_RANGE')
        if range_header:
            import re
            range_match = re.match(r'bytes=(\d+)-(\d*)', range_header)
            if range_match:
                start = int(range_match.group(1))
                end = int(range_match.group(2)) if range_match.group(2) else file_size - 1

                # Validate range
                if start >= file_size:
                    logger.warning(f"Range start {start} exceeds file size {file_size}")
                    response = HttpResponse(status=416)
                    response['Content-Range'] = f'bytes */{file_size}'
                    response['Access-Control-Allow-Origin'] = '*'
                    return response
                    
                if end >= file_size:
                    end = file_size - 1
                    
                if start > end:
                    response = HttpResponse(status=416)
                    response['Content-Range'] = f'bytes */{file_size}'
                    response['Access-Control-Allow-Origin'] = '*'
                    return response

                range_size = end - start + 1
                logger.info(f"Range request: {start}-{end} ({range_size} bytes)")

                # For large ranges, use streaming response
                if range_size > 5 * 1024 * 1024:  # 5MB+
                    def stream_large_range():
                        chunk_size = 2 * 1024 * 1024  # 2MB chunks
                        current = start
                        chunks_sent = 0
                        while current <= end:
                            try:
                                chunk_end = min(current + chunk_size - 1, end)
                                chunk = stream_from_s3(s3_key, current, chunk_end)
                                if chunk:
                                    yield chunk
                                    chunks_sent += 1
                                    current = chunk_end + 1
                                else:
                                    logger.warning(f"Empty chunk received at byte {current}")
                                    break
                            except Exception as chunk_error:
                                logger.error(f"Error streaming chunk: {chunk_error}")
                                break
                        logger.info(f"Streamed {chunks_sent} chunks for range request")
                    
                    response = StreamingHttpResponse(stream_large_range(), status=206, content_type=content_type)
                else:
                    # Small ranges - get all at once
                    content = stream_from_s3(s3_key, start, end)
                    if content is None:
                        logger.error(f"Failed to stream range {start}-{end}")
                        return JsonResponse({"Error": "Failed to stream video range"}, status=500)
                    response = HttpResponse(content, status=206, content_type=content_type)

                response['Content-Range'] = f'bytes {start}-{end}/{file_size}'
                response['Accept-Ranges'] = 'bytes'
                response['Content-Length'] = str(range_size)
                response['Cache-Control'] = 'public, max-age=3600'
                response['Access-Control-Allow-Origin'] = '*'
                response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
                response['Access-Control-Allow-Headers'] = 'Range, Content-Type, Accept'
                response['Access-Control-Expose-Headers'] = 'Content-Range, Accept-Ranges, Content-Length'
                
                return response

        # For full file requests, use streaming response
        def stream_full_file():
            chunk_size = 2 * 1024 * 1024  # 2MB chunks for smooth streaming
            start = 0
            chunks_sent = 0
            while start < file_size:
                try:
                    end = min(start + chunk_size - 1, file_size - 1)
                    chunk = stream_from_s3(s3_key, start, end)
                    if chunk:
                        yield chunk
                        chunks_sent += 1
                        start = end + 1
                    else:
                        logger.warning(f"Empty chunk received at byte {start}")
                        break
                except Exception as chunk_error:
                    logger.error(f"Error streaming full file chunk: {chunk_error}")
                    break
            logger.info(f"Streamed {chunks_sent} chunks for full file request")

        response = StreamingHttpResponse(stream_full_file(), content_type=content_type)
        response['Accept-Ranges'] = 'bytes'
        response['Content-Length'] = str(file_size)
        response['Cache-Control'] = 'public, max-age=3600'
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Range, Content-Type, Accept'
        response['Access-Control-Expose-Headers'] = 'Content-Range, Accept-Ranges, Content-Length'

        logger.info(f"Streaming full video file {id} ({file_size} bytes)")
        return response
        
    except Exception as e:
        logger.error(f"Stream error for video {id}: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": "Internal server error"}, status=500)

# === 6. HANDLE DOCUMENTS (FIXED FOR DOCX - FORCE DOWNLOAD) ===
@require_http_methods(["GET"])
def handle_document(request, id, doc_type):
    """Handle document requests - SUPPORTS BOTH PDF AND DOCX
    
    IMPORTANT: DOCX files cannot be displayed inline in browsers.
    They will always be downloaded, regardless of 'action' parameter.
    Only PDF files can be viewed inline.
    """
    try:
        logger.info(f"Document request - ID: {id}, Type: {doc_type}, Action: {request.GET.get('action', 'view')}")

        action = request.GET.get('action', 'view').lower()
        if action not in ['view', 'download']:
            logger.error(f"Invalid action: {action}")
            return JsonResponse({"Error": "Invalid action. Use 'view' or 'download'"}, status=400)

        try:
            video_id = ObjectId(id)
        except Exception as e:
            logger.error(f"Invalid video ID format: {id}, Error: {e}")
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)

        video = collection.find_one({"_id": video_id})
        if not video:
            logger.error(f"Video ID {id} not found in MongoDB")
            return JsonResponse({"Error": f"Video ID {id} not found"}, status=404)

        # Enhanced access control check
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        meeting_id = video.get("meeting_id", "")
        
        if not is_user_allowed(meeting_id, email, user_id):
            logger.warning(f"Access denied for user {user_id} to document {doc_type} of video {id}")
            return JsonResponse({"Error": "Access denied: You are not authorized to view this document"}, status=403)

        if doc_type not in ["transcript", "summary"]:
            logger.error(f"Invalid document type: {doc_type}")
            return JsonResponse({"Error": "Invalid document type. Use 'transcript' or 'summary'"}, status=400)

        # Get document URL based on type
        if doc_type == "transcript":
            doc_url = video.get("transcript_url")
            doc_name = "Transcript"
        else:  # summary
            doc_url = video.get("summary_url")
            doc_name = "Summary"
            
        if not doc_url:
            logger.error(f"No {doc_type}_url found in video document for ID {id}")
            return JsonResponse({"Error": f"No {doc_name.lower()} document found for ID {id}"}, status=404)

        # Extract S3 key with improved method
        s3_key = extract_s3_key_from_url(doc_url, AWS_S3_BUCKET)
        if not s3_key:
            logger.error(f"Failed to extract S3 key from document URL: {doc_url}")
            return JsonResponse({"Error": f"Invalid {doc_name.lower()} URL format"}, status=400)

        logger.info(f"S3 Key for document: {s3_key}")

        # Check if document exists
        doc_size = get_s3_object_size(s3_key)
        if doc_size <= 0:
            logger.error(f"Document file not found in S3: {s3_key}")
            return JsonResponse({"Error": f"{doc_name} document not accessible in S3"}, status=404)

        # Download content
        try:
            content = stream_from_s3(s3_key)
            if content is None:
                raise Exception("Failed to download from S3")
        except Exception as e:
            logger.error(f"Failed to download {doc_type} document: {e}")
            return JsonResponse({"Error": f"Failed to access {doc_name.lower()} document"}, status=500)

        # ✅ CRITICAL FIX: Auto-detect file format from S3 key
        file_extension = os.path.splitext(s3_key)[1].lower()
        
        if file_extension == '.docx':
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_ext_display = 'docx'
            # ⚠️ FORCE DOWNLOAD FOR DOCX - Browsers cannot display DOCX inline!
            force_download = True
            logger.info(f"Document format: DOCX (forcing download - cannot display inline)")
        elif file_extension == '.pdf':
            content_type = 'application/pdf'
            file_ext_display = 'pdf'
            force_download = False  # PDF can be displayed inline
            logger.info(f"Document format: PDF")
        else:
            # Fallback: check video metadata for document_format field
            doc_format = video.get('document_format', 'pdf')
            if doc_format == 'docx':
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                file_ext_display = 'docx'
                force_download = True  # DOCX cannot be displayed inline
                logger.info(f"Document format from metadata: DOCX (forcing download)")
            else:
                content_type = 'application/pdf'
                file_ext_display = 'pdf'
                force_download = False
                logger.info(f"Document format from metadata: PDF")

        # Build filename for download
        meeting_name = video.get("custom_recording_name") or video.get("filename", "meeting")
        # Clean the meeting name for filename
        clean_name = re.sub(r'[^\w\s-]', '', str(meeting_name)).strip()
        clean_name = re.sub(r'\s+', '_', clean_name)
        if not clean_name:
            clean_name = f"meeting_{id}"
        
        filename = f"{doc_type}_{clean_name}.{file_ext_display}"

        # Determine response based on action AND file type
        # ⚠️ KEY FIX: DOCX files ALWAYS download because browsers can't render them
        if force_download or action == 'download':
            response = HttpResponse(content, content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            log_action = f"Downloaded {doc_name} document for video ID {id} ({file_ext_display})"
        else:
            # Only PDF can be viewed inline
            response = HttpResponse(content, content_type=content_type)
            response['Content-Disposition'] = f'inline; filename="{filename}"'
            log_action = f"Displayed {doc_name} document for video ID {id} ({file_ext_display})"

        # Set response headers
        response['Content-Length'] = str(len(content))
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Accept'
        response['Access-Control-Expose-Headers'] = 'Content-Length, Content-Disposition'
        response['Cache-Control'] = 'public, max-age=3600'

        logger.info(log_action)
        return response

    except Exception as e:
        logger.error(f"[ERROR] Failed to handle document for video ID {id}: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

@require_http_methods(["GET"])
def get_transcript_content(request, id):
    """Return transcript content as JSON for frontend to render."""
    try:
        video_id = ObjectId(id)
        video = collection.find_one({"_id": video_id})
        
        if not video:
            return JsonResponse({"error": "Video not found"}, status=404)

        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        
        if not (email or user_id):
            return JsonResponse({"error": "Missing credentials"}, status=400)

        transcript_url = video.get("transcript_url")
        if not transcript_url:
            return JsonResponse({"error": "Transcript not available"}, status=404)

        s3_key = extract_s3_key_from_url(transcript_url, AWS_S3_BUCKET)
        content = stream_from_s3(s3_key)
        
        if not content:
            return JsonResponse({"error": "Failed to fetch"}, status=500)

        # Parse DOCX and extract text
        from io import BytesIO
        from docx import Document
        
        doc = Document(BytesIO(content))
        
        paragraphs = []
        full_text = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = "heading" if para.style and 'Heading' in para.style.name else "paragraph"
                paragraphs.append({"text": text, "style": style})
                full_text += text + "\n\n"

        return JsonResponse({
            "status": "success",
            "data": {
                "meeting_name": video.get("custom_recording_name") or video.get("filename", "Meeting"),
                "meeting_id": video.get("meeting_id"),
                "date": video.get("timestamp").isoformat() if video.get("timestamp") else None,
                "paragraphs": paragraphs,
                "full_text": full_text.strip(),
                "word_count": len(full_text.split())
            }
        })

    except Exception as e:
        logger.error(f"Get transcript error: {e}")
        return JsonResponse({"error": str(e)}, status=500)

@require_http_methods(["GET"])
def view_mindmap(request, id):
    """Serve the mind map image with relaxed access control."""
    try:
        # Validate ObjectId
        try:
            video_id = ObjectId(id)
        except Exception:
            return JsonResponse({"error": "Invalid video ID format"}, status=400)
        
        video = collection.find_one({"_id": video_id})
        if not video:
            logger.warning(f"Video document not found for ID {id}")
            return JsonResponse({"error": "Video not found"}, status=404)

        # Get request parameters
        email = request.GET.get('email', '').strip().lower()
        user_id = request.GET.get('user_id', '')
        
        # ============ SIMPLIFIED ACCESS CONTROL ============
        # Allow access if ANY of these conditions are true:
        access_granted = False
        access_reason = ""
        
        # 1. User is the video owner (uploaded the video)
        video_owner_id = video.get("user_id")
        if user_id and video_owner_id and str(video_owner_id) == str(user_id):
            access_granted = True
            access_reason = "video_owner"
        
        # 2. Email is in visible_to list
        if not access_granted and email:
            visible_to = video.get("visible_to", [])
            visible_emails = [e.strip().lower() for e in visible_to if e]
            if email in visible_emails:
                access_granted = True
                access_reason = "in_visible_to"
        
        # 3. User has valid credentials (fallback - more permissive)
        if not access_granted and (email or user_id):
            # Allow if user provides valid credentials
            # This is a fallback for cases where access control data is incomplete
            access_granted = True
            access_reason = "authenticated_user"
            logger.info(f"⚠️ Granting access via fallback for user {user_id}/{email} to video {id}")
        
        if not access_granted:
            logger.warning(f"❌ Access denied: No valid credentials provided for video {id}")
            return JsonResponse({"error": "Access denied: Please provide email or user_id"}, status=403)
        
        logger.info(f"✅ Access granted ({access_reason}) for user {user_id}/{email} to mindmap {id}")
        # ============ END ACCESS CONTROL ============

        # Check if mindmap exists
        image_url = video.get("image_url")
        if not image_url:
            logger.warning(f"Mind map not available for video {id}")
            return JsonResponse({
                "error": "Mind map not available",
                "reason": "Mind map was not generated for this recording"
            }, status=404)

        # Extract S3 key
        s3_key = extract_s3_key_from_url(image_url, AWS_S3_BUCKET)
        if not s3_key:
            logger.error(f"Failed to extract S3 key from mindmap URL: {image_url}")
            return JsonResponse({"error": "Invalid mindmap URL format"}, status=400)

        # Check if file exists in S3
        file_size = get_s3_object_size(s3_key)
        if file_size <= 0:
            logger.error(f"Mindmap file not found in S3: {s3_key}")
            return JsonResponse({"error": "Mindmap file not found in S3"}, status=404)

        # Download and serve
        content = stream_from_s3(s3_key)
        if content is None:
            return JsonResponse({"error": "Failed to download mindmap"}, status=500)

        response = HttpResponse(content, content_type='image/png')
        response['Content-Length'] = str(len(content))
        response['Content-Disposition'] = f'inline; filename="mindmap_{id}.png"'
        response['Cache-Control'] = 'public, max-age=3600'
        response['Access-Control-Allow-Origin'] = '*'

        logger.info(f"✅ Served mind map for video {id} ({file_size} bytes)")
        return response
        
    except Exception as e:
        logger.error(f"[ERROR] Mind map error for ID {id}: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return JsonResponse({"error": str(e)}, status=500)

# === 8. NEW: GET SUBTITLES ===
@require_http_methods(["GET"])
def get_subtitles(request, id, lang):
    """Serve subtitle file with improved S3 handling."""
    try:
        video = collection.find_one({"_id": ObjectId(id)})
        if not video:
            logger.warning(f"Video ID {id} not found")
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Access control check
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        if not is_user_allowed(video.get("meeting_id", ""), email, user_id):
            logger.warning(f"Access denied for user {user_id} to subtitles of video {id}")
            return JsonResponse({"Error": "Access denied"}, status=403)

        subtitles = video.get("subtitles", {})
        if lang not in subtitles:
            logger.error(f"Subtitles not available for language: {lang}")
            return JsonResponse({"Error": f"Subtitles not available for language: {lang}"}, status=404)

        subtitle_url = subtitles[lang]
        if not subtitle_url:
            logger.error(f"Subtitle URL is empty for language {lang}")
            return JsonResponse({"Error": "Subtitle URL not found"}, status=404)

        # Extract S3 key with improved method
        s3_key = extract_s3_key_from_url(subtitle_url, AWS_S3_BUCKET)
        if not s3_key:
            logger.error(f"Failed to extract S3 key from subtitle URL: {subtitle_url}")
            return JsonResponse({"Error": "Invalid subtitle URL format"}, status=400)

        logger.info(f"S3 Key for subtitles ({lang}): {s3_key}")

        # Check if file exists
        file_size = get_s3_object_size(s3_key)
        if file_size <= 0:
            logger.error(f"Subtitle file not found in S3: {s3_key}")
            return JsonResponse({"Error": "Subtitle file not accessible"}, status=404)

        # Download content
        try:
            content = stream_from_s3(s3_key)
            if content is None:
                raise Exception("Failed to download from S3")
        except Exception as e:
            logger.error(f"Failed to download subtitle file: {e}")
            return JsonResponse({"Error": "Failed to access subtitle file"}, status=500)

        response = HttpResponse(content, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="subtitles_{lang}_{id}.srt"'
        response['Content-Length'] = str(len(content))
        response['Access-Control-Allow-Origin'] = '*'
        response['Cache-Control'] = 'public, max-age=3600'

        logger.info(f"Served subtitle file for video {id} in language {lang}")
        return response

    except Exception as e:
        logger.error(f"[ERROR] Subtitle serving error: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 9. UPLOAD RECORDING API ===
@require_http_methods(["POST"])
@csrf_exempt
def upload_recording(request):
    """Upload recording file and process it with enhanced features."""
    try:
        # Get form data
        meeting_id = request.POST.get('meeting_id')
        user_id = request.POST.get('user_id')
        
        # Get uploaded file
        if 'recording_file' not in request.FILES:
            return JsonResponse({"Error": "No recording file uploaded"}, status=400)
        
        uploaded_file = request.FILES['recording_file']
        
        if not meeting_id or not user_id:
            return JsonResponse({"Error": "Missing meeting_id or user_id"}, status=400)
        
        # Check if already processed
        existing = collection.find_one({
            "meeting_id": meeting_id, 
            "user_id": user_id, 
            "filename": uploaded_file.name
        })
        
        if existing:
            return JsonResponse({
                "status": "already_processed",
                "file": uploaded_file.name,
                "video_url": existing.get("video_url"),
                "transcript_url": existing.get("transcript_url"),
                "summary_url": existing.get("summary_url"),
                "summary_image_url": existing.get("image_url"),
                "subtitle_urls": existing.get("subtitles"),
                "message": "This video has already been processed."
            })
        
        # Validate file type
        allowed_extensions = ['.mp4', '.webm', '.mkv', '.avi']
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        if file_extension not in allowed_extensions:
            return JsonResponse({"Error": "Invalid file type. Only video files are allowed."}, status=400)
        
        # Save uploaded file temporarily and process
        with TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            
            with open(temp_file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            
            # Calculate file size for routing decision
            file_size = os.path.getsize(temp_file_path)
            
            # For long videos, run processing in background thread
            video_size_mb = file_size / (1024 * 1024)
            
            if video_size_mb > 500:  # Likely > 1 hour video
                import threading
                import shutil
                
                # Copy file to persistent location (TemporaryDirectory will be deleted)
                persistent_dir = "/tmp/long_videos"
                os.makedirs(persistent_dir, exist_ok=True)
                persistent_path = os.path.join(persistent_dir, f"{meeting_id}_{user_id}_{int(time.time())}.mp4")
                shutil.copy(temp_file_path, persistent_path)
                
                def background_process():
                    try:
                        bg_result = process_video_sync(persistent_path, meeting_id, user_id)
                        logging.info(f"✅ Background processing complete: {bg_result.get('status')}")
                    except Exception as bg_err:
                        logging.error(f"❌ Background processing failed: {bg_err}")
                    finally:
                        try:
                            os.remove(persistent_path)
                        except:
                            pass
                
                thread = threading.Thread(target=background_process, daemon=True)
                thread.start()
                
                result = {
                    "status": "processing_in_background",
                    "message": "Long video accepted. Processing in background. Check status in a few minutes.",
                    "meeting_id": meeting_id,
                    "user_id": user_id,
                    "estimated_minutes": int(video_size_mb / 50)
                }
            else:
                logging.info(f"🔄 Starting video processing...")
                result = process_video_sync(temp_file_path, meeting_id, user_id)
                logging.info(f"✅ Video processing completed: {result}")
            
            return JsonResponse(result)
                
    except Exception as e:
        logger.error(f"[ERROR] Recording upload failed: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 10. UPLOAD RECORDING BLOB ===
@require_http_methods(["POST"])
@csrf_exempt
def upload_recording_blob(request):
    """Upload recording blob and process it with enhanced error handling."""
    try:
        # Debug request data
        logging.info(f"📥 Recording upload request received")
        logging.info(f"📄 Request FILES: {list(request.FILES.keys())}")
        logging.info(f"📄 Request POST: {dict(request.POST)}")
        
        # Validate required parameters
        if 'recording_blob' not in request.FILES:
            logging.error("❌ No recording_blob in request.FILES")
            return JsonResponse({
                "status": "error",
                "error": "No recording blob provided",
                "Error": "No recording blob provided"
            }, status=400)
        
        recording_file = request.FILES['recording_blob']
        meeting_id = request.POST.get('meeting_id')
        user_id = request.POST.get('user_id')
        recording_id = request.POST.get('recording_id')
        
        logging.info(f"📹 Processing upload: meeting_id={meeting_id}, user_id={user_id}, file_size={recording_file.size}")
        
        if not all([meeting_id, user_id]):
            return JsonResponse({
                "status": "error",
                "error": "Missing required parameters: meeting_id and user_id",
                "Error": "Missing required parameters"
            }, status=400)
        
        if recording_file.size == 0:
            return JsonResponse({
                "status": "error", 
                "error": "Recording file is empty",
                "Error": "Recording file is empty"
            }, status=400)
        
        # Update recording status to "uploading"
        try:
            if recording_id:
                collection.update_one(
                    {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                    {"$set": {
                        "recording_status": "uploading",
                        "upload_start_time": datetime.now(),
                        "file_size": recording_file.size
                    }}
                )
        except Exception as update_error:
            logging.warning(f"⚠ Failed to update recording status: {update_error}")
        
        # Process the recording
        with TemporaryDirectory() as temp_dir:
            # Determine file extension
            file_ext = '.webm'
            if recording_file.content_type:
                if 'mp4' in recording_file.content_type:
                    file_ext = '.mp4'
                elif 'webm' in recording_file.content_type:
                    file_ext = '.webm'
            
            temp_file_path = os.path.join(temp_dir, f"recording_{meeting_id}{file_ext}")
            
            logging.info(f"💾 Saving recording to: {temp_file_path}")
            
            # Save the uploaded file
            with open(temp_file_path, 'wb+') as destination:
                for chunk in recording_file.chunks():
                    destination.write(chunk)
            
            file_size = os.path.getsize(temp_file_path)
            logging.info(f"✅ File saved successfully, size: {file_size} bytes")
            
            if file_size == 0:
                return JsonResponse({
                    "status": "error",
                    "error": "Uploaded file is empty after saving",
                    "Error": "Upload failed - empty file"
                }, status=400)
            
            # Update status to "processing"
            try:
                if recording_id:
                    collection.update_one(
                        {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                        {"$set": {
                            "recording_status": "processing",
                            "processing_start_time": datetime.now()
                        }}
                    )
            except Exception:
                pass
            
            # Decide: background thread for long videos, synchronous for short ones
            video_size_mb = file_size / (1024 * 1024)
            
            if video_size_mb > 500:  # Likely > 1 hour video
                import threading
                import shutil
                
                # Copy file to persistent location (TemporaryDirectory will be deleted)
                persistent_dir = "/tmp/long_videos"
                os.makedirs(persistent_dir, exist_ok=True)
                persistent_path = os.path.join(persistent_dir, f"{meeting_id}_{user_id}_{int(time.time())}.mp4")
                shutil.copy(temp_file_path, persistent_path)
                
                def background_process():
                    try:
                        bg_result = process_video_sync(persistent_path, meeting_id, user_id)
                        logging.info(f"✅ Background processing complete: {bg_result.get('status')}")
                    except Exception as bg_err:
                        logging.error(f"❌ Background processing failed: {bg_err}")
                    finally:
                        try:
                            os.remove(persistent_path)
                        except:
                            pass
                
                thread = threading.Thread(target=background_process, daemon=True)
                thread.start()
                
                result = {
                    "status": "processing_in_background",
                    "message": "Long video accepted. Processing in background. Check status in a few minutes.",
                    "meeting_id": meeting_id,
                    "user_id": user_id,
                    "estimated_minutes": int(video_size_mb / 50)
                }
            else:
                logging.info(f"🔄 Starting video processing...")
                result = process_video_sync(temp_file_path, meeting_id, user_id)
                logging.info(f"✅ Video processing completed: {result.get('status')}")
            
            # Update final recording metadata
            if recording_id and result.get("status") == "success":
                try:
                    final_metadata = {
                        "recording_status": "completed",
                        "processing_completed_time": datetime.now(),
                        "video_url": result.get("video_url"),
                        "transcript_url": result.get("transcript_url"),
                        "summary_url": result.get("summary_url"),
                        "image_url": result.get("summary_image_url"),
                        "subtitles": result.get("subtitle_urls", {}),
                        "final_file_size": file_size,
                        "original_filename": recording_file.name,
                        "processing_notes": result.get("processing_notes", {}),
                        "upload_successful": True
                    }
                    
                    collection.update_one(
                        {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                        {"$set": final_metadata}
                    )
                    
                    logging.info(f"✅ Final metadata updated for recording: {recording_id}")
                    
                except Exception as final_update_error:
                    logging.warning(f"⚠ Failed to update final metadata: {final_update_error}")
            
            # Add additional metadata to result
            result.update({
                "meeting_id": meeting_id,
                "user_id": user_id,
                "original_filename": recording_file.name,
                "file_size": file_size,
                "upload_timestamp": datetime.now().isoformat(),
                "recording_id": recording_id,
                "upload_successful": True,
                "processing_completed": True
            })
            
            logging.info(f"✅ Recording upload and processing completed successfully")
            return JsonResponse(result)
                
    except Exception as e:
        logging.error(f"❌ Recording blob upload failed: {e}")
        import traceback
        logging.error(f"❌ Full traceback: {traceback.format_exc()}")
        
        # Update recording status to "failed" if recording_id available
        try:
            recording_id = request.POST.get('recording_id')
            if recording_id:
                collection.update_one(
                    {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                    {"$set": {
                        "recording_status": "failed",
                        "error_message": str(e),
                        "error_timestamp": datetime.now()
                    }}
                )
        except Exception:
            pass
        
        return JsonResponse({
            "status": "error",
            "error": str(e),
            "Error": f"Server error: {str(e)}",
            "upload_successful": False
        }, status=500)
              
# === 11. START RECORDING WITH METADATA ===
@require_http_methods(["POST"])
@csrf_exempt
def start_recording_with_metadata(request, id):
    """Start recording and store initial metadata."""
    try:
        # Check if meeting exists
        video = collection.find_one({"_id": ObjectId(id)})
        if not video:
            # Check if it's a meeting record instead
            video = collection.find_one({"meeting_id": id})
        
        if not video:
            logger.error(f"Meeting ID {id} not found")
            return JsonResponse({"Error": "Meeting not found"}, status=404)

        # Get user from request or use default
        data = json.loads(request.body) if request.body else {}
        user_id = data.get('user_id', video.get('user_id', 'unknown'))

        # Create initial recording metadata
        recording_metadata = {
            "meeting_id": id,
            "user_id": user_id,
            "recording_status": "active",
            "start_time": datetime.now(),
            "video_url": None,
            "transcript_url": None,
            "summary_url": None,
            "image_url": None,
            "subtitles": {},
            "file_size": 0,
            "duration": 0
        }

        # Insert or update recording metadata
        existing_recording = collection.find_one({"meeting_id": id, "recording_status": "active"})
        if existing_recording:
            collection.update_one(
                {"_id": existing_recording["_id"]},
                {"$set": recording_metadata}
            )
            recording_id = str(existing_recording["_id"])
        else:
            result = collection.insert_one(recording_metadata)
            recording_id = str(result.inserted_id)

        logger.info(f"Recording started for meeting {id} with metadata ID {recording_id}")
        
        return JsonResponse({
            "Message": "Recording started with metadata",
            "recording_id": recording_id,
            "meeting_id": id
        }, status=200)
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to start recording with metadata for meeting ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 12. STOP RECORDING AND FINALIZE ===
@require_http_methods(["POST"])
@csrf_exempt
def stop_recording_and_finalize(request, id):
    """Stop recording and prepare for upload."""
    try:
        # Find active recording
        recording = collection.find_one({"meeting_id": id, "recording_status": "active"})
        if not recording:
            logger.error(f"No active recording found for meeting ID {id}")
            return JsonResponse({"Error": "No active recording found"}, status=404)

        # Update recording status
        collection.update_one(
            {"_id": recording["_id"]},
            {"$set": {
                "recording_status": "stopped",
                "end_time": datetime.now()
            }}
        )

        logger.info(f"Recording stopped for meeting {id}")
        
        return JsonResponse({
            "Message": "Recording stopped successfully",
            "recording_id": str(recording["_id"]),
            "meeting_id": id,
            "ready_for_upload": True
        }, status=200)
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to stop recording for meeting ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 13. GET RECORDING STATUS ===
@require_http_methods(["GET"])
def get_recording_status(request, meeting_id):
    """Get recording status for a meeting."""
    try:
        recording = collection.find_one({"meeting_id": meeting_id}, sort=[("start_time", -1)])
        
        if not recording:
            return JsonResponse({"Error": "No recording found for this meeting"}, status=404)
        
        # Convert ObjectId to string
        recording['_id'] = str(recording['_id'])
        if recording.get('start_time'):
            recording['start_time'] = recording['start_time'].isoformat()
        if recording.get('end_time'):
            recording['end_time'] = recording['end_time'].isoformat()
        if recording.get('upload_time'):
            recording['upload_time'] = recording['upload_time'].isoformat()
        
        return JsonResponse({
            "status": "success",
            "recording": recording
        })
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to get recording status for meeting {meeting_id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 14. PROCESS MEETING RECORDING ===
def process_meeting_recording(video_blob, meeting_id, user_id):
    """Process meeting recording blob and upload to S3."""
    try:
        with TemporaryDirectory() as workdir:
            # Save blob to temporary file
            temp_video_path = os.path.join(workdir, f"meeting_recording_{meeting_id}.mp4")
            
            with open(temp_video_path, 'wb') as f:
                f.write(video_blob)
            
            # Use the synchronous process_video function
            result = process_video_sync(temp_video_path, meeting_id, user_id)
            
            logger.info(f"✅ Meeting recording processed successfully for meeting {meeting_id}")
            return result
            
    except Exception as e:
        logger.error(f"❌ Failed to process meeting recording: {e}")
        raise e

# Replace the existing Start_Recording function with this:
@require_http_methods(["POST"])
@csrf_exempt
def Start_Recording(request, id):
    """Start LiveKit stream recording - calls recording service via HTTP"""
    try:
        # Parse request body
        recording_settings = {}
        if request.body:
            try:
                recording_settings = json.loads(request.body)
            except json.JSONDecodeError:
                pass

        with connection.cursor() as cursor:
            select_query = """
            SELECT Host_ID, Is_Recording_Enabled, Meeting_Name
            FROM tbl_Meetings
            WHERE ID = %s
            """
            cursor.execute(select_query, [id])
            row = cursor.fetchone()
            if not row:
                logging.error(f"Meeting ID {id} not found")
                return JsonResponse({"Error": "Meeting not found"}, status=404)

            host_id, is_recording_enabled, meeting_name = row
            
            # Check if recording is already active
            if is_recording_enabled:
                logging.info(f"Recording is already active for meeting {id}")
                return JsonResponse({
                    "Message": "Recording is already active",
                    "success": True,
                    "already_recording": True,
                    "is_recording": True,
                    "meeting_id": id,
                    "recording_type": "livekit_stream"
                }, status=200)

            # ✅ NEW: Call recording service via HTTP
            try:
                room_name = recording_settings.get('room_name', f"meeting_{id}")
                
                # Make HTTP request to recording service
                response = requests.post(
                    f"{RECORDING_SERVICE_URL}/start",
                    json={
                        "meeting_id": id,
                        "host_user_id": str(host_id),
                        "room_name": room_name
                    },
                    timeout=30
                )
                
                if response.status_code == 200:
                    result = response.json()
                    
                    # Update database
                    started_at = timezone.now()
                    cursor.execute(
                        "UPDATE tbl_Meetings SET Is_Recording_Enabled = 1, Started_At = %s WHERE ID = %s",
                        [started_at, id]
                    )
                    
                    logging.info(f"Stream recording started for meeting {id}")
                    
                    return JsonResponse({
                        "Message": "Stream recording started successfully",
                        "success": True,
                        "is_recording": True,
                        "meeting_id": id,
                        "recording_id": result.get("recording_id"),
                        "session_id": result.get("session_id"),  # ✅ Pass session_id to frontend
                        "recording_type": "livekit_stream",
                        "service_mode": "standalone",
                        **result
                    })
                else:
                    error_msg = response.json().get("message", "Recording service error")
                    return JsonResponse({
                        "Error": error_msg,
                        "success": False,
                        "meeting_id": id
                    }, status=500)
                    
            except requests.exceptions.Timeout:
                return JsonResponse({
                    "Error": "Recording service timeout",
                    "success": False,
                    "meeting_id": id
                }, status=504)
            except requests.exceptions.ConnectionError:
                return JsonResponse({
                    "Error": "Cannot connect to recording service",
                    "success": False,
                    "meeting_id": id
                }, status=503)
            except Exception as e:
                logging.error(f"Error calling recording service: {e}")
                return JsonResponse({
                    "Error": f"Recording service error: {str(e)}",
                    "success": False,
                    "meeting_id": id
                }, status=500)
            
    except Exception as e:
        logging.error(f"Database error: {e}")
        return JsonResponse({
            "Error": f"Database error: {str(e)}",
            "success": False,
            "meeting_id": id
        }, status=500)

# Replace the existing Stop_Recording function with this:
@require_http_methods(["POST"])
@csrf_exempt
def Stop_Recording(request, id):
    """Stop LiveKit stream recording - calls recording service via HTTP"""
    try:
        # Get meeting info
        with connection.cursor() as cursor:
            cursor.execute("""
            SELECT Host_ID, Is_Recording_Enabled, Meeting_Name
            FROM tbl_Meetings
            WHERE ID = %s
            """, [id])
            
            row = cursor.fetchone()
            if not row:
                return JsonResponse({"Error": "Meeting not found"}, status=404)

            host_id, is_recording_enabled, meeting_name = row
            
            if not is_recording_enabled:
                return JsonResponse({
                    "Message": "Recording was not active",
                    "meeting_id": id,
                    "is_recording": False,
                    "success": True
                }, status=200)

        # ✅ NEW: Call recording service via HTTP
        try:
            # ✅ FIX: Accept session_id from frontend to stop the correct recording
            stop_data = {}
            if request.body:
                try:
                    stop_data = json.loads(request.body)
                except json.JSONDecodeError:
                    pass
            
            session_id = stop_data.get('session_id')
            
            response = requests.post(
                f"{RECORDING_SERVICE_URL}/stop",
                json={
                    "meeting_id": id,
                    "session_id": session_id  # ✅ Tell recording service WHICH recording to stop
                },
                timeout=60
            )
            
            # Update database immediately
            ended_at = timezone.now()
            with connection.cursor() as cursor:
                cursor.execute("""
                UPDATE tbl_Meetings
                SET Is_Recording_Enabled = 0, Ended_At = %s
                WHERE ID = %s
                """, [ended_at, id])
            
            if response.status_code == 200:
                result = response.json()
                logging.info(f"Recording stopped for meeting {id}")
                
                return JsonResponse({
                    "Message": "Recording stopped successfully",
                    "success": True,
                    "meeting_id": id,
                    "meeting_name": meeting_name,
                    "is_recording": False,
                    "service_mode": "standalone",
                    **result
                })
            else:
                error_msg = response.json().get("message", "Recording service error")
                return JsonResponse({
                    "Error": error_msg,
                    "success": False,
                    "meeting_id": id
                }, status=500)
                
        except requests.exceptions.Timeout:
            return JsonResponse({
                "Error": "Recording service timeout",
                "success": False,
                "meeting_id": id
            }, status=504)
        except requests.exceptions.ConnectionError:
            return JsonResponse({
                "Error": "Cannot connect to recording service",
                "success": False,
                "meeting_id": id
            }, status=503)
        except Exception as e:
            logging.error(f"Error calling recording service: {e}")
            return JsonResponse({
                "Error": f"Recording service error: {str(e)}",
                "success": False,
                "meeting_id": id
            }, status=500)

    except Exception as e:
        logging.error(f"Critical failure: {e}")
        return JsonResponse({
            "Error": f"Critical error: {str(e)}", 
            "success": False,
            "meeting_id": id
        }, status=500)
           
# === 17. NEW: UPLOAD SINGLE FILE (FROM FASTAPI) ===
@require_http_methods(["POST"])
@csrf_exempt
def upload_single_file(request):
    """Upload and process a single video file (FastAPI equivalent)."""
    try:
        # Get form data
        meeting_id = request.POST.get('meeting_id')
        user_id = request.POST.get('user_id')
        
        if 'file' not in request.FILES:
            return JsonResponse({"error": "No file uploaded"}, status=400)
        
        uploaded_file = request.FILES['file']
        
        if not meeting_id or not user_id:
            return JsonResponse({"error": "Missing meeting_id or user_id"}, status=400)
        
        # Check if already processed
        existing = collection.find_one({
            "meeting_id": meeting_id, 
            "user_id": user_id, 
            "filename": uploaded_file.name
        })
        
        if existing:
            return JsonResponse({
                "status": "already_processed",
                "file": uploaded_file.name,
                "video_url": existing.get("video_url"),
                "transcript_url": existing.get("transcript_url"),
                "summary_url": existing.get("summary_url"),
                "summary_image_url": existing.get("image_url"),
                "subtitle_urls": existing.get("subtitles"),
                "message": "This video has already been processed."
            })
        
        # Save uploaded file temporarily and process
        with TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            
            with open(temp_file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            
            # Process the recording
            result = process_video_sync(temp_file_path, meeting_id, user_id)
            result["file"] = uploaded_file.name
            result["meeting_id"] = meeting_id
            
            return JsonResponse(result)
                
    except Exception as e:
        logger.error(f"[ERROR] Single file upload failed: {e}")
        return JsonResponse({"error": f"Server error: {str(e)}"}, status=500)

# === 18. NEW: LIST ALL VIDEOS WITH METADATA ===
@require_http_methods(["GET"])
def list_all_videos_detailed(request):
    """List all videos with detailed metadata."""
    try:
        # Query parameters
        page = int(request.GET.get('page', 1))
        limit = int(request.GET.get('limit', 20))
        meeting_id = request.GET.get('meeting_id')
        user_id = request.GET.get('user_id')
        
        # Build query filter
        query_filter = {}
        if meeting_id:
            query_filter['meeting_id'] = meeting_id
        if user_id:
            query_filter['user_id'] = user_id
        
        # Calculate skip value for pagination
        skip = (page - 1) * limit
        
        # Fetch videos
        videos = list(collection.find(query_filter).sort("timestamp", -1).skip(skip).limit(limit))
        total_count = collection.count_documents(query_filter)
        
        # Process videos
        for video in videos:
            video['_id'] = str(video['_id'])
            video['timestamp'] = video['timestamp'].isoformat() if video.get('timestamp') else None
            if video.get('start_time'):
                video['start_time'] = video['start_time'].isoformat()
            if video.get('end_time'):
                video['end_time'] = video['end_time'].isoformat()
            if video.get('upload_time'):
                video['upload_time'] = video['upload_time'].isoformat()
            
            # ✅ CRITICAL FIX: Compute subtitles_available field
            subtitles_dict = video.get('subtitles', {})
            video['subtitles_available'] = bool(subtitles_dict and len(subtitles_dict) > 0)

        return JsonResponse({
            "status": "success",
            "data": videos,
            "pagination": {
                "page": page,
                "limit": limit,
                "total": total_count,
                "total_pages": (total_count + limit - 1) // limit
            }
        })
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to list videos: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === MEETING TRASH MANAGEMENT FUNCTIONS (USING EXISTING TABLES ONLY) ===
@require_http_methods(["DELETE"])
@csrf_exempt
def move_video_to_trash(request, id):
    """Move a video recording to trash."""
    try:
        # Validate ObjectId format
        try:
            if len(id) != 24:
                return JsonResponse({"Error": "Invalid video ID format"}, status=400)
            video_id = ObjectId(id)
        except Exception:
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)

        # Find video in MongoDB
        video = collection.find_one({"_id": video_id})
        if not video:
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Check if already trashed
        if video.get("is_trashed"):
            return JsonResponse({"Error": "Video is already in trash"}, status=400)

        # Check permissions
        user_id = request.GET.get("user_id")
        email = request.GET.get("email", "")
        
        if not user_id:
            return JsonResponse({"Error": "Missing user_id"}, status=400)

        meeting_id = video.get("meeting_id")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "Permission denied: Only the meeting host can delete this recording"}, status=403)

        # Move to trash
        trash_date = datetime.now()
        permanent_delete_date = trash_date + timedelta(days=TRASH_RETENTION_DAYS)
        
        collection.update_one(
            {"_id": video_id},
            {"$set": {
                "is_trashed": True,
                "trashed_at": trash_date,
                "permanent_delete_at": permanent_delete_date
            }}
        )

        logger.info(f"Video {id} moved to trash by user {user_id}")

        return JsonResponse({
            "Message": "Recording moved to trash successfully",
            "video_id": id,
            "trashed": True,
            "trashed_at": trash_date.isoformat(),
            "permanent_delete_at": permanent_delete_date.isoformat()
        })

    except Exception as e:
        logger.error(f"Failed to move video to trash: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

@require_http_methods(["GET"])
def list_trash_videos(request):
    """List trashed videos for the user - HOST ONLY ACCESS."""
    try:
        user_id = request.GET.get("user_id")
        
        if not user_id:
            return JsonResponse({"Error": "Missing user_id"}, status=400)

        # Get trashed videos where user is the HOST/OWNER only
        trashed_videos = list(collection.find({
            "is_trashed": True,
            "user_id": user_id  # Only show videos uploaded by this user (host)
        }).sort("trashed_at", -1))

        # Format response
        for video in trashed_videos:
            video["_id"] = str(video["_id"])
            for date_field in ["trashed_at", "permanent_delete_at", "timestamp"]:
                if date_field in video and video[date_field]:
                    video[date_field] = video[date_field].isoformat()

        return JsonResponse({
            "videos": trashed_videos,
            "total": len(trashed_videos),
            "page": 1,
            "pages": 1
        })

    except Exception as e:
        logger.error(f"Failed to list trash videos: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

@require_http_methods(["POST"])
@csrf_exempt  
def restore_video(request, id):
    """Restore video from trash."""
    try:
        video_id = ObjectId(id)
        video = collection.find_one({"_id": video_id})
        
        if not video:
            return JsonResponse({"Error": "Video not found"}, status=404)
            
        if not video.get("is_trashed"):
            return JsonResponse({"Error": "Video is not in trash"}, status=400)

        # Check permissions
        user_id = request.GET.get("user_id")
        email = request.GET.get("email", "")
        
        meeting_id = video.get("meeting_id")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "Permission denied"}, status=403)

        # Restore video
        collection.update_one(
            {"_id": video_id},
            {"$unset": {
                "is_trashed": "",
                "trashed_at": "",
                "permanent_delete_at": ""
            }}
        )

        return JsonResponse({
            "Message": "Recording restored successfully",
            "video_id": id,
            "restored": True
        })

    except Exception as e:
        logger.error(f"Failed to restore video: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

@require_http_methods(["DELETE"])
@csrf_exempt
def permanent_delete_video(request, id):
    """Permanently delete a video and its S3 files."""
    try:
        video_id = ObjectId(id)
        video = collection.find_one({"_id": video_id})
        
        if not video:
            return JsonResponse({"Error": "Video not found"}, status=404)
            
        if not video.get("is_trashed"):
            return JsonResponse({"Error": "Video must be in trash first"}, status=400)

        # Check permissions
        user_id = request.GET.get("user_id")
        email = request.GET.get("email", "")
        
        meeting_id = video.get("meeting_id")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "Permission denied"}, status=403)

        # Delete S3 files
        s3_keys_to_delete = []
        for url_field in ['video_url', 'transcript_url', 'summary_url', 'image_url', 'thumbnail_url']:
            url = video.get(url_field)
            if url and AWS_S3_BUCKET in url:
                try:
                    s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    s3_keys_to_delete.append(s3_key)
                except:
                    pass

        # Delete subtitle files
        subtitles = video.get('subtitles', {})
        for lang, url in subtitles.items():
            if url and AWS_S3_BUCKET in url:
                try:
                    s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    s3_keys_to_delete.append(s3_key)
                except:
                    pass

        # Delete S3 files
        deleted_count = 0
        for s3_key in s3_keys_to_delete:
            if delete_from_s3(s3_key):
                deleted_count += 1

        # Delete from MongoDB
        collection.delete_one({"_id": video_id})

        return JsonResponse({
            "Message": "Recording permanently deleted",
            "video_id": id,
            "deleted_s3_files": deleted_count,
            "total_s3_files": len(s3_keys_to_delete)
        })

    except Exception as e:
        logger.error(f"Failed to permanently delete video: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === NEW: STORE CUSTOM RECORDING NAME ===
@require_http_methods(["POST"])
@csrf_exempt
def store_custom_recording_name(request):
    """
    Store custom recording name immediately when user enters it.
    This happens while recording is still processing in background.
    """
    try:
        data = json.loads(request.body)
        meeting_id = data.get('meeting_id')
        custom_name = data.get('custom_name')
        user_id = data.get('user_id')  # For permission checking
        recording_timestamp = data.get('recording_timestamp')  # ✅ NEW: To identify specific recording

        if not meeting_id or not custom_name:
            return JsonResponse({
                "status": "error",
                "error": "Missing meeting_id or custom_name"
            }, status=400)
        
        # Validate custom name
        custom_name = custom_name.strip()
        if len(custom_name) == 0:
            return JsonResponse({
                "status": "error",
                "error": "Recording name cannot be empty"
            }, status=400)
        
        if len(custom_name) > 200:
            return JsonResponse({
                "status": "error",
                "error": "Recording name too long (max 200 characters)"
            }, status=400)
        
        # Optional: Check if user has permission for this meeting
        if user_id:
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                    row = cursor.fetchone()
                    if row and str(row[0]) != str(user_id):
                        # Not the host, check if participant
                        cursor.execute(
                            "SELECT COUNT(*) FROM tbl_Participants WHERE Meeting_ID = %s AND User_ID = %s",
                            [meeting_id, user_id]
                        )
                        if cursor.fetchone()[0] == 0:
                            return JsonResponse({
                                "status": "error",
                                "error": "You don't have permission to name this recording"
                            }, status=403)
            except Exception as perm_error:
                logger.warning(f"Permission check failed: {perm_error}")
        
        # ✅ FIX: Accept session_id from frontend (got it from start-recording response)
        # Only generate random one if frontend didn't provide it (legacy support)
        session_id = data.get('session_id')
        if not session_id:
            import uuid
            session_id = str(uuid.uuid4())[:8]
            logger.warning(f"No session_id provided by frontend, generated fallback: {session_id}")

        # Store the custom name with a pending flag
        custom_name_document = {
            "meeting_id": meeting_id,
            "custom_recording_name": custom_name,
            "pending_name_update": True,
            "custom_name_session_id": session_id,  # ✅ Linked to actual recording session
            "recording_timestamp": recording_timestamp or int(time.time()),
            "name_stored_at": datetime.now(),
            "user_id": user_id
        }

        # Always insert new document (one per recording session)
        result = collection.insert_one(custom_name_document)
        logger.info(f"Stored custom name for meeting {meeting_id}, session {session_id}: {custom_name}")

        return JsonResponse({
            "status": "success",
            "message": "Recording name stored successfully",
            "custom_name": custom_name,
            "meeting_id": meeting_id,
            "session_id": session_id
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            "status": "error",
            "error": "Invalid JSON in request body"
        }, status=400)
    except Exception as e:
        logger.error(f"Error storing custom recording name: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return JsonResponse({
            "status": "error",
            "error": f"Server error: {str(e)}"
        }, status=500)

@require_http_methods(["GET"])
def get_trainer_evaluation_doc(request, id):
    """Download trainer evaluation DOCX. HOST ONLY."""
    try:
        try:
            video_id = ObjectId(id)
        except Exception:
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)

        video = collection.find_one({"_id": video_id})
        if not video:
            return JsonResponse({"Error": "Video not found"}, status=404)

        user_id = request.GET.get('user_id', '')
        if not user_id:
            return JsonResponse({"Error": "Missing user_id"}, status=400)

        meeting_id = video.get("meeting_id", "")

        # Strict host-only check
        is_host = False
        if str(video.get("user_id", "")) == str(user_id):
            is_host = True
        else:
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                    row = cursor.fetchone()
                    if row and str(row[0]) == str(user_id):
                        is_host = True
            except Exception as db_err:
                logger.error(f"Host check DB error: {db_err}")

        if not is_host:
            return JsonResponse(
                {"Error": "Access denied: trainer evaluation is host-only"},
                status=403
            )

        eval_url = video.get("trainer_evaluation_url")
        if not eval_url:
            return JsonResponse({"Error": "Trainer evaluation not available"}, status=404)

        s3_key = extract_s3_key_from_url(eval_url, AWS_S3_BUCKET)
        if not s3_key:
            return JsonResponse({"Error": "Invalid evaluation URL"}, status=400)

        content = stream_from_s3(s3_key)
        if content is None:
            return JsonResponse({"Error": "Failed to fetch evaluation"}, status=500)

        response = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="trainer_evaluation_{id}.docx"'
        response['Content-Length'] = str(len(content))
        response['Cache-Control'] = 'private, no-cache'
        return response

    except Exception as e:
        logger.error(f"Trainer evaluation download error: {e}")
        return JsonResponse({"Error": str(e)}, status=500)


@require_http_methods(["GET"])
def get_summary_content(request, id):
    """Return summary content with trainer evaluation as JSON for frontend to render."""
    try:
        video_id = ObjectId(id)
        video = collection.find_one({"_id": video_id})
        
        if not video:
            return JsonResponse({"error": "Video not found"}, status=404)

        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        
        if not (email or user_id):
            return JsonResponse({"error": "Missing credentials"}, status=400)

        # Check access permission
        meeting_id = video.get("meeting_id", "")
        if not is_user_allowed(meeting_id, email, user_id):
            return JsonResponse({"error": "Access denied"}, status=403)

        # Check if user is the host (trainer) — only host sees evaluation
        is_host = False
        if user_id:
            # Check 1: Is user the video uploader (recorder/host)?
            if str(video.get("user_id", "")) == str(user_id):
                is_host = True
            else:
                # Check 2: Is user the meeting host in database?
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                        row = cursor.fetchone()
                        if row and str(row[0]) == str(user_id):
                            is_host = True
                except Exception:
                    pass

        summary_url = video.get("summary_url")
        summary_text = video.get("summary_text", "")
        trainer_evaluation = video.get("trainer_evaluation", {})
        image_url = video.get("image_url")  # Mind map image
        
        paragraphs = []
        full_text = ""
        
        # Try to get summary from S3 DOCX file
        if summary_url:
            try:
                s3_key = extract_s3_key_from_url(summary_url, AWS_S3_BUCKET)
                content = stream_from_s3(s3_key)
                
                if content:
                    # Parse DOCX and extract text
                    from io import BytesIO
                    from docx import Document
                    
                    doc = Document(BytesIO(content))
                    
                    for para in doc.paragraphs:
                        text = para.text.strip()
                        if text:
                            # Detect style
                            style = "paragraph"
                            if para.style and para.style.name:
                                if 'Heading 1' in para.style.name:
                                    style = "heading1"
                                elif 'Heading 2' in para.style.name:
                                    style = "heading2"
                                elif 'Heading 3' in para.style.name:
                                    style = "heading3"
                                elif 'Heading' in para.style.name:
                                    style = "heading"
                            
                            paragraphs.append({"text": text, "style": style})
                            full_text += text + "\n\n"
                    
                    logger.info(f"✅ Summary loaded from S3 DOCX for video {id}")
                    
            except Exception as s3_error:
                logger.warning(f"⚠️ Failed to load summary from S3: {s3_error}")
        
        # Fallback: Use stored summary_text if S3 fetch failed
        if not full_text and summary_text:
            full_text = summary_text
            for line in summary_text.split('\n'):
                line = line.strip()
                if line:
                    # Simple heading detection from markdown-style text
                    style = "paragraph"
                    if line.startswith('### '):
                        style = "heading3"
                        line = line[4:]
                    elif line.startswith('## '):
                        style = "heading2"
                        line = line[3:]
                    elif line.startswith('# '):
                        style = "heading1"
                        line = line[2:]
                    
                    paragraphs.append({"text": line, "style": style})
            
            logger.info(f"✅ Summary loaded from MongoDB summary_text for video {id}")
        
        # If still no summary available
        if not full_text:
            return JsonResponse({"error": "Summary not available"}, status=404)

        # Build response with all data
        response_data = {
            "status": "success",
            "data": {
                # Meeting info
                "meeting_name": video.get("custom_recording_name") or video.get("filename", "Meeting"),
                "meeting_id": video.get("meeting_id"),
                "date": video.get("timestamp").isoformat() if video.get("timestamp") else None,
                "duration": video.get("duration", 0),
                
                # Summary content
                "paragraphs": paragraphs,
                "full_text": full_text.strip(),
                "word_count": len(full_text.split()),
                
                # Mind map image
                "image_url": image_url,
                "has_mindmap": bool(image_url),
                
                # Trainer evaluation scores — ONLY visible to host (trainer)
                "trainer_evaluation": {
                    "technical_content": trainer_evaluation.get("technical_content", 0),
                    "explanation_clarity": trainer_evaluation.get("explanation_clarity", 0),
                    "friendliness": trainer_evaluation.get("friendliness", 0),
                    "communication": trainer_evaluation.get("communication", 0),
                    "overall_feedback": trainer_evaluation.get("overall_feedback", ""),
                    "has_evaluation": bool(trainer_evaluation and not trainer_evaluation.get("error")),
                    "download_url": video.get("trainer_evaluation_url")
                } if is_host else None,
                
                # Additional metadata
                "user_name": video.get("user_name", ""),
                "transcription_engine": video.get("transcription_engine", ""),
                "document_format": video.get("document_format", "docx"),
            }
        }

        return JsonResponse(response_data)

    except Exception as e:
        logger.error(f"Get summary content error: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return JsonResponse({"error": str(e)}, status=500)

@require_http_methods(["GET"])
@csrf_exempt
def health_check(request):
    """Lightweight health check endpoint - NO DATABASE QUERIES"""
    return JsonResponse({
        "status": "healthy",
        "service": "imeetpro-backend",
        "timestamp": datetime.now().isoformat()
    })


# Alias for root-level health check
def health(request):
    """Root level health check - alias for health_check"""
    return health_check(request)