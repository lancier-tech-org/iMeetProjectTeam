import os
import uuid
import shutil
import subprocess
import json
import re
import asyncio
from urllib.parse import quote_plus
import time
from datetime import datetime, timedelta
from tempfile import TemporaryDirectory
from typing import List
from django.db import connection
from pymongo import MongoClient
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.utils.decorators import method_decorator
from django.views import View

# Graphviz - wrap in try/except
try:
    from graphviz import Source
except ImportError:
    Source = None
    print("Warning: graphviz not available")

# PyTorch - wrap in try/except (GPU features optional)
try:
    import torch
except ImportError:
    torch = None
    print("Warning: torch not available - GPU features disabled")

# Transformers - wrap in try/except (GPU features optional)
try:
    from transformers import MarianMTModel, MarianTokenizer
except ImportError:
    MarianMTModel = None
    MarianTokenizer = None
    print("Warning: transformers not available")

import logging
from urllib.parse import quote_plus
from django.http import StreamingHttpResponse
import openai
from django.urls import path
from asgiref.sync import sync_to_async
from django.http import StreamingHttpResponse, FileResponse
import requests
from io import BytesIO
from wsgiref.util import FileWrapper
import mimetypes
from bson import ObjectId
import boto3
from botocore.exceptions import NoCredentialsError
from pydub import AudioSegment
from deep_translator import GoogleTranslator
from django.utils import timezone
from core.WebSocketConnection.meetings import BAD_REQUEST_STATUS, NOT_FOUND_STATUS, SERVER_ERROR_STATUS, SUCCESS_STATUS, TBL_MEETINGS, create_meetings_table
from openai import OpenAI
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.WebSocketConnection.notifications import (
    ensure_notification_tables,
    _get_recording_meeting_info,
    _get_recording_participants,
    short_id
)
import pytz  # For timezone handling in notifications

# Check GPU availability safely
if torch is not None:
    print("Using GPU:", torch.cuda.is_available())
else:
    print("PyTorch not available - GPU features disabled")

# === CONFIGURATION ===
# AWS Configuration
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
AWS_REGION = os.getenv("AWS_REGION", "ap-south-1")
AWS_S3_BUCKET = os.getenv("AWS_S3_BUCKET", "imeetpro-prod-recordings")

# S3 Client
s3_client = boto3.client(
    "s3",
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
    region_name=AWS_REGION
)

S3_FOLDERS = {
    "videos": os.getenv("S3_FOLDER_VIDEOS", "videos"),
    "transcripts": os.getenv("S3_FOLDER_TRANSCRIPTS", "transcripts"),
    "summary": os.getenv("S3_FOLDER_SUMMARY", "summary"),
    "images": os.getenv("S3_FOLDER_IMAGES", "summary_image"),
    "subtitles": os.getenv("S3_FOLDER_SUBTITLES", "subtitles")
}

try:
    client = OpenAI()
except Exception as e:
    client = None
    print(f"Warning: OpenAI client not initialized: {e}")
    
try:
    groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
except Exception as e:
    groq_client = None
    print(f"Warning: Groq client not initialized: {e}")
    
openai.api_key = os.getenv("OPENAI_API_KEY")

# === FIXED FUNCTION 1: Enhanced S3 Key Extraction ===
def extract_s3_key_from_url(s3_url: str, s3_bucket: str) -> str:
    """
    Extract S3 key from URL with better handling for nested paths and special characters.
    Fixes issue with folder names like: 4b77cc81-9064-434c-8e79-0f9fab31cbdf_Updates
    """
    try:
        if not s3_url:
            return None
        
        logger.debug(f"Extracting S3 key from URL: {s3_url}")
        
        # Remove protocol if present
        if s3_url.startswith('s3://'):
            key = s3_url.replace(f's3://{s3_bucket}/', '')
            logger.debug(f"S3 protocol: Extracted key = {key}")
            return key
        
        # HTTP(S) URL handling - MAIN STRATEGY
        if 'http' in s3_url and '.amazonaws.com/' in s3_url:
            # Split on .amazonaws.com/ - everything after is the S3 key
            # This cleanly separates domain from path regardless of bucket name or special chars
            key = s3_url.split('.amazonaws.com/')[-1]
            # Remove query parameters if any
            key = key.split('?')[0]
            if key:
                logger.debug(f"✅ Extracted S3 key: {key}")
                return key
        
        # Fallback for non-standard URLs
        if 'http' in s3_url and s3_bucket in s3_url:
            # Find first occurrence of bucket
            idx = s3_url.find(s3_bucket)
            after_bucket = s3_url[idx + len(s3_bucket):]
            
            # Look for the path part (after domain suffixes)
            if '/' in after_bucket:
                # Take everything after first slash
                key = after_bucket.split('/', 1)[1]
                key = key.split('?')[0]
                if key:
                    logger.debug(f"✅ Extracted S3 key (fallback): {key}")
                    return key
        
        # If it looks like already a key (no URL patterns)
        if not s3_url.startswith('http') and '://' not in s3_url:
            logger.debug(f"Direct key format: {s3_url}")
            return s3_url
        
        logger.warning(f"⚠️ Could not extract S3 key from URL: {s3_url}")
        return None
        
    except Exception as e:
        logger.error(f"❌ Exception in extract_s3_key_from_url: {e}")
        logger.error(f"URL was: {s3_url}")
        return None

# === FIXED FUNCTION 2: Robust S3 Key Builder (DOCX) ===
def build_s3_key_from_parts(
    base_folder: str,
    meeting_id: str,
    user_id: str,
    meeting_type: str = None,
    doc_type: str = None,
    lang: str = None
) -> str:
    """
    Build complete S3 key path with proper validation (DOCX for documents).
    """
    try:
        # Sanitize inputs
        meeting_id = str(meeting_id).strip() if meeting_id else ""
        user_id = str(user_id).strip() if user_id else ""
        base_folder = str(base_folder).strip() if base_folder else "videos"

        if not meeting_id or not user_id:
            logger.error(f"Invalid input: meeting_id={meeting_id}, user_id={user_id}")
            return None

        # ================= Schedule Meeting =================
        if meeting_type and meeting_type.lower() == "schedulemeeting":
            try:
                schedule_meta = get_schedule_meeting_metadata(meeting_id)

                if schedule_meta.get("is_scheduled"):
                    folder_path = schedule_meta.get("folder_path", "")

                    if doc_type == "transcript":
                        return f"transcripts/{folder_path}/{meeting_id}_{user_id}_transcript.docx"

                    elif doc_type == "summary":
                        return f"summary/{folder_path}/{meeting_id}_{user_id}_summary.docx"

                    elif doc_type == "subtitles":
                        if lang:
                            return f"subtitles/{folder_path}/{meeting_id}_{user_id}_{lang}.srt"
                        return f"subtitles/{folder_path}"

                    # Video
                    return f"{base_folder}/{folder_path}/{meeting_id}_{user_id}_recording.mp4"

            except Exception as schedule_error:
                logger.warning(
                    f"Failed to get schedule metadata: {schedule_error}, falling back to default"
                )

        # ================= Default / Fallback =================
        if doc_type == "transcript":
            return f"transcripts/{meeting_id}_{user_id}_transcript.docx"

        elif doc_type == "summary":
            return f"summary/{meeting_id}_{user_id}_summary.docx"

        elif doc_type == "subtitles":
            if lang:
                return f"subtitles/{meeting_id}_{user_id}_{lang}.srt"
            return f"subtitles/{meeting_id}_{user_id}"

        # Video
        return f"{base_folder}/{meeting_id}_{user_id}_recording.mp4"

    except Exception as e:
        logger.error(f"Error building S3 key: {e}")
        return None


# === FIXED FUNCTION 3: Verify and Repair Video URL ===
def verify_and_repair_video_url(video_doc: dict) -> dict:
    """
    Verify video URL matches actual S3 content.
    If mismatch found, attempt to rebuild the correct path.
    
    Returns updated video document with corrected URL if needed.
    """
    try:
        meeting_id = video_doc.get("meeting_id")
        user_id = video_doc.get("user_id")
        stored_url = video_doc.get("video_url")
        meeting_type = video_doc.get("meeting_type", "InstantMeeting")
        
        if not all([meeting_id, user_id, stored_url]):
            logger.warning(f"Incomplete video metadata: meeting_id={meeting_id}, user_id={user_id}")
            return video_doc
        
        # Check if stored URL exists in S3
        try:
            s3_key = extract_s3_key_from_url(stored_url, AWS_S3_BUCKET)
            if s3_key:
                size = get_s3_object_size(s3_key)
                if size > 0:
                    logger.info(f"✅ Video URL verified: {s3_key} ({size} bytes)")
                    return video_doc  # URL is valid
                else:
                    logger.warning(f"⚠️ Video file not found at S3 key: {s3_key}")
        except Exception as check_error:
            logger.warning(f"⚠️ Failed to verify URL: {check_error}")
        
        # URL is invalid, attempt to rebuild correct path
        logger.info(f"🔧 Attempting to rebuild S3 path for meeting_id={meeting_id}, user_id={user_id}, type={meeting_type}")
        
        rebuilt_key = build_s3_key_from_parts(
            base_folder="videos",
            meeting_id=meeting_id,
            user_id=user_id,
            meeting_type=meeting_type
        )
        
        if not rebuilt_key:
            logger.error(f"Failed to rebuild S3 key")
            return video_doc
        
        # Check if rebuilt key exists
        try:
            size = get_s3_object_size(rebuilt_key)
            if size > 0:
                # Rebuilt path is valid! Update the document
                rebuilt_url = f"https://{AWS_S3_BUCKET}.s3.{AWS_REGION}.amazonaws.com/{rebuilt_key}"
                logger.info(f"✅ Found video at rebuilt path: {rebuilt_key}")
                
                # Update MongoDB
                try:
                    collection.update_one(
                        {"_id": video_doc.get("_id")},
                        {"$set": {"video_url": rebuilt_url}}
                    )
                    logger.info(f"✅ Updated video URL in MongoDB")
                except Exception as update_error:
                    logger.warning(f"⚠️ Failed to update MongoDB: {update_error}")
                
                video_doc["video_url"] = rebuilt_url
                video_doc["url_repaired"] = True
                return video_doc
            else:
                logger.warning(f"⚠️ Rebuilt path also not found: {rebuilt_key}")
                
        except Exception as rebuild_check_error:
            logger.warning(f"⚠️ Failed to check rebuilt path: {rebuild_check_error}")
        
        # Could not find video at any location
        logger.error(f"❌ Cannot locate video file for meeting_id={meeting_id}, user_id={user_id}")
        video_doc["url_error"] = "Video file not found in S3"
        return video_doc
        
    except Exception as e:
        logger.error(f"Error in verify_and_repair_video_url: {e}")
        return video_doc

# MongoDB Configuration - Updated to match FastAPI
mongo_user = os.getenv("MONGO_USER", "connectly")
mongo_password = os.getenv("MONGO_PASSWORD", "password")
mongo_host = os.getenv("MONGO_HOST", "192.168.48.201")
mongo_port = os.getenv("MONGO_PORT", "27017")
mongo_db = os.getenv("MONGO_DB", "imeetpro")

MONGO_URI = os.getenv("MONGO_URI")
mongo_client = MongoClient(MONGO_URI)
db = mongo_client[mongo_db]
collection = db["test"]
TRASH_RETENTION_DAYS = 15
# === LOGGING SETUP ===
logger = logging.getLogger("video_processor")
logging.basicConfig(level=logging.INFO)

# === UTILITY FUNCTIONS ===
def upload_to_aws_s3(local_file_path: str, s3_key: str) -> str:
    """Upload file to AWS S3 and return the URL."""
    try:
        s3_client.upload_file(local_file_path, AWS_S3_BUCKET, s3_key)
        logger.info(f"File uploaded to s3://{AWS_S3_BUCKET}/{s3_key}")
        return f"https://{AWS_S3_BUCKET}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"
    except NoCredentialsError:
        logger.error("AWS credentials not available.")
        return None
    except Exception as e:
        logger.error(f"S3 upload failed: {e}")
        return None

def download_from_s3(s3_key: str, local_file_path: str) -> bool:
    """Download file from S3."""
    try:
        s3_client.download_file(AWS_S3_BUCKET, s3_key, local_file_path)
        return True
    except Exception as e:
        logger.error(f"S3 download failed: {e}")
        return False

def delete_from_s3(s3_key: str) -> bool:
    """Delete file from S3."""
    try:
        s3_client.delete_object(Bucket=AWS_S3_BUCKET, Key=s3_key)
        logger.info(f"Deleted from S3: {s3_key}")
        return True
    except Exception as e:
        logger.error(f"S3 delete failed: {e}")
        return False

def get_s3_object_size(s3_key: str) -> int:
    """Get object size from S3 - optimized version."""
    try:
        response = s3_client.head_object(Bucket=AWS_S3_BUCKET, Key=s3_key)
        return response['ContentLength']
    except Exception as e:
        logger.error(f"Failed to get S3 object size for {s3_key}: {e}")
        return 0

def stream_from_s3(s3_key: str, start: int = None, end: int = None) -> bytes:
    """Stream content from S3 - optimized version."""
    try:
        if start is not None and end is not None:
            range_header = f'bytes={start}-{end}'
            response = s3_client.get_object(Bucket=AWS_S3_BUCKET, Key=s3_key, Range=range_header)
        else:
            response = s3_client.get_object(Bucket=AWS_S3_BUCKET, Key=s3_key)
        
        return response['Body'].read()
    except Exception as e:
        logger.error(f"S3 streaming failed for {s3_key}: {e}")
        return None

def get_meeting_type(meeting_id: str) -> str:
    """Get meeting type from database - checks all meeting tables"""
    try:
        with connection.cursor() as cursor:
            # First check tbl_Meetings for Meeting_Type
            cursor.execute("SELECT Meeting_Type FROM tbl_Meetings WHERE ID = %s", [meeting_id])
            row = cursor.fetchone()
            
            if row and row[0]:
                meeting_type = row[0]
                logger.info(f"Meeting {meeting_id} type: {meeting_type}")
                return meeting_type
            
            # Fallback: Check specific tables if Meeting_Type is NULL
            # Check CalendarMeetings
            cursor.execute("SELECT id FROM tbl_CalendarMeetings WHERE id = %s", [meeting_id])
            if cursor.fetchone():
                logger.info(f"Meeting {meeting_id} found in CalendarMeetings")
                return "CalendarMeeting"
            
            # Check ScheduledMeetings
            cursor.execute("SELECT id FROM tbl_ScheduledMeetings WHERE id = %s", [meeting_id])
            if cursor.fetchone():
                logger.info(f"Meeting {meeting_id} found in ScheduledMeetings")
                return "ScheduleMeeting"
            
            # Default to InstantMeeting if found in tbl_Meetings but not in specific tables
            cursor.execute("SELECT ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
            if cursor.fetchone():
                logger.info(f"Meeting {meeting_id} defaulting to InstantMeeting")
                return "InstantMeeting"
            
            logger.warning(f"Meeting {meeting_id} not found in any table")
            return "InstantMeeting"  # Default fallback
            
    except Exception as e:
        logger.error(f"Error getting meeting type for {meeting_id}: {e}")
        return "InstantMeeting"  # Safe default

def get_schedule_meeting_metadata(meeting_id: str) -> dict:
    """
    Get schedule meeting metadata for organizing recordings into folders.
    Returns schedule_id, schedule_title, and folder path.
    
    This function retrieves metadata about a scheduled meeting so recordings
    can be organized in S3 folders by schedule.
    """
    try:
        with connection.cursor() as cursor:
            # Get the schedule info from tbl_ScheduledMeetings
            cursor.execute("""
                SELECT id, title 
                FROM tbl_ScheduledMeetings 
                WHERE id = %s
            """, [meeting_id])
            
            row = cursor.fetchone()
            if row:
                schedule_id, title = row[0], row[1]
                logger.info(f"✅ Found schedule meeting metadata: schedule_id={schedule_id}, title={title}")
                
                # Sanitize title for folder name (remove special chars, replace spaces with underscores)
                sanitized_title = re.sub(r'[^a-zA-Z0-9\s_-]', '', title)
                sanitized_title = re.sub(r'\s+', '_', sanitized_title.strip())
                
                return {
                    "is_scheduled": True,
                    "schedule_id": schedule_id,
                    "schedule_title": title,
                    "sanitized_folder_name": f"{schedule_id}_{sanitized_title}",
                    "folder_path": f"schedule_meetings/{schedule_id}_{sanitized_title}"
                }
            else:
                logger.warning(f"Meeting {meeting_id} not found in tbl_ScheduledMeetings")
                return {
                    "is_scheduled": False,
                    "schedule_id": None,
                    "schedule_title": None,
                    "sanitized_folder_name": None,
                    "folder_path": None
                }
                
    except Exception as e:
        logger.error(f"Error getting schedule meeting metadata for {meeting_id}: {e}")
        return {
            "is_scheduled": False,
            "schedule_id": None,
            "schedule_title": None,
            "sanitized_folder_name": None,
            "folder_path": None
        }

def build_s3_video_path(meeting_id: str, user_id: str, meeting_type: str) -> str:
    """
    Build S3 path for video based on meeting type.
    
    For ScheduleMeetings: 
        videos/schedule_meetings/{schedule_id}_{sanitized_title}/{meeting_id}_{user_id}_recording.mp4
    
    For InstantMeetings/CalendarMeetings: 
        videos/{meeting_id}_{user_id}_recording.mp4
    
    This organizes scheduled meeting recordings into folders automatically.
    """
    try:
        base_folder = S3_FOLDERS['videos']
        
        if meeting_type == "ScheduleMeeting":
            # Get schedule metadata and create folder structure
            schedule_meta = get_schedule_meeting_metadata(meeting_id)
            
            if schedule_meta.get("is_scheduled"):
                # Path: videos/schedule_meetings/schedule_id_title/meeting_id_user_id_recording.mp4
                s3_key = f"{base_folder}/{schedule_meta['folder_path']}/{meeting_id}_{user_id}_recording.mp4"
                logger.info(f"✅ ScheduleMeeting S3 path: {s3_key}")
                return s3_key
        
        # Default path for InstantMeeting, CalendarMeeting, or if schedule fetch fails
        s3_key = f"{base_folder}/{meeting_id}_{user_id}_recording.mp4"
        logger.info(f"Default S3 path: {s3_key}")
        return s3_key
        
    except Exception as e:
        logger.error(f"Error building S3 path: {e}")
        return f"{S3_FOLDERS['videos']}/{meeting_id}_{user_id}_recording.mp4"
    
def build_s3_document_path(
    meeting_id: str,
    user_id: str,
    meeting_type: str,
    doc_type: str
) -> str:
    """
    Build S3 path for transcript/summary based on meeting type.

    For ScheduleMeetings: organizes documents in schedule folder
    For others: flat structure

    doc_type can be: 'transcript', 'summary', or 'subtitles'
    """
    try:
        # Default meeting type
        if not meeting_type:
            meeting_type = "InstantMeeting"

        logger.debug(
            f"Building S3 path: type={doc_type}, meeting_type={meeting_type}"
        )

        # ================= Schedule Meeting =================
        if meeting_type == "ScheduleMeeting":
            try:
                schedule_meta = get_schedule_meeting_metadata(meeting_id)

                if schedule_meta and schedule_meta.get("is_scheduled"):
                    folder_prefix = schedule_meta.get("folder_path", "")

                    if doc_type == "transcript":
                        base_folder = S3_FOLDERS.get("transcripts", "transcripts")
                        return (
                            f"{base_folder}/{folder_prefix}/"
                            f"{meeting_id}_{user_id}_transcript.docx"
                        )

                    elif doc_type == "summary":
                        base_folder = S3_FOLDERS.get("summary", "summary")
                        return (
                            f"{base_folder}/{folder_prefix}/"
                            f"{meeting_id}_{user_id}_summary.docx"
                        )

                    elif doc_type == "subtitles":
                        base_folder = S3_FOLDERS.get("subtitles", "subtitles")
                        return f"{base_folder}/{folder_prefix}"  # lang appended later

                else:
                    logger.warning(
                        f"Schedule metadata not found or not scheduled for {meeting_id}"
                    )

            except Exception as schedule_error:
                logger.warning(
                    f"Error fetching schedule metadata: {schedule_error}, "
                    f"using default path"
                )

        # ================= Default / Fallback =================
        if doc_type == "transcript":
            return (
                f"{S3_FOLDERS.get('transcripts', 'transcripts')}/"
                f"{meeting_id}_{user_id}_transcript.docx"
            )

        elif doc_type == "summary":
            return (
                f"{S3_FOLDERS.get('summary', 'summary')}/"
                f"{meeting_id}_{user_id}_summary.docx"
            )

        elif doc_type == "subtitles":
            return f"{S3_FOLDERS.get('subtitles', 'subtitles')}"

        logger.warning(f"Unknown doc_type: {doc_type}")
        return ""

    except Exception as e:
        logger.error(f"Error building document S3 path: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return ""

### ✅ FIXED `is_user_allowed` FUNCTION
def is_user_allowed(meeting_id: str, email: str = "", user_id: str = "") -> bool:
    """Check if user is allowed to access meeting recording - FIXED OR LOGIC"""
    try:
        if not email and not user_id:
            return False
            
        with connection.cursor() as cursor:
            # Check 1: ScheduledMeeting email
            if email:
                cursor.execute("SELECT email FROM tbl_ScheduledMeetings WHERE id = %s", [meeting_id])
                row = cursor.fetchone()
                if row and row[0]:
                    scheduled_emails = [e.strip().lower() for e in row[0].split(',') if e.strip()]
                    if email.strip().lower() in scheduled_emails:
                        logger.info(f"✅ Access granted via ScheduledMeeting: {email}")
                        return True  # ✅ RETURN IMMEDIATELY

            # Check 2: CalendarMeeting email
            if email:
                cursor.execute("SELECT email, guestEmails, attendees FROM tbl_CalendarMeetings WHERE id = %s", [meeting_id])
                row = cursor.fetchone()
                if row:
                    for field in row:
                        if field:
                            emails = [e.strip().lower() for e in re.split(r'[;,]', field) if e.strip()]
                            if email.strip().lower() in emails:
                                logger.info(f"✅ Access granted via CalendarMeeting: {email}")
                                return True  # ✅ RETURN IMMEDIATELY

            # Check 3: InstantMeeting participant (only for instant meetings)
            if user_id:
                cursor.execute("""
                    SELECT COUNT(*) FROM tbl_Participants 
                    WHERE Meeting_ID = %s AND User_ID = %s
                """, [meeting_id, user_id])
                count = cursor.fetchone()[0]
                if count > 0:
                    logger.info(f"✅ Access granted: User {user_id} joined meeting")
                    return True  # ✅ RETURN IMMEDIATELY

            # Check 4: Meeting host
            if user_id:
                cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                row = cursor.fetchone()
                if row and str(row[0]) == str(user_id):
                    logger.info(f"✅ Access granted: User {user_id} is host")
                    return True  # ✅ RETURN IMMEDIATELY

        # ALL checks failed
        logger.debug(f"❌ Access denied: No authorization found for {user_id}/{email}")
        return False
        
    except Exception as e:
        logger.error(f"Error in is_user_allowed: {e}")
        return False
            
def is_user_allowed_debug(meeting_id: str, email: str = "", user_id: str = "") -> bool:
    """Debug version - more permissive for testing."""
    try:
        if email or user_id:
            return True
        return False
    except Exception:
        return False
        
def get_meeting_participants_emails(meeting_id: str) -> list:
    """Get all participant emails for a meeting from different meeting types."""
    visible_to_emails = []
    
    try:
        with connection.cursor() as cursor:
            # Check Scheduled Meetings
            cursor.execute("SELECT email FROM tbl_ScheduledMeetings WHERE id = %s", [meeting_id])
            row = cursor.fetchone()
            if row and row[0]:
                visible_to_emails += [e.strip() for e in row[0].split(',') if e.strip()]

            # Check Calendar Meetings
            cursor.execute("SELECT email, guestEmails, attendees FROM tbl_CalendarMeetings WHERE id = %s", [meeting_id])
            row = cursor.fetchone()
            if row:
                for field in row:
                    if field:
                        emails = [e.strip() for e in re.split(r'[;,]', field) if e.strip()]
                        visible_to_emails += emails

            # Check Instant Meeting Participants
            cursor.execute("SELECT User_ID FROM tbl_Participants WHERE Meeting_ID = %s", [meeting_id])
            user_ids = [r[0] for r in cursor.fetchall() if r[0]]

            if user_ids:
                format_strings = ','.join(['%s'] * len(user_ids))
                cursor.execute(f"SELECT Email FROM tbl_Users WHERE ID IN ({format_strings})", user_ids)
                visible_to_emails += [r[0].strip() for r in cursor.fetchall() if r[0]]

        # Remove duplicates and empty strings
        visible_to_emails = list(set([email for email in visible_to_emails if email]))
        
        logger.info(f"Found {len(visible_to_emails)} participant emails for meeting {meeting_id}")
        return visible_to_emails
        
    except Exception as e:
        logger.error(f"Failed to get meeting participants emails for meeting {meeting_id}: {e}")
        return []

def get_user_details(user_id: str) -> dict:
    """Get user details (name, email) from database."""
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT ID, full_name, email 
                FROM tbl_Users 
                WHERE ID = %s
            """, [user_id])
            row = cursor.fetchone()
            
            if row:
                user_id_db, full_name, email = row
                
                # Priority: full_name > email prefix > User ID
                if full_name and full_name.strip():
                    display_name = full_name.strip()
                elif email and email.strip():
                    display_name = email.split('@')[0]
                else:
                    display_name = f"User {user_id}"
                
                logger.info(f"✅ Found user {user_id}: full_name='{full_name}', email='{email}', display='{display_name}'")
                
                return {
                    "user_id": str(user_id_db),
                    "user_name": display_name,
                    "user_email": email or ""
                }
            
            logger.warning(f"❌ User {user_id} not found in tbl_Users")
            return {
                "user_id": str(user_id),
                "user_name": f"User {user_id}",
                "user_email": ""
            }
            
    except Exception as e:
        logger.error(f"❌ Error getting user details for {user_id}: {e}")
        return {
            "user_id": str(user_id),
            "user_name": f"User {user_id}",
            "user_email": ""
        }

def get_user_details(user_id: str) -> dict:
    """Get user details (name, email) from database."""
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT ID, full_name, email 
                FROM tbl_Users 
                WHERE ID = %s
            """, [user_id])
            row = cursor.fetchone()
            
            if row:
                user_id_db, full_name, email = row
                
                # Priority: full_name > email prefix > User ID
                if full_name and full_name.strip():
                    display_name = full_name.strip()
                elif email and email.strip():
                    display_name = email.split('@')[0]
                else:
                    display_name = f"User {user_id}"
                
                logger.info(f"✅ Found user {user_id}: full_name='{full_name}', email='{email}', display='{display_name}'")
                
                return {
                    "user_id": str(user_id_db),
                    "user_name": display_name,
                    "user_email": email or ""
                }
            
            logger.warning(f"❌ User {user_id} not found in tbl_Users")
            return {
                "user_id": str(user_id),
                "user_name": f"User {user_id}",
                "user_email": ""
            }
            
    except Exception as e:
        logger.error(f"❌ Error getting user details for {user_id}: {e}")
        return {
            "user_id": str(user_id),
            "user_name": f"User {user_id}",
            "user_email": ""
        }

def clean_markdown(text: str) -> str:
    """
    Cleans markdown syntax from LLM summary while preserving hierarchy and readability.
    """
    import re

    # Allow code blocks. Only strip markdown syntax, not code.
    text = re.sub(r"```(\w+)?", "", text)
    text = text.replace("```", "")

    # remove bold
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    # inline code cleanup
    text = re.sub(r"`([^`]*)`", r"\1", text)

    # convert headings (#, ##, ###)
    text = re.sub(r"^### (.*)", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"^## (.*)", r"\1", text, flags=re.MULTILINE)
    text = re.sub(r"^# (.*)", r"\1", text, flags=re.MULTILINE)

    # normalize spacing
    text = re.sub(r"\n{2,}", "\n\n", text)

    return text.strip()

def format_srt_time(seconds: float) -> str:
    from datetime import timedelta

    td = timedelta(seconds=seconds)
    total_seconds = int(td.total_seconds())
    millis = int((td.total_seconds() - total_seconds) * 1000)

    return f"{str(timedelta(seconds=total_seconds)).zfill(8)},{millis:03}"

def create_srt_from_segments(segments: list, output_path: str):
    import logging
    logger = logging.getLogger("video_processor")

    if not segments:
        logger.warning(f"[WARN] No segments found. Writing fallback subtitle: {output_path}")
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(
                "1\n"
                "00:00:00,000 --> 00:00:05,000\n"
                "[No speech detected]\n\n"
            )
        return

    logger.info(f"[INFO] Writing {len(segments)} subtitle segments → {output_path}")

    with open(output_path, "w", encoding="utf-8") as f:
        for idx, seg in enumerate(segments, start=1):
            start = format_srt_time(float(seg.get("start", 0)))
            end = format_srt_time(float(seg.get("end", 0)))
            text = seg.get("text", "").strip()

            if not text:
                text = "[Silence]"

            f.write(f"{idx}\n")
            f.write(f"{start} --> {end}\n")
            f.write(f"{text}\n\n")

def generate_graph(dot_code: str, output_path: str):
    from graphviz import Source
    s = Source(dot_code)
    return s.render(filename=output_path, format="png", cleanup=True)

def save_docx(content: str, path: str, image_path: str = None, title: str = ""):
    from docx.shared import RGBColor

    doc = Document()

    # === Title ===
    if title:
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === Main Content ===
    for line in content.splitlines():
        line = line.strip()
        if not line:
            continue

        # Headings
        if line.startswith("### "):
            run = doc.add_heading(line[4:].strip(), level=3).runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith("## "):
            run = doc.add_heading(line[3:].strip(), level=2).runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif line.startswith("# "):
            run = doc.add_heading(line[2:].strip(), level=1).runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Bold Example headings
        elif line.lower().startswith("example"):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        # Normal text
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(12)

    # === Add Mind Map (ONE TIME ONLY) ===
    if image_path and os.path.exists(image_path):
        doc.add_page_break()
        doc.add_heading("Mind Map", level=2)
        doc.add_picture(image_path, width=Inches(6))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # === Save once ===
    doc.save(path)

def send_recording_completion_notifications(meeting_id, video_url, transcript_url=None, summary_url=None):
    """
    Send professional notifications to all meeting participants when recording is ready.
    Different messages for host vs participants.
    
    INCLUDES:
    - Duplicate prevention at function level (prevents double-calling)
    - Duplicate prevention at participant level (prevents same email twice)
    - Case-insensitive email handling
    - Professional HTML formatting (NO EMOJIS - MySQL utf8 compatible)
    """
    if not meeting_id or not video_url:
        logging.warning("Missing meeting_id or video_url for recording notifications")
        return 0

    # ✅ CRITICAL FIX 1: Check if notifications already sent for this meeting
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT COUNT(*) FROM tbl_Notifications 
                WHERE meeting_id = %s 
                  AND notification_type IN ('recording_completed', 'recording_completed_host')
            """, [str(meeting_id)])
            
            existing_count = cursor.fetchone()[0]
            
            if existing_count > 0:
                logging.warning(f"Recording notifications already sent for meeting {meeting_id} ({existing_count} exist). Skipping to prevent duplicates.")
                return 0
    except Exception as e:
        logging.warning(f"Could not check for existing notifications: {e}")
        # Continue anyway - better to risk duplicate than miss notification

    try:
        ensure_notification_tables()
    except Exception as e:
        logging.error(f"Failed to ensure notification tables: {e}")
        return 0

    ist = pytz.timezone("Asia/Kolkata")
    now = datetime.now(ist)
    sent_count = 0
    skipped_count = 0

    # ========== FETCH MEETING DETAILS ==========
    meeting_info = _get_recording_meeting_info(meeting_id)
    
    meeting_title = meeting_info.get('title') or f"Meeting {meeting_id[:8]}"
    meeting_type = meeting_info.get('meeting_type') or 'Meeting'
    host_name = meeting_info.get('host_name') or 'Host'
    host_email = meeting_info.get('host_email')
    video_name = meeting_info.get('custom_recording_name') or f"Recording - {meeting_title}"
    meeting_date = meeting_info.get('start_time')
    
    # Format date nicely
    if meeting_date:
        if isinstance(meeting_date, str):
            try:
                meeting_date = datetime.strptime(meeting_date, '%Y-%m-%d %H:%M:%S')
            except:
                meeting_date = now
        formatted_date = meeting_date.strftime('%B %d, %Y at %I:%M %p')
    else:
        formatted_date = now.strftime('%B %d, %Y')

    # ========== GET ALL PARTICIPANTS ==========
    participants = _get_recording_participants(meeting_id)
    
    if not participants:
        logging.warning(f"No participants found for meeting {meeting_id}")
        return 0

    # ✅ CRITICAL FIX 2: Deduplicate participants list (case-insensitive)
    unique_participants = []
    seen_emails = set()
    for email in participants:
        if email and '@' in email:
            normalized_email = email.strip().lower()
            if normalized_email not in seen_emails:
                seen_emails.add(normalized_email)
                unique_participants.append(normalized_email)
    
    if len(unique_participants) != len(participants):
        logging.info(f"Deduplicated participants: {len(participants)} -> {len(unique_participants)}")

    logging.info(f"Sending recording notifications to {len(unique_participants)} unique participants")

    # ========== BUILD RESOURCE LIST (Professional - NO EMOJIS) ==========
    resources = ["Video Recording"]
    if transcript_url:
        resources.append("Full Transcript")
    if summary_url:
        resources.append("AI-Generated Summary & Mind Map")

    resources_text = "\n".join([f"  * {r}" for r in resources])

    # ========== SEND NOTIFICATIONS ==========
    for participant_email in unique_participants:
        if not participant_email or '@' not in participant_email:
            continue

        try:
            participant_email = participant_email.strip().lower()
            
            # ✅ CRITICAL FIX 3: Check if notification already exists for this specific user
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT COUNT(*) FROM tbl_Notifications 
                    WHERE meeting_id = %s 
                      AND LOWER(recipient_email) = %s
                      AND notification_type IN ('recording_completed', 'recording_completed_host')
                """, [str(meeting_id), participant_email])
                
                if cursor.fetchone()[0] > 0:
                    logging.info(f"Notification already exists for {participant_email}, skipping")
                    skipped_count += 1
                    continue

            is_host = (
                host_email and 
                participant_email == host_email.strip().lower()
            )

            # Calculate other participants count for host message
            other_participants_count = len(unique_participants) - 1

            if is_host:
                # ========== HOST NOTIFICATION (Professional Format - NO EMOJIS) ==========
                notification_type = "recording_completed_host"
                title = "Your Recording is Ready"
                message = (
                    f"<b>Great news!</b> Your meeting recording has been successfully processed and is now available.\n\n"
                    f"<b>Meeting Details</b>\n"
                    f"  * <b>Title:</b> {meeting_title}\n"
                    f"  * <b>Date:</b> {formatted_date}\n"
                    f"  * <b>Recording:</b> {video_name}\n\n"
                    f"<b>Available Resources</b>\n"
                    f"{resources_text}\n\n"
                    f"<b>Notification Status</b>\n"
                    f"  * {other_participants_count} participant(s) have been notified.\n\n"
                    f"<i>Click here to access your recording dashboard.</i>"
                )
                priority = "normal"
            else:
                # ========== PARTICIPANT NOTIFICATION (Professional Format - NO EMOJIS) ==========
                notification_type = "recording_completed"
                title = "Meeting Recording Available"
                message = (
                    f"<b>Hello!</b> A recording from your recent meeting is now available for you to view.\n\n"
                    f"<b>Meeting Details</b>\n"
                    f"  * <b>Title:</b> {meeting_title}\n"
                    f"  * <b>Host:</b> {host_name}\n"
                    f"  * <b>Date:</b> {formatted_date}\n"
                    f"  * <b>Recording:</b> {video_name}\n\n"
                    f"<b>You Have Access To</b>\n"
                    f"{resources_text}\n\n"
                    f"<i>Click here to view the recording in your dashboard.</i>"
                )
                priority = "normal"

            # ========== INSERT NOTIFICATION ==========
            notification_id = short_id()
            
            with connection.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO tbl_Notifications (
                        id, recipient_email, meeting_id, notification_type,
                        title, message, meeting_title, meeting_url,
                        is_read, priority, created_at
                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, [
                    notification_id,
                    participant_email,
                    str(meeting_id),
                    notification_type,
                    title,
                    message,
                    meeting_title,
                    video_url,
                    False,
                    priority,
                    now
                ])
                
                # ✅ Explicit commit
                connection.commit()

                if cursor.rowcount > 0:
                    sent_count += 1
                    logging.info(f"Recording notification sent to {participant_email} ({'host' if is_host else 'participant'})")
                else:
                    logging.error(f"INSERT returned 0 rows for {participant_email}")

        except Exception as e:
            logging.error(f"Failed to send recording notification to {participant_email}: {e}", exc_info=True)

    logging.info(f"Recording notifications completed: {sent_count} sent, {skipped_count} skipped (duplicates)")
    return sent_count

# === NEW: CHUNKED TRANSCRIPTION FROM FASTAPI ===
async def transcribe_chunk(chunk_file: str, offset: float):
    try:
        with open(chunk_file, "rb") as f:
            result = groq_client.audio.transcriptions.create(
                model="whisper-large-v3-turbo",
                file=f,
                response_format="verbose_json"
            )

        segments = []

        # CASE 1 — Whisper returned real timestamped segments
        if hasattr(result, "segments") and result.segments:
            for seg in result.segments:
                segments.append({
                    "start": offset + float(seg["start"]),
                    "end": offset + float(seg["end"]),
                    "text": seg["text"].strip()
                })

        # CASE 2 — No segments, fallback
        else:
            text = (result.text or "").strip()
            if not text:
                text = "[No speech detected]"

            segments.append({
                "start": offset,
                "end": offset + 5,
                "text": text
            })

        return segments

    except Exception as e:
        logger.error(f"[GROQ ERROR] {e}")
        return [{
            "start": offset,
            "end": offset + 5,
            "text": "[Transcription failed]"
        }]

def summarize_segment(transcript: str, context: str = ""):
    prompt = f"""
You are a senior documentation and technical writing expert. Your task is to convert the following raw transcript segment into a comprehensive, highly accurate, and formal implementation or study guide based on the subject matter discussed.

The final output must:

- Be structured and formatted according to professional standards for enterprise-level training, onboarding, line pictures, and technical enablement.
- Include step-by-step procedures, clearly numbered and logically ordered.
- Provide real-world tools, technologies, configurations, commands, and screenshots/images (placeholders if needed) relevant to the topic.
- Embed technical examples, use cases, CLI/GUI instructions, and expected outputs or screenshots where applicable.
- Cover common pitfalls, troubleshooting tips, and best practices to ensure full practical understanding.
- Use terminology and instructional depth suitable for readers to gain 100% conceptual and hands-on knowledge of the subject.
- The final document should resemble internal documentation used at organizations like SAP, Oracle, Java, Selenium, AI/ML, Data Science, AWS, Microsoft, or Google — clear, comprehensive, and instructional in tone.
---
---------------------------------------------------------------------
EXPLANATION STYLE RULE (VERY IMPORTANT – LARGE SCALE):

The output must follow a TEACHING STYLE similar to real instructors on YouTube:
- Step-by-step explanation
- Beginner-friendly structure
- Concepts explained in simple language
- Real examples in EVERY section
- Coding examples ONLY when transcript includes coding
- If NOT coding → produce real-life scenario examples (not code)

The explanation must look like a proper "Getting Started Guide" or "Beginner-Friendly Training Material".

---------------------------------------------------------------------
FINAL FULL CODE OUTPUT RULE (LARGE SCALE – ALWAYS CORRECT):

Before generating the Conclusion and Mind Map, create a new section titled:

"Complete Code Example (Only When Coding Is Detected)"

TRIGGER CONDITIONS (coding = TRUE):
Coding is TRUE if ANY of the following occur:
1. Transcript or explanation includes ANY code snippet.
2. Transcript mentions ANY programming language:
   (Python, Java, C++, JavaScript, SQL, React, Node, HTML, etc.)
3. Transcript mentions ANY programming concepts:
   (function, class, variable, loop, API, script, project, compiler, debugging, IDE)
4. The assistant generated ANY inline code earlier in the explanation.

IF coding = TRUE:
    • ALWAYS generate the "Complete Code Example" section.
    • Provide ONE clean, runnable, complete code example that represents the main topic.
    • Include:
        - Imports (if applicable)
        - Functions or classes
        - Main execution block
        - Comments explaining the logic
        - Proper indentation & syntax
        - Output example (if applicable)

IF coding = FALSE:
    • Do NOT generate this section.

Placement requirements:
This section must appear:
    ✔ RIGHT BEFORE "Conclusion"
    ✔ After all explanation sections
    ✔ Before "Mind Map" and "Diagram Placeholder"
---------------------------------------------------------------------

OBJECTIVE:

Create a detailed, real-world step-by-step implementation or process guide for [INSERT TOPIC/SUBJECT], designed specifically to support the creation of over 100 technical or comprehension questions. The guide must:

- Reflect real-world tools, technologies, workflows, and industry terminology.
- Break down each phase of the implementation or process logically and sequentially.
- Include practical examples, code snippets (if applicable), key decisions, best practices, and commonly used tools at each step.
- Highlight common challenges or misconceptions, and how they’re addressed in real practice.
- Use terminology and structure that would support SMEs or instructional designers in generating high-quality technical questions based on the guide.
- Avoid abstract or overly generic statements — focus on precision, clarity, and applied knowledge.
- If the transcript is fully or mostly coding-related (Python, Java, JavaScript, React, SQL, DevOps, automation, APIs, backend, etc.), you must include complete coding examples with correct indentation, working syntax, and explanatory comments. You must also expand the matter to look like a proper coding implementation guide, including functions, scripts, API samples, folder structures, command-line usage, debugging steps, and real-world coding workflows.

---

DOCUMENT FORMAT & STRUCTURE RULES:

1. STRUCTURE
- Use numbered sections and sub-sections (e.g., 1, 1.1, 1.2.1)
- No markdown, emojis, or decorative formatting
- Use plain, formal, enterprise-grade language

2. EACH SECTION MUST INCLUDE:
- A *clear title* and *brief purpose statement*
- *Step-by-step technical or procedural instructions*, including:
    - All relevant tools, platforms, or interfaces used (if any)
    - Any paths, commands, actions, configurations, or API calls involved
    - All required inputs, values, parameters, or dependencies
    - A logical sequence of operations, clearly numbered or separated by actionable steps
    - Tips, warnings, and Important Notes, or expected outcomes where necessary
- **5-10 sentence description** of each main topic, explaining what the concept is, its use cases, and real-world applications. This should be clear and concise for technical audiences to understand why the topic is essential and how it fits into practical workflows.

3. VALIDATION

- Describe how to confirm success (e.g., Expected Outputs, System or Health Checks, Technical and Functional Verifications, Visual Indicators, Fallback/Error Conditions indicators)

4. TROUBLESHOOTING (if applicable)

- Clearly list frequent or known issues that may arise during or after the procedure
- Describe the conditions or misconfigurations that typically lead to each issue
- Provide step-by-step corrective actions or configuration changes needed to resolve each problem
- Mention specific file paths, log viewer tools, console commands, or dashboard areas where errors and diagnostics can be found
- Include example error codes or system messages that help in identifying the issue

5. BEST PRACTICES

- You are a senior technical writer. Based on the following transcript or topic, create a BEST PRACTICES section suitable for formal technical documentation, onboarding materials, or enterprise IT guides.
- Efficiency improvements (e.g., time-saving configurations, automation tips)
- Security or compliance tips (e.g., encryption, IAM roles, audit logging)
- Standard operating procedures (SOPs) used in enterprise environments
- Avoided pitfalls and why they should be avoided
- Format the content using bullet points or short sections for clarity and actionability.
- Avoid vague, obvious, or overly general suggestions — focus on real-world, practical insights derived from field experience or best-in-class implementation norms.

6. CONCLUSION
- Summarize what was implemented or discussed
- Confirm expected outcomes and readiness indicators

---

IMPORTANT:
If the input contains any values such as usernames, IP addresses, server names, passwords, port numbers, or similar technical identifiers — replace their actual content with generic XML-style tags, while preserving the sentence structure and purpose. For example:

- Replace any specific IP address with: <ip>
- Replace any actual password or secret with: <password>
- Replace any actual hostname with: <hostname>
- Replace any actual port number with: <port>
- Replace any username with: <username>
- Replace any email with: <email>

Do NOT alter the sentence structure, meaning, or flow — keep the language intact while swapping the actual values with tags.
Do not display or retain real values — just show the placeholder tag. Maintain the original meaning and flow of the instructions.
Format the output as clean, professional documentation, suitable for inclusion in implementation guides, SOPs, or training materials.
Highlight any placeholders in a way that makes it easy for the user to identify where to substitute their own values later.

---

Also:
- Cross-check all tools, commands, file paths, service names, APIs, and utilities with reliable, real-world sources (e.g., official vendor documentation, widely accepted best practices).

 1. If something appears ambiguous, incorrect, or outdated, correct it to its current, supported version.
 2. Use only commands, APIs, or tool names that are verifiably valid and relevant to the topic context.
- Consolidate duplicate or f                          ragmented instructions:
 1. If a step or process is repeated across segments, merge them into a single, complete, and accurate version.
 2. Remove redundancy and preserve the most detailed and correct version of each step.
 3. Do NOT include deprecated or unverifiable content:
 4. Exclude outdated commands, legacy references, or tools no longer maintained.
 5. Replace such content with modern equivalents where available.

- Output the final result as a formal technical guide, with:
  1. Clear section headings
  2. Correct and tested commands/scripts
  3. Accurate tool names and workflows
  4. Logical flow suitable for developers, engineers, or IT teams

---

---------------------------------------------------------------------
FINAL FULL CODE OUTPUT RULE (ADD AT END OF SUMMARY)
Before generating the Conclusion and Mind Map, create a new section titled:

"Complete Code Example (Only When Coding Is Detected)"

RULES:
1. This section must appear at the END of the summary, right before Conclusion.
2. Add this section ONLY if the transcript contains ANY coding-related content.
3. Detect coding language automatically (Python, Java, C++, JavaScript, SQL, etc.)
4. Provide ONE clean, runnable, complete code example:
      - Full imports / package statements
      - Functions or classes
      - Main execution block
      - Comments explaining the logic
      - Correct indentation and syntax
5. The code must represent the MAIN concept taught in the transcript.
6. If transcript is NOT coding related → do NOT generate this section.
7. After this code block, continue with:
      • Conclusion
      • Mind map (DOT)
      • Diagram placeholder
---------------------------------------------------------------------


COMBINED INPUT:
\"\"\"{transcript}\n\n{context}\"\"\"

---

FINAL INSTRUCTION:
Return only the fully formatted implementation or process guide that includes below:

- A clear, descriptive title
- A concise purpose statement or overview
- Prerequisites and tools required
- Numbered step-by-step instructions with:
   1. Commands, paths, configuration settings, or code blocks (as needed)
   2. GUI or CLI actions explained clearly
   3. Expected inputs, parameters, or options
   4. Confirmation of success (outputs, logs, tests, or validation steps)
   5. Troubleshooting (common issues, causes, and resolutions — if applicable)
   6. Best Practices (efficiency, reliability, security — if applicable)
   7. **Include a mind map diagram in DOT format enclosed in triple backticks at the end**
   8. **Insert chart/diagram placeholders inline to represent where the visual mind map image should appear**

- Replace any real usernames, IP addresses, passwords, ports, or hostnames with <username>, <ip>, <password>, <port>, or <hostname> where needed.
- Eliminate all redundant or outdated, abused content. Only use valid and current tools and commands.

End Document with Standardized "Suggested Next Steps" Note  
*Suggested next steps: No specific next steps mentioned in this segment.*
"""


    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a technical documentation assistant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=14000 # new param in v1.x
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"[ERROR] Summary generation failed: {e}")
        return "Summary generation failed."
    
def analyze_trainer_performance(transcript: str) -> dict:
    """
    Analyze trainer's technical content, explanation clarity, friendliness, and communication
    based purely on transcript semantics and delivery style markers.
    Uses GPT-4o for text-based behavioral and technical scoring.
    """
    if not transcript.strip():
        return {
            "technical_content": 0,
            "explanation_clarity": 0,
            "friendliness": 0,
            "communication": 0,
            "overall_feedback": "No speech detected for evaluation."
        } 

    prompt = f"""
You are an expert communication and training evaluator.
Evaluate the trainer's communication quality, tone, and content in the transcript below.

TRANSCRIPT:
\"\"\"{transcript}\"\"\"

Evaluate across the following dimensions (each scored 0–100%):
1. Technical Content — accuracy, depth, and domain clarity.
2. Explanation Clarity — how logically and simply ideas are explained.
3. Friendliness — warmth, politeness, and positive tone.
4. Communication – evaluate using Indian English standards:
   - Focus on fluency, confidence, and comfort with Indian accent.
   - Accept light Indianisms such as “basically”, “ok na”, “ya”, etc.
   - Penalize only unclear speech or excessive filler use (“umm”, “like”, “you know”).
   - Do not reduce marks for accent style — evaluate clarity, not foreign pronunciation.
Output a short, factual JSON report with numeric scores and 1–2 lines of feedback.
Format exactly as:

{{
  "technical_content": <number>,
  "explanation_clarity": <number>,
  "friendliness": <number>,
  "communication": <number>,
  "overall_feedback": "<short summary>"
}}
"""

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a behavioral analytics and technical communication evaluator."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=500
        )

        raw = response.choices[0].message.content.strip()
        # Extract JSON safely
        match = re.search(r"\{.*\}", raw, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        else:
            logger.warning("⚠️ Unable to parse trainer evaluation JSON.")
            return {"error": "Failed to parse evaluation output."}

    except Exception as e:
        logger.error(f"[ERROR] Trainer performance evaluation failed: {e}")
        return {"error": str(e)}

# === ENHANCED VIDEO PROCESSING WITH FASTAPI FEATURES ===
# Add these imports at the very top of your file
try:
    import torch
except ImportError:
    torch = None

try:
    from transformers import MarianMTModel, MarianTokenizer
except ImportError:
    MarianMTModel = None
    MarianTokenizer = None
import logging

class LocalIndianLanguageTranslator:
    """Fast local translation for Hindi and Telugu using Helsinki-NLP models"""
    
    def __init__(self):
        self.models = {}
        self.tokenizers = {}
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        # Model names for Indian languages
        self.model_names = {
            "hi": "Helsinki-NLP/opus-mt-en-hi",  # English to Hindi
            "te": "Helsinki-NLP/opus-mt-en-hi",  # Use Hindi model for Telugu (fallback)
        }
        
        logging.info(f"🔧 LocalTranslator initialized on device: {self.device}")
    
    def load_model(self, lang):
        """Lazy load model only when needed"""
        if lang in self.models:
            return  # Already loaded
        
        try:
            model_name = self.model_names.get(lang)
            if not model_name:
                raise Exception(f"No model available for {lang}")
            
            logging.info(f"📥 Loading {lang} translation model: {model_name}")
            
            self.tokenizers[lang] = MarianTokenizer.from_pretrained(model_name)
            self.models[lang] = MarianMTModel.from_pretrained(model_name)
            
            # Move to GPU if available
            if self.device == "cuda":
                self.models[lang] = self.models[lang].cuda()
                logging.info(f"✅ {lang} model loaded on GPU")
            else:
                logging.info(f"✅ {lang} model loaded on CPU")
                
        except Exception as e:
            logging.error(f"❌ Failed to load {lang} model: {e}")
            raise
    
    def translate_batch(self, texts, target_lang):
        """Translate a batch of texts to target language"""
        if target_lang not in self.model_names:
            raise Exception(f"Unsupported language: {target_lang}")
        
        # Load model if not already loaded
        self.load_model(target_lang)
        
        try:
            # Tokenize input texts
            inputs = self.tokenizers[target_lang](
                texts, 
                return_tensors="pt", 
                padding=True, 
                truncation=True, 
                max_length=512
            )
            
            # Move to GPU if available
            if self.device == "cuda":
                inputs = {k: v.cuda() for k, v in inputs.items()}
            
            # Generate translations
            with torch.no_grad():  # Disable gradient calculation for inference
                outputs = self.models[target_lang].generate(
                    **inputs,
                    max_length=512,
                    num_beams=4,  # Better quality
                    early_stopping=True
                )
            
            # Decode translations
            translations = self.tokenizers[target_lang].batch_decode(
                outputs, 
                skip_special_tokens=True
            )
            
            return translations
            
        except Exception as e:
            logging.error(f"Translation error for {target_lang}: {e}")
            # Return fallback translations
            return ["[Translation unavailable]"] * len(texts)
    
    def translate_segments(self, segments, target_lang, batch_size=32):
        """Translate all segments for a language in batches"""
        if target_lang == "en":
            return segments  # No translation needed
        
        logging.info(f"🌐 Translating {len(segments)} segments to {target_lang} using local model...")
        
        translated_segments = []
        total_batches = (len(segments) + batch_size - 1) // batch_size
        
        for batch_idx in range(0, len(segments), batch_size):
            batch_segments = segments[batch_idx:batch_idx + batch_size]
            batch_texts = [seg["text"].strip() for seg in batch_segments]
            
            # Translate batch
            translations = self.translate_batch(batch_texts, target_lang)
            
            # Create translated segments
            for i, seg in enumerate(batch_segments):
                translated_segments.append({
                    "start": seg["start"],
                    "end": seg["end"],
                    "text": translations[i] if i < len(translations) else "[Translation failed]"
                })
            
            # Progress logging
            current_batch = (batch_idx // batch_size) + 1
            if current_batch % 5 == 0 or current_batch == total_batches:
                progress = (len(translated_segments) / len(segments)) * 100
                logging.info(f"   {target_lang}: {progress:.1f}% ({len(translated_segments)}/{len(segments)} segments)")
        
        return translated_segments
    
    def cleanup(self):
        """Free GPU memory"""
        for model in self.models.values():
            del model
        self.models.clear()
        self.tokenizers.clear()
        
        if torch.cuda.is_available():
            torch.cuda.empty_cache()

def process_video_sync(video_path: str, meeting_id: str, user_id: str):
    """Process video with GPU acceleration - UPDATED WITH GROQ + DOCX + TRAINER EVAL"""
    import logging
    import subprocess
    from tempfile import TemporaryDirectory
    import os
    import json
    from datetime import datetime

    logging.info(f"🎬 Starting video processing: {video_path}")
    
    with TemporaryDirectory() as workdir:
        try:
            # Check input file
            if not os.path.exists(video_path):
                raise Exception(f"Input video file not found: {video_path}")
            
            input_size = os.path.getsize(video_path)
            logging.info(f"📁 Input file size: {input_size} bytes")
            
            if input_size == 0:
                raise Exception("Input video file is empty")
            
            input_ext = os.path.splitext(video_path)[1].lower()
            logging.info(f"📹 Processing {input_ext} file: {video_path}")
            
            # ========== PROBE INPUT FILE ==========
            compressed = os.path.join(workdir, "compressed.mp4")
            
            probe_cmd = [
                "ffprobe", "-v", "quiet", "-print_format", "json", 
                "-show_streams", "-show_format", video_path
            ]
            
            try:
                probe_result = subprocess.run(probe_cmd, capture_output=True, text=True, check=True, timeout=30)
                streams_info = json.loads(probe_result.stdout)
                
                has_audio = any(stream.get('codec_type') == 'audio' for stream in streams_info.get('streams', []))
                has_video = any(stream.get('codec_type') == 'video' for stream in streams_info.get('streams', []))
                video_duration = float(streams_info.get('format', {}).get('duration', 0))
                
                video_stream = next((s for s in streams_info.get('streams', []) if s.get('codec_type') == 'video'), {})
                audio_stream = next((s for s in streams_info.get('streams', []) if s.get('codec_type') == 'audio'), {})
                
                logging.info(f"📊 File analysis: has_video={has_video}, has_audio={has_audio}, duration={video_duration:.2f}s")
                
                if video_stream:
                    logging.info(f"📺 Video: {video_stream.get('codec_name', 'unknown')} {video_stream.get('width', 0)}x{video_stream.get('height', 0)}")
                if audio_stream:
                    logging.info(f"🔊 Audio: {audio_stream.get('codec_name', 'unknown')} {audio_stream.get('sample_rate', 0)}Hz")
                
            except Exception as probe_error:
                logging.warning(f"⚠ Failed to probe input file: {probe_error}")
                has_audio = True
                has_video = True
                video_duration = 0
            
            if not has_video:
                raise Exception("Input file does not contain a valid video stream")
            
            if video_duration <= 0:
                logging.warning(f"⚠ Duration detection failed")
                video_duration = 30.0

            # ========== GPU DETECTION ==========
            nvenc_available = False
            try:
                check_nvenc = subprocess.run(
                    ['ffmpeg', '-h', 'encoder=h264_nvenc'],
                    capture_output=True, text=True, timeout=5
                )
                nvenc_available = (check_nvenc.returncode == 0)
                logging.info(f"{'🚀 GPU (NVENC) detected' if nvenc_available else 'ℹ️ GPU not available - using CPU'}")
            except:
                nvenc_available = False
            
            # ========== VIDEO COMPRESSION ==========
            skip_compression = False
            
            if "_final.mp4" in video_path:
                logging.info("✅ Input is already optimized - skipping compression")
                compressed = video_path
                skip_compression = True
            else:
                logging.info(f"🔄 {'Optimizing' if input_ext == '.mp4' else 'Converting'} video...")
                
                if nvenc_available:
                    if has_audio:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-c:v", "h264_nvenc", "-preset", "p1", "-tune", "hq",
                            "-rc", "vbr", "-cq", "23", "-b:v", "5M",
                            "-maxrate", "10M", "-bufsize", "20M",
                            "-c:a", "aac" if input_ext == '.webm' else "copy",
                            "-ar", "44100", "-ac", "2", "-b:a", "192k",
                            "-movflags", "+faststart", "-pix_fmt", "yuv420p",
                            "-avoid_negative_ts", "make_zero", compressed
                        ]
                    else:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-f", "lavfi", "-i", "anullsrc=channel_layout=stereo:sample_rate=44100",
                            "-c:v", "h264_nvenc", "-preset", "p1", "-tune", "hq",
                            "-rc", "vbr", "-cq", "23", "-b:v", "3M",
                            "-maxrate", "5M", "-bufsize", "6M",
                            "-c:a", "aac", "-ar", "44100", "-ac", "2", "-b:a", "64k",
                            "-shortest", "-movflags", "+faststart",
                            "-pix_fmt", "yuv420p", "-avoid_negative_ts", "make_zero",
                            compressed
                        ]
                else:
                    # CPU fallback
                    if has_audio:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-c:v", "libx264", "-preset", "fast", "-crf", "23",
                            "-maxrate", "10M", "-bufsize", "20M",
                            "-c:a", "aac" if input_ext == '.webm' else "copy",
                            "-ar", "44100", "-ac", "2", "-b:a", "192k",
                            "-movflags", "+faststart", "-pix_fmt", "yuv420p",
                            "-avoid_negative_ts", "make_zero", compressed
                        ]
                    else:
                        ffmpeg_cmd = [
                            "ffmpeg", "-y", "-i", video_path,
                            "-f", "lavfi", "-i", "anullsrc=channel_layout=stereo:sample_rate=44100",
                            "-c:v", "libx264", "-preset", "fast", "-crf", "23",
                            "-maxrate", "3M", "-bufsize", "6M",
                            "-c:a", "aac", "-ar", "44100", "-ac", "2", "-b:a", "64k",
                            "-shortest", "-movflags", "+faststart",
                            "-pix_fmt", "yuv420p", "-avoid_negative_ts", "make_zero",
                            compressed
                        ]

            if not skip_compression:
                try:
                    logging.info(f"🔄 Running compression...")
                    subprocess.run(ffmpeg_cmd, check=True, capture_output=True, text=True, timeout=600)
                    logging.info(f"✅ Video compressed using {'GPU' if nvenc_available else 'CPU'}")
                except subprocess.TimeoutExpired:
                    raise Exception("Video compression timed out")
                except subprocess.CalledProcessError as e:
                    logging.error(f"❌ Compression failed: {e.stderr}")
                    raise Exception(f"Video compression failed: {e.stderr}")

            if not os.path.exists(compressed) or os.path.getsize(compressed) == 0:
                raise Exception("Compressed video file is empty")

            compressed_size = os.path.getsize(compressed)
            logging.info(f"✅ Compressed file: {compressed_size} bytes")

            # ========== AUDIO EXTRACTION FOR TRANSCRIPTION (RAW & ACCURATE) ==========
            audio = os.path.join(workdir, "audio.wav")

            audio_extract_cmd = [
                "ffmpeg", "-y",
                "-i", video_path,          # 🔥 IMPORTANT: ORIGINAL VIDEO, NOT COMPRESSED
                "-vn",
                "-ac", "1",
                "-ar", "16000",
                "-c:a", "pcm_s16le",       # 🔥 LOSSLESS AUDIO (BEST FOR ASR)
                audio
            ]

            try:
                subprocess.run(
                    audio_extract_cmd,
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=120
                )
                logging.info(f"✅ Raw audio extracted for transcription: {audio}")
            except Exception as e:
                raise Exception(f"Audio extraction for transcription failed: {e}")

            # ========== TRANSCRIPTION ==========
            transcript_text = ""
            segments = []

            from pydub import AudioSegment

            audio_file = AudioSegment.from_file(audio)
            chunk_ms = 5 * 60 * 1000  # 5 minutes
            offset = 0.0

            for i in range(0, len(audio_file), chunk_ms):
                chunk = audio_file[i:i + chunk_ms]
                chunk_path = os.path.join(workdir, f"chunk_{i}.wav")
                chunk.export(chunk_path, format="wav")

                with open(chunk_path, "rb") as f:
                    result = groq_client.audio.transcriptions.create(
                        model="whisper-large-v3-turbo",
                        file=f,
                        response_format="verbose_json"
                    )

                # ✅ NORMAL CASE
                if hasattr(result, "segments") and result.segments:
                    for seg in result.segments:
                        segments.append({
                            "start": offset + float(seg["start"]),
                            "end": offset + float(seg["end"]),
                            "text": seg["text"].strip()
                        })
                # ✅ FALLBACK
                else:
                    text = (result.text or "[No speech detected]").strip()
                    segments.append({
                        "start": offset,
                        "end": offset + 5,
                        "text": text
                    })

                offset += len(chunk) / 1000
                os.remove(chunk_path)

            # 🔒 HARD GUARANTEE — NEVER EMPTY
            if not segments:
                segments = [{
                    "start": 0.0,
                    "end": max(5.0, video_duration),
                    "text": "No speech detected."
                }]

            transcript_text = " ".join(seg["text"] for seg in segments)

            # ========== CLEAN TRANSCRIPT & SUBTITLE SEGMENTS ==========
            clean_lines = []
            subtitle_segments = []

            last_text = ""
            MIN_DURATION = 0.6  # seconds (filters noise / music bleed)

            for seg in segments:
                start = float(seg["start"])
                end = float(seg["end"])
                duration = end - start
                text = seg["text"].strip()

                # ❌ Skip very short detections (noise / music)
                if duration < MIN_DURATION:
                    continue

                # ✅ Valid subtitle segment
                subtitle_segments.append({
                    "start": start,
                    "end": end,
                    "text": text
                })

                # ❌ Skip empty or duplicate text for transcript
                if not text or text.lower() == last_text.lower():
                    continue

                clean_lines.append(text)
                last_text = text

            # ✅ Final transcript text
            if clean_lines:
                transcript_text = "\n".join(clean_lines)
            else:
                transcript_text = "No clear spoken dialogue detected."


            # ========== TRAINER PERFORMANCE EVALUATION ==========
            trainer_evaluation = {}
            try:
                if transcript_text and len(transcript_text.strip()) > 50:
                    logging.info("📊 Analyzing trainer performance...")
                    trainer_evaluation = analyze_trainer_performance(transcript_text)
                    
                    if trainer_evaluation and not trainer_evaluation.get("error"):
                        logging.info(f"✅ Trainer evaluation completed:")
                        logging.info(f"   Technical Content: {trainer_evaluation.get('technical_content', 0)}%")
                        logging.info(f"   Explanation Clarity: {trainer_evaluation.get('explanation_clarity', 0)}%")
                        logging.info(f"   Friendliness: {trainer_evaluation.get('friendliness', 0)}%")
                        logging.info(f"   Communication: {trainer_evaluation.get('communication', 0)}%")
                    else:
                        logging.warning(f"⚠ Trainer evaluation returned error: {trainer_evaluation.get('error')}")
                else:
                    logging.info("ℹ️ Insufficient content for trainer evaluation")
                    trainer_evaluation = {
                        "technical_content": 0,
                        "explanation_clarity": 0,
                        "friendliness": 0,
                        "communication": 0,
                        "overall_feedback": "Insufficient content for evaluation"
                    }
            except Exception as eval_error:
                logging.error(f"❌ Trainer evaluation failed: {eval_error}")
                trainer_evaluation = {
                    "error": str(eval_error),
                    "technical_content": 0,
                    "explanation_clarity": 0,
                    "friendliness": 0,
                    "communication": 0,
                    "overall_feedback": "Evaluation failed"
                }

            # ========== SUBTITLES GENERATION ==========
            subtitle_urls = {}
            meeting_type = get_meeting_type(meeting_id)

            for lang in ["en", "hi", "te"]:
                translated_segments = []

                for seg in subtitle_segments:
                    text = seg["text"]

                    if lang != "en":
                        try:
                            text = GoogleTranslator(source="en", target=lang).translate(text)
                        except Exception:
                            text = "[Translation failed]"

                    translated_segments.append({
                        "start": seg["start"],
                        "end": seg["end"],
                        "text": text
                    })

                srt_path = os.path.join(workdir, f"subs_{lang}.srt")
                create_srt_from_segments(translated_segments, srt_path)

                if os.path.exists(srt_path) and os.path.getsize(srt_path) > 0:
                    subtitle_urls[lang] = upload_to_aws_s3(
                        srt_path,
                        f"{meeting_id}/{user_id}/subtitles/{lang}.srt"
                    )

            # ========== SUMMARY GENERATION ==========
            summary = "Processing summary..."
            try:
                if transcript_text and len(transcript_text.strip()) > 10:
                    summary = summarize_segment(transcript_text)
                    logging.info(f"✅ Summary generated ({len(summary)} chars)")
                else:
                    summary = "No sufficient content available."
            except Exception as summary_error:
                logging.warning(f"⚠ Summary generation failed: {summary_error}")
                summary = "Summary generation failed."

            # ========== SUMMARY + MIND MAP (END OF DOCUMENT) ==========
            import re
            from docx import Document
            from docx.shared import Inches

            # --- Generate summary text ---
            if transcript_text.strip():
                summary = summarize_segment(transcript_text)
            else:
                summary = "No sufficient content available."

            summary_doc_path = os.path.join(workdir, "summary.docx")
            summary_doc = Document()

            # Title
            summary_doc.add_heading("Meeting Summary", level=1)

            # Summary content
            for line in summary.split("\n"):
                summary_doc.add_paragraph(line)

            # ---- Mind Map extraction (DOT must come from summary) ----
            image_url = None
            dot_match = re.search(r"```dot\s*(.*?)```", summary, re.DOTALL)

            if dot_match:
                dot_code = dot_match.group(1).strip()
                mindmap_png = os.path.join(workdir, "mindmap.png")

                generate_graph(dot_code, mindmap_png[:-4])

                image_url = upload_to_aws_s3(
                    mindmap_png,
                    f"{meeting_id}/{user_id}/mindmap.png"
                )

                # ✅ ALWAYS AT END
                summary_doc.add_page_break()
                summary_doc.add_heading("Mind Map", level=2)
                summary_doc.add_picture(mindmap_png, width=Inches(6))

            summary_doc.save(summary_doc_path)


            # ========== UPLOAD TO S3 ==========
            logging.info("☁ Starting S3 uploads...")

            video_s3_key = build_s3_video_path(meeting_id, user_id, meeting_type)
            video_url = upload_to_aws_s3(compressed, video_s3_key)

            if not video_url:
                raise Exception("Failed to upload video to S3")

            logging.info(f"✅ Video uploaded: {video_url}")

            transcript_url = None
            summary_url = None

            # --- Transcript DOCX ---
            transcript_path = os.path.join(workdir, "transcript.docx")
            save_docx(transcript_text or "No speech detected.", transcript_path, title="Meeting Transcript")

            # 🔧 FIX: bind correct summary path
            summary_path = summary_doc_path

            if os.path.exists(transcript_path) and os.path.getsize(transcript_path) > 0:
                transcript_s3_key = build_s3_document_path(meeting_id, user_id, meeting_type, "transcript")
                transcript_url = upload_to_aws_s3(transcript_path, transcript_s3_key)

            if os.path.exists(summary_path) and os.path.getsize(summary_path) > 0:
                summary_s3_key = build_s3_document_path(meeting_id, user_id, meeting_type, "summary")
                summary_url = upload_to_aws_s3(summary_path, summary_s3_key)

            logging.info("✅ S3 uploads completed")


            # ========== GET PARTICIPANT EMAILS ==========
            visible_to_emails = get_meeting_participants_emails(meeting_id)
            logging.info(f"✅ Recording visible to {len(visible_to_emails)} users")

            # ✅ GET MEETING TYPE BEFORE CREATING DOCUMENT
            meeting_type = get_meeting_type(meeting_id)
            logging.info(f"📋 Meeting type determined: {meeting_type}")

            # ✅ NEW: Get user details (name, email) for the recording
            user_details = get_user_details(user_id)
            logging.info(f"👤 User details for recording: {user_details}")

            # ========== CHECK CUSTOM NAME ==========
            custom_recording_name = None
            try:
                custom_name_doc = collection.find_one({
                    "meeting_id": meeting_id,
                    "pending_name_update": True
                })
                if custom_name_doc:
                    custom_recording_name = custom_name_doc.get("custom_recording_name")
                    logging.info(f"Found custom name: {custom_recording_name}")
            except Exception as e:
                logging.warning(f"Failed to check custom name: {e}")

            # ========== SCHEDULE METADATA ==========
            schedule_meta = {}
            if meeting_type == "ScheduleMeeting":
                schedule_meta = get_schedule_meeting_metadata(meeting_id)

            # Determine filename
            if custom_recording_name:
                display_filename = f"{custom_recording_name}.mp4"
                original_filename = f"{custom_recording_name}.mp4"
            else:
                display_filename = os.path.basename(video_path)
                original_filename = os.path.basename(video_path)

            # ========== SAVE TO MONGODB ==========
            video_document = {
                "meeting_id": meeting_id,
                "user_id": user_id,
                "user_name": user_details.get("user_name", f"User {user_id}"),  # ✅ ADD THIS
                "user_email": user_details.get("user_email", ""),  # ✅ ADD THIS
                "meeting_type": meeting_type,
                "schedule_id": schedule_meta.get("schedule_id"),
                "schedule_title": schedule_meta.get("schedule_title"),
                "schedule_folder": schedule_meta.get("folder_path"),
                "filename": display_filename,
                "original_filename": original_filename,
                "custom_recording_name": custom_recording_name,
                "video_url": video_url,
                "transcript_url": transcript_url,
                "summary_url": summary_url,
                "summary_text": summary,
                "image_url": image_url,
                "subtitles": subtitle_urls,
                "timestamp": datetime.now(),
                "visible_to": visible_to_emails,
                "file_size": compressed_size,
                "duration": video_duration,
                "transcription_available": bool(transcript_url),
                "summary_available": bool(summary_url),
                "processing_status": "completed",
                "subtitle_format": "srt_multi_language",
                "subtitle_languages": list(subtitle_urls.keys()),
                "embedded_subtitles": False,
                "audio_processing_status": "success" if os.path.exists(audio) else "fallback",
                "audio_preserved": has_audio,
                "source_format": "pyav_mp4",
                "smooth_playback": True,
                "file_type": "video/mp4",
                "is_final_video": True,
                "encoder_used": "GPU (NVENC)" if nvenc_available else "CPU (libx264)",
                "gpu_accelerated": nvenc_available,
                "document_format": "docx",
                "transcription_engine": "groq_whisper_large_v3_turbo",
                "trainer_evaluation": trainer_evaluation  # ✅ ADDED
            }
            
            logging.info(f"💾 Saving video data to MongoDB...")
            
            existing_doc = collection.find_one({
                "meeting_id": meeting_id,
                "user_id": user_id,
                "is_final_video": True
            })

            if existing_doc:
                collection.update_one({"_id": existing_doc["_id"]}, {"$set": video_document})
                logging.info(f"✅ Updated existing document")
            else:
                collection.insert_one(video_document)
                logging.info(f"✅ Created new document")

            # Clean up custom name document
            if custom_recording_name:
                try:
                    collection.delete_one({
                        "meeting_id": meeting_id,
                        "pending_name_update": True
                    })
                    logging.info(f"✅ Cleaned up custom name document")
                except Exception as cleanup_error:
                    logging.warning(f"Cleanup failed: {cleanup_error}")

            # ========== SEND NOTIFICATIONS ==========
            logging.info(f"📧 Sending notifications...")
            try:
                notification_count = send_recording_completion_notifications(
                    meeting_id=meeting_id,
                    video_url=video_url,
                    transcript_url=transcript_url,
                    summary_url=summary_url
                )
                logging.info(f"✅ Sent {notification_count} notifications")
            except Exception as notif_error:
                logging.error(f"⚠ Notifications failed: {notif_error}")

            # ========== RETURN SUCCESS ==========
            result_dict = {
                "status": "success",
                "video_url": video_url,
                "transcript_url": transcript_url,
                "summary_url": summary_url,
                "summary_image_url": image_url,
                "subtitle_urls": subtitle_urls,
                "subtitle_languages": list(subtitle_urls.keys()),
                "file_size": compressed_size,
                "meeting_id": meeting_id,
                "user_id": user_id,
                "subtitle_format": "srt_multi_language",
                "authorized_users_count": len(visible_to_emails),
                "encoder_used": "GPU (NVENC)" if nvenc_available else "CPU (libx264)",
                "gpu_accelerated": nvenc_available,
                "document_format": "docx",
                "transcription_engine": "groq",
                "trainer_evaluation": trainer_evaluation,  # ✅ ADDED
                "processing_notes": {
                    "audio_extracted": os.path.exists(audio),
                    "transcription_successful": bool(segments),
                    "subtitles_generated": len(subtitle_urls),
                    "summary_generated": len(summary) > 50,
                    "original_had_audio": has_audio,
                    "audio_preserved": has_audio,
                    "source_format": "pyav_mp4",
                    "smooth_playback_enabled": True,
                    "duration_preserved": video_duration > 0,
                    "only_final_mp4_saved": True,
                    "gpu_acceleration_used": nvenc_available,
                    "trainer_evaluated": bool(trainer_evaluation and not trainer_evaluation.get("error"))
                }
            }
            
            logging.info(f"✅ Processing completed using {'GPU' if nvenc_available else 'CPU'}")
            return result_dict
            
        except Exception as e:
            logging.error(f"❌ Video processing failed: {e}")
            import traceback
            logging.error(f"❌ Traceback: {traceback.format_exc()}")
            
            return {
                "status": "error",
                "error": str(e),
                "error_type": type(e).__name__,
                "video_url": None,
                "transcript_url": None,
                "summary_url": None,
                "subtitle_urls": {},
                "file_size": 0,
                "meeting_id": meeting_id,
                "user_id": user_id,
                "processing_notes": {
                    "failed_at": "video_processing",
                    "input_file_size": input_size if 'input_size' in locals() else 0,
                    "input_extension": input_ext if 'input_ext' in locals() else "unknown"
                }
            }
          
# === 1. GET ALL VIDEOS ===
@require_http_methods(["GET"])
def get_all_videos(request):
    """Get all video documents with pagination and filtering - STRICT ACCESS CONTROL + MEETING TYPE FILTER"""
    try:
        # Query parameters for pagination and filtering
        page = int(request.GET.get('page', 1))
        limit = int(request.GET.get('limit', 10))
        user_id = request.GET.get('user_id')
        email = request.GET.get('email', '')
        meeting_id = request.GET.get('meeting_id')
        meeting_type = request.GET.get('meeting_type')  # ✅ NEW PARAMETER

        # Build query filter - ONLY SHOW FINAL VIDEOS
        query_filter = {"is_final_video": True}
        
        if meeting_id:
            query_filter['meeting_id'] = meeting_id
        
        # ✅ NEW: Filter by meeting type
        if meeting_type:
            if meeting_type in ['CalendarMeeting', 'ScheduleMeeting', 'InstantMeeting']:
                query_filter['meeting_type'] = meeting_type
                logger.info(f"Filtering videos by meeting_type: {meeting_type}")
            else:
                logger.warning(f"Invalid meeting_type parameter: {meeting_type}")

        # Calculate skip value for pagination
        skip = (page - 1) * limit

        # Debug logging
        logger.info(f"📋 Query params: email={email}, user_id={user_id}, meeting_id={meeting_id}, meeting_type={meeting_type}")
        
        # Fetch raw results
        raw_videos = list(collection.find(query_filter).sort("timestamp", -1).skip(skip).limit(limit))
        logger.info(f"📹 Raw videos found in MongoDB: {len(raw_videos)}")

        allowed_videos = []
        for video in raw_videos:
            video_user_id = video.get("user_id")
            visible_to_emails = video.get("visible_to", [])
            video_meeting_id = video.get("meeting_id")
            
            # STRICT ACCESS CONTROL - ONLY 2 WAYS TO ACCESS:
            
            # Way 1: User uploaded/created this video (is the original uploader)
            user_uploaded = False
            if user_id and video_user_id:
                user_uploaded = str(video_user_id) == str(user_id)
            
            # Way 2: User's email is in the visible_to list (was authorized during processing)
            email_authorized = False
            if email and visible_to_emails:
                # Case-insensitive email comparison
                user_email_lower = email.strip().lower()
                visible_emails_lower = [e.strip().lower() for e in visible_to_emails if e]
                email_authorized = user_email_lower in visible_emails_lower
            
            # ALLOW ACCESS ONLY if user uploaded OR email is in visible_to list
            has_access = user_uploaded or email_authorized
            
            # Debug logging for each video (detailed)
            logger.debug(f"🎬 Video {video.get('_id')} (Meeting: {video_meeting_id}, Type: {video.get('meeting_type')}):")
            logger.debug(f"   - Video uploader: {video_user_id}, Current user: {user_id}, Match: {user_uploaded}")
            logger.debug(f"   - Visible to: {visible_to_emails}")
            logger.debug(f"   - User email: {email}, Match: {email_authorized}")
            logger.debug(f"   - Final access decision: {has_access}")
            
            if has_access:
                # Process video data
                video['_id'] = str(video['_id'])
                video['timestamp'] = video['timestamp'].isoformat() if video.get('timestamp') else None
                
                # ✅ Ensure meeting_type is present (backfill if missing)
                if not video.get('meeting_type'):
                    video['meeting_type'] = get_meeting_type(video_meeting_id)
                    # Update MongoDB with the meeting type
                    try:
                        collection.update_one(
                            {"_id": ObjectId(video['_id'])},
                            {"$set": {"meeting_type": video['meeting_type']}}
                        )
                        logger.info(f"Backfilled meeting_type for video {video['_id']}: {video['meeting_type']}")
                    except Exception as backfill_error:
                        logger.warning(f"Failed to backfill meeting_type: {backfill_error}")
                
                # ✅ Set display name (prioritize custom name over filename)
                if video.get('custom_recording_name'):
                    video['display_name'] = video['custom_recording_name']
                elif video.get('display_name'):
                    video['display_name'] = video['display_name']
                else:
                    # Fallback: clean up filename for display
                    filename = video.get('filename', 'Unnamed Recording')
                    # Remove technical prefixes like "raw_video_" and file extensions
                    if filename.startswith('raw_video_'):
                        filename = filename.replace('raw_video_', '').replace('_final.mp4', '').replace('.mp4', '')
                    video['display_name'] = filename if filename else 'Unnamed Recording'
                
                # Add access reason for debugging
                if user_uploaded:
                    video['access_reason'] = 'uploader'
                    logger.info(f"✅ Video {video.get('_id')}: Access granted (uploader)")
                elif email_authorized:
                    video['access_reason'] = 'authorized_email'
                    logger.info(f"✅ Video {video.get('_id')}: Access granted (authorized email)")
                
                allowed_videos.append(video)
            else:
                logger.debug(f"❌ Video {video.get('_id')}: Access denied for user {user_id}/{email}")

        logger.info(f"✅ Videos after access filtering: {len(allowed_videos)}/{len(raw_videos)}")

        return JsonResponse({
            "status": "success",
            "data": allowed_videos,
            "videos": allowed_videos,  # Alternative key for compatibility
            "pagination": {
                "page": page,
                "limit": limit,
                "total": len(allowed_videos),
                "total_pages": (len(allowed_videos) + limit - 1) // limit
            },
            "debug": {
                "raw_count": len(raw_videos),
                "filtered_count": len(allowed_videos),
                "query_params": {
                    "email": email,
                    "user_id": user_id,
                    "meeting_id": meeting_id,
                    "meeting_type": meeting_type  # ✅ NEW
                }
            }
        })

    except Exception as e:
        logger.error(f"[ERROR] Failed to get videos: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 2. GET VIDEO BY ID ===
@require_http_methods(["GET"])
def get_video_by_id(request, id):
    """Get a specific video document by ID with access control."""
    try:
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')

        video = collection.find_one({"_id": ObjectId(id)})
        if not video:
            logger.error(f"Video ID {id} not found")
            return JsonResponse({"Error": "Video not found"}, status=404)

        meeting_id = video.get("meeting_id", "")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "You are not authorized to view this video"}, status=403)

        video['_id'] = str(video['_id'])
        video['timestamp'] = video['timestamp'].isoformat() if video.get('timestamp') else None

        return JsonResponse({
            "status": "success",
            "data": video
        })

    except Exception as e:
        logger.error(f"[ERROR] Failed to get video by ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 3. UPDATE VIDEO (PUT) ===
@require_http_methods(["PUT"])
@csrf_exempt
def update_video(request, id):
    """Update a video document."""
    try:
        data = json.loads(request.body)
        
        update_data = {}
        allowed_fields = ['meeting_id', 'user_id', 'video_url', 'transcript_url','summary_text', 'summary_url', 'image_url', 'subtitles']
        
        for field in allowed_fields:
            if field in data:
                update_data[field] = data[field]
        
        if not update_data:
            return JsonResponse({"Error": "No valid fields to update"}, status=400)
        
        result = collection.update_one(
            {"_id": ObjectId(id)},
            {"$set": update_data}
        )
        
        if result.matched_count == 0:
            logger.error(f"Video ID {id} not found for update")
            return JsonResponse({"Error": "Video not found"}, status=404)
        
        updated_video = collection.find_one({"_id": ObjectId(id)})
        updated_video['_id'] = str(updated_video['_id'])
        updated_video['timestamp'] = updated_video['timestamp'].isoformat() if updated_video.get('timestamp') else None
        
        logger.info(f"Updated video ID {id}")
        return JsonResponse({
            "Message": "Video updated successfully",
            "data": updated_video
        })
    except Exception as e:
        logger.error(f"[ERROR] Failed to update video ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 4. DELETE VIDEO ===
@require_http_methods(["DELETE"])
@csrf_exempt
def delete_video(request, id):
    """Delete a video document and associated S3 files (host-only permission)."""
    try:
        # Validate ObjectId format first
        try:
            if len(id) != 24:  # MongoDB ObjectId is always 24 characters
                return JsonResponse({"Error": "Invalid video ID format"}, status=400)
            video_id = ObjectId(id)
        except Exception as e:
            logger.error(f"Invalid video ID format: {id}")
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)

        video = collection.find_one({"_id": video_id})
        if not video:
            logger.error(f"Video ID {id} not found")
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Check if user is host
        user_id = request.GET.get("user_id")
        if not user_id:
            return JsonResponse({"Error": "Missing user_id for authorization"}, status=400)

        meeting_id = video.get("meeting_id")
        if not meeting_id:
            return JsonResponse({"Error": "Invalid video metadata (missing meeting_id)"}, status=400)

        # FIXED: More flexible meeting host check
        is_host = False
        
        # First check if user uploaded this video (direct ownership)
        if str(video.get("user_id", "")) == str(user_id):
            is_host = True
            logger.info(f"User {user_id} is the uploader of video {id}")
        else:
            # Then check if user is host in tbl_Meetings
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                    row = cursor.fetchone()
                    if row:
                        host_id = str(row[0])
                        if host_id == str(user_id):
                            is_host = True
                            logger.info(f"User {user_id} is the host of meeting {meeting_id}")
                    else:
                        # Meeting not found in tbl_Meetings, check if user uploaded the video
                        logger.warning(f"Meeting {meeting_id} not found in tbl_Meetings table")
                        if str(video.get("user_id", "")) == str(user_id):
                            is_host = True
                            logger.info(f"Meeting not in DB but user {user_id} is video uploader - allowing delete")
            except Exception as db_error:
                logger.error(f"Database error checking host: {db_error}")
                # Fallback: allow if user uploaded the video
                if str(video.get("user_id", "")) == str(user_id):
                    is_host = True
                    logger.info(f"DB error but user {user_id} is video uploader - allowing delete")

        if not is_host:
            logger.warning(f"User {user_id} is not authorized to delete video {id}")
            return JsonResponse({"Error": "Only the meeting host or video uploader can delete this recording"}, status=403)

        # Delete associated S3 files
        s3_keys_to_delete = []
        
        # Video, transcript, summary, image URLs
        for url_field in ['video_url', 'transcript_url', 'summary_url', 'image_url']:
            url = video.get(url_field)
            if url:
                try:
                    # Extract S3 key from URL more reliably
                    if AWS_S3_BUCKET in url:
                        # Format: https://bucket.s3.region.amazonaws.com/folder/file
                        s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    else:
                        # Fallback: take last two parts
                        url_parts = url.split('/')
                        if len(url_parts) >= 2:
                            s3_key = '/'.join(url_parts[-2:])
                        else:
                            continue
                    s3_keys_to_delete.append(s3_key)
                except Exception as e:
                    logger.warning(f"Failed to extract S3 key from {url}: {e}")
        
        # Subtitle URLs
        subtitles = video.get('subtitles', {})
        for lang, url in subtitles.items():
            if url:
                try:
                    if AWS_S3_BUCKET in url:
                        s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    else:
                        url_parts = url.split('/')
                        if len(url_parts) >= 2:
                            s3_key = '/'.join(url_parts[-2:])
                        else:
                            continue
                    s3_keys_to_delete.append(s3_key)
                except Exception as e:
                    logger.warning(f"Failed to extract S3 key from subtitle {url}: {e}")

        # Delete files from S3 (track success/failure)
        deleted_count = 0
        for s3_key in s3_keys_to_delete:
            try:
                if delete_from_s3(s3_key):
                    deleted_count += 1
            except Exception as e:
                logger.warning(f"Failed to delete S3 file {s3_key}: {e}")

        # Delete video document from MongoDB
        delete_result = collection.delete_one({"_id": video_id})
        if delete_result.deleted_count == 0:
            return JsonResponse({"Error": "Failed to delete video document"}, status=500)
            
        logger.info(f"Successfully deleted video document ID {id}, removed {deleted_count}/{len(s3_keys_to_delete)} S3 files")
        return JsonResponse({
            "Message": "Video deleted successfully",
            "deleted_s3_files": deleted_count,
            "total_s3_files": len(s3_keys_to_delete)
        }, status=200)

    except Exception as e:
        logger.error(f"[ERROR] Failed to delete video ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 5. STREAM VIDEO ===
@require_http_methods(["GET", "HEAD", "OPTIONS"])
@csrf_exempt
def stream_video(request, id):
    """Optimized video streaming with automatic URL repair and better error handling."""
    
    # Handle CORS preflight
    if request.method == 'OPTIONS':
        response = HttpResponse()
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Range, Content-Type, Accept, Authorization'
        response['Access-Control-Expose-Headers'] = 'Content-Range, Accept-Ranges, Content-Length'
        response['Access-Control-Max-Age'] = '86400'
        return response
    
    try:
        # Find video document
        try:
            video = collection.find_one({"_id": ObjectId(id)})
        except Exception as id_error:
            logger.error(f"Invalid video ID format: {id}, Error: {id_error}")
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)
            
        if not video:
            logger.warning(f"Video ID {id} not found in MongoDB")
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Access control check
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        meeting_id = video.get("meeting_id", id)
        
        try:
            access_allowed = is_user_allowed(meeting_id, email=email, user_id=user_id)
            if not access_allowed:
                access_allowed = is_user_allowed_debug(meeting_id, email=email, user_id=user_id)
                
            if not access_allowed:
                logger.warning(f"Access denied for user {user_id}/{email} to video {id}")
                return JsonResponse({"Error": "Access denied"}, status=403)
        except Exception as access_error:
            logger.warning(f"Access control check failed: {access_error}")

        # Get video URL and verify/repair if needed
        video_url = video.get("video_url")
        if not video_url:
            logger.error(f"Video URL not found in MongoDB for ID {id}")
            return JsonResponse({"Error": "Video URL not found"}, status=404)

        # Verify URL and attempt repair if broken
        logger.info(f"Verifying video URL for streaming: {id}")
        video = verify_and_repair_video_url(video)
        
        video_url = video.get("video_url")
        if not video_url:
            logger.error(f"Video URL still missing after repair attempt for ID {id}")
            return JsonResponse({"Error": "Video not accessible"}, status=404)

        # Extract S3 key with improved method
        s3_key = extract_s3_key_from_url(video_url, AWS_S3_BUCKET)
        if not s3_key:
            logger.error(f"Failed to extract S3 key from URL: {video_url}")
            return JsonResponse({"Error": "Invalid video URL format"}, status=400)

        logger.info(f"S3 Key extracted: {s3_key}")

        # Check file size
        file_size = get_s3_object_size(s3_key)
        if file_size <= 0:
            logger.error(f"Video file not found or empty in S3: {s3_key}")
            return JsonResponse({"Error": "Video file not accessible in S3"}, status=404)

        logger.info(f"Video file found in S3: {s3_key} ({file_size} bytes)")

        # Determine content type
        file_ext = os.path.splitext(s3_key)[1].lower()
        content_type_map = {
            '.mp4': 'video/mp4',
            '.avi': 'video/x-msvideo',
            '.mov': 'video/quicktime',
            '.wmv': 'video/x-ms-wmv',
            '.webm': 'video/webm',
            '.mkv': 'video/x-matroska'
        }
        content_type = content_type_map.get(file_ext, 'video/mp4')
        logger.info(f"Content type: {content_type}")

        # Handle HEAD requests
        if request.method == 'HEAD':
            response = HttpResponse(content_type=content_type)
            response['Accept-Ranges'] = 'bytes'
            response['Content-Length'] = str(file_size)
            response['Cache-Control'] = 'public, max-age=3600'
            response['Access-Control-Allow-Origin'] = '*'
            logger.info(f"HEAD request served for video {id}")
            return response

        # Handle range requests
        range_header = request.META.get('HTTP_RANGE')
        if range_header:
            import re
            range_match = re.match(r'bytes=(\d+)-(\d*)', range_header)
            if range_match:
                start = int(range_match.group(1))
                end = int(range_match.group(2)) if range_match.group(2) else file_size - 1

                # Validate range
                if start >= file_size:
                    logger.warning(f"Range start {start} exceeds file size {file_size}")
                    response = HttpResponse(status=416)
                    response['Content-Range'] = f'bytes */{file_size}'
                    response['Access-Control-Allow-Origin'] = '*'
                    return response
                    
                if end >= file_size:
                    end = file_size - 1
                    
                if start > end:
                    response = HttpResponse(status=416)
                    response['Content-Range'] = f'bytes */{file_size}'
                    response['Access-Control-Allow-Origin'] = '*'
                    return response

                range_size = end - start + 1
                logger.info(f"Range request: {start}-{end} ({range_size} bytes)")

                # For large ranges, use streaming response
                if range_size > 5 * 1024 * 1024:  # 5MB+
                    def stream_large_range():
                        chunk_size = 2 * 1024 * 1024  # 2MB chunks
                        current = start
                        chunks_sent = 0
                        while current <= end:
                            try:
                                chunk_end = min(current + chunk_size - 1, end)
                                chunk = stream_from_s3(s3_key, current, chunk_end)
                                if chunk:
                                    yield chunk
                                    chunks_sent += 1
                                    current = chunk_end + 1
                                else:
                                    logger.warning(f"Empty chunk received at byte {current}")
                                    break
                            except Exception as chunk_error:
                                logger.error(f"Error streaming chunk: {chunk_error}")
                                break
                        logger.info(f"Streamed {chunks_sent} chunks for range request")
                    
                    response = StreamingHttpResponse(stream_large_range(), status=206, content_type=content_type)
                else:
                    # Small ranges - get all at once
                    content = stream_from_s3(s3_key, start, end)
                    if content is None:
                        logger.error(f"Failed to stream range {start}-{end}")
                        return JsonResponse({"Error": "Failed to stream video range"}, status=500)
                    response = HttpResponse(content, status=206, content_type=content_type)

                response['Content-Range'] = f'bytes {start}-{end}/{file_size}'
                response['Accept-Ranges'] = 'bytes'
                response['Content-Length'] = str(range_size)
                response['Cache-Control'] = 'public, max-age=3600'
                response['Access-Control-Allow-Origin'] = '*'
                response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
                response['Access-Control-Allow-Headers'] = 'Range, Content-Type, Accept'
                response['Access-Control-Expose-Headers'] = 'Content-Range, Accept-Ranges, Content-Length'
                
                return response

        # For full file requests, use streaming response
        def stream_full_file():
            chunk_size = 2 * 1024 * 1024  # 2MB chunks for smooth streaming
            start = 0
            chunks_sent = 0
            while start < file_size:
                try:
                    end = min(start + chunk_size - 1, file_size - 1)
                    chunk = stream_from_s3(s3_key, start, end)
                    if chunk:
                        yield chunk
                        chunks_sent += 1
                        start = end + 1
                    else:
                        logger.warning(f"Empty chunk received at byte {start}")
                        break
                except Exception as chunk_error:
                    logger.error(f"Error streaming full file chunk: {chunk_error}")
                    break
            logger.info(f"Streamed {chunks_sent} chunks for full file request")

        response = StreamingHttpResponse(stream_full_file(), content_type=content_type)
        response['Accept-Ranges'] = 'bytes'
        response['Content-Length'] = str(file_size)
        response['Cache-Control'] = 'public, max-age=3600'
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Range, Content-Type, Accept'
        response['Access-Control-Expose-Headers'] = 'Content-Range, Accept-Ranges, Content-Length'

        logger.info(f"Streaming full video file {id} ({file_size} bytes)")
        return response
        
    except Exception as e:
        logger.error(f"Stream error for video {id}: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": "Internal server error"}, status=500)

# === 6. HANDLE DOCUMENTS ===
# === 6. HANDLE DOCUMENTS (FIXED FOR DOCX SUPPORT) ===
@require_http_methods(["GET"])
def handle_document(request, id, doc_type):
    """Handle document requests - SUPPORTS BOTH PDF AND DOCX"""
    try:
        logger.info(f"Document request - ID: {id}, Type: {doc_type}, Action: {request.GET.get('action', 'view')}")

        action = request.GET.get('action', 'view').lower()
        if action not in ['view', 'download']:
            logger.error(f"Invalid action: {action}")
            return JsonResponse({"Error": "Invalid action. Use 'view' or 'download'"}, status=400)

        try:
            video_id = ObjectId(id)
        except Exception as e:
            logger.error(f"Invalid video ID format: {id}, Error: {e}")
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)

        video = collection.find_one({"_id": video_id})
        if not video:
            logger.error(f"Video ID {id} not found in MongoDB")
            return JsonResponse({"Error": f"Video ID {id} not found"}, status=404)

        # Enhanced access control check
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        meeting_id = video.get("meeting_id", "")
        
        if not is_user_allowed(meeting_id, email, user_id):
            logger.warning(f"Access denied for user {user_id} to document {doc_type} of video {id}")
            return JsonResponse({"Error": "Access denied: You are not authorized to view this document"}, status=403)

        if doc_type not in ["transcript", "summary"]:
            logger.error(f"Invalid document type: {doc_type}")
            return JsonResponse({"Error": "Invalid document type. Use 'transcript' or 'summary'"}, status=400)

        # Get document URL based on type
        if doc_type == "transcript":
            doc_url = video.get("transcript_url")
            doc_name = "Transcript"
        else:  # summary
            doc_url = video.get("summary_url")
            doc_name = "Summary"
            
        if not doc_url:
            logger.error(f"No {doc_type}_url found in video document for ID {id}")
            return JsonResponse({"Error": f"No {doc_name.lower()} document found for ID {id}"}, status=404)

        # Extract S3 key with improved method
        s3_key = extract_s3_key_from_url(doc_url, AWS_S3_BUCKET)
        if not s3_key:
            logger.error(f"Failed to extract S3 key from document URL: {doc_url}")
            return JsonResponse({"Error": f"Invalid {doc_name.lower()} URL format"}, status=400)

        logger.info(f"S3 Key for document: {s3_key}")

        # Check if document exists
        doc_size = get_s3_object_size(s3_key)
        if doc_size <= 0:
            logger.error(f"Document file not found in S3: {s3_key}")
            return JsonResponse({"Error": f"{doc_name} document not accessible in S3"}, status=404)

        # Download content
        try:
            content = stream_from_s3(s3_key)
            if content is None:
                raise Exception("Failed to download from S3")
        except Exception as e:
            logger.error(f"Failed to download {doc_type} document: {e}")
            return JsonResponse({"Error": f"Failed to access {doc_name.lower()} document"}, status=500)

        # ✅ CRITICAL FIX: Auto-detect file format from S3 key
        file_extension = os.path.splitext(s3_key)[1].lower()
        
        if file_extension == '.docx':
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            file_ext_display = 'docx'
            logger.info(f"Document format: DOCX")
        elif file_extension == '.pdf':
            content_type = 'application/pdf'
            file_ext_display = 'pdf'
            logger.info(f"Document format: PDF")
        else:
            # Fallback: check video metadata for document_format field
            doc_format = video.get('document_format', 'pdf')
            if doc_format == 'docx':
                content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                file_ext_display = 'docx'
                logger.info(f"Document format from metadata: DOCX")
            else:
                content_type = 'application/pdf'
                file_ext_display = 'pdf'
                logger.info(f"Document format from metadata: PDF")

        # Determine response based on action
        if action == 'view':
            response = HttpResponse(content, content_type=content_type)
            response['Content-Disposition'] = f'inline; filename="{doc_type}_{id}.{file_ext_display}"'
            log_action = f"Displayed {doc_name} document for video ID {id} ({file_ext_display})"
        else:  # download
            response = HttpResponse(content, content_type=content_type)
            response['Content-Disposition'] = f'attachment; filename="{doc_type}_{video.get("meeting_name", id)}.{file_ext_display}"'
            log_action = f"Downloaded {doc_name} document for video ID {id} ({file_ext_display})"

        # Set response headers
        response['Content-Length'] = str(len(content))
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Accept'
        response['Access-Control-Expose-Headers'] = 'Content-Length, Content-Disposition'
        response['Cache-Control'] = 'public, max-age=3600'

        logger.info(log_action)
        return response

    except Exception as e:
        logger.error(f"[ERROR] Failed to handle document for video ID {id}: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)
    
# === 7. VIEW MINDMAP ===
@require_http_methods(["GET"])
def view_mindmap(request, id):
    """Serve the mind map image with improved S3 handling."""
    try:
        video = collection.find_one({"_id": ObjectId(id)})
        if not video:
            logger.warning(f"Video document not found for ID {id}")
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Access control check
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        if not is_user_allowed(video.get("meeting_id", ""), email, user_id):
            logger.warning(f"Access denied for user {user_id} to mindmap of video {id}")
            return JsonResponse({"Error": "Access denied: You are not authorized to view this image"}, status=403)

        image_url = video.get("image_url")
        if not image_url:
            logger.error(f"Mind map image not found for ID {id}")
            return JsonResponse({"Error": "Mind map image not found"}, status=404)

        # Extract S3 key with improved method
        s3_key = extract_s3_key_from_url(image_url, AWS_S3_BUCKET)
        if not s3_key:
            logger.error(f"Failed to extract S3 key from mindmap URL: {image_url}")
            return JsonResponse({"Error": "Invalid mindmap URL format"}, status=400)

        logger.info(f"S3 Key for mindmap: {s3_key}")

        # Check if file exists
        file_size = get_s3_object_size(s3_key)
        if file_size <= 0:
            logger.error(f"Mindmap file not found in S3: {s3_key}")
            return JsonResponse({"Error": "Mindmap image not accessible"}, status=404)

        # Download content
        try:
            content = stream_from_s3(s3_key)
            if content is None:
                raise Exception("Failed to download from S3")
        except Exception as e:
            logger.error(f"Failed to download mindmap image: {e}")
            return JsonResponse({"Error": "Failed to access mindmap image"}, status=500)

        response = HttpResponse(content, content_type='image/png')
        response['Content-Length'] = str(len(content))
        response['Cache-Control'] = 'public, max-age=3600'
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'GET, HEAD, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Accept'

        logger.info(f"Served mind map image for video ID {id}")
        return response
        
    except Exception as e:
        logger.error(f"[ERROR] Mind map serving error for ID {id}: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return HttpResponse("Internal server error", status=500)

# === 8. NEW: GET SUBTITLES ===
@require_http_methods(["GET"])
def get_subtitles(request, id, lang):
    """Serve subtitle file with improved S3 handling."""
    try:
        video = collection.find_one({"_id": ObjectId(id)})
        if not video:
            logger.warning(f"Video ID {id} not found")
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Access control check
        email = request.GET.get('email', '')
        user_id = request.GET.get('user_id', '')
        if not is_user_allowed(video.get("meeting_id", ""), email, user_id):
            logger.warning(f"Access denied for user {user_id} to subtitles of video {id}")
            return JsonResponse({"Error": "Access denied"}, status=403)

        subtitles = video.get("subtitles", {})
        if lang not in subtitles:
            logger.error(f"Subtitles not available for language: {lang}")
            return JsonResponse({"Error": f"Subtitles not available for language: {lang}"}, status=404)

        subtitle_url = subtitles[lang]
        if not subtitle_url:
            logger.error(f"Subtitle URL is empty for language {lang}")
            return JsonResponse({"Error": "Subtitle URL not found"}, status=404)

        # Extract S3 key with improved method
        s3_key = extract_s3_key_from_url(subtitle_url, AWS_S3_BUCKET)
        if not s3_key:
            logger.error(f"Failed to extract S3 key from subtitle URL: {subtitle_url}")
            return JsonResponse({"Error": "Invalid subtitle URL format"}, status=400)

        logger.info(f"S3 Key for subtitles ({lang}): {s3_key}")

        # Check if file exists
        file_size = get_s3_object_size(s3_key)
        if file_size <= 0:
            logger.error(f"Subtitle file not found in S3: {s3_key}")
            return JsonResponse({"Error": "Subtitle file not accessible"}, status=404)

        # Download content
        try:
            content = stream_from_s3(s3_key)
            if content is None:
                raise Exception("Failed to download from S3")
        except Exception as e:
            logger.error(f"Failed to download subtitle file: {e}")
            return JsonResponse({"Error": "Failed to access subtitle file"}, status=500)

        response = HttpResponse(content, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="subtitles_{lang}_{id}.srt"'
        response['Content-Length'] = str(len(content))
        response['Access-Control-Allow-Origin'] = '*'
        response['Cache-Control'] = 'public, max-age=3600'

        logger.info(f"Served subtitle file for video {id} in language {lang}")
        return response

    except Exception as e:
        logger.error(f"[ERROR] Subtitle serving error: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 9. UPLOAD RECORDING API ===
@require_http_methods(["POST"])
@csrf_exempt
def upload_recording(request):
    """Upload recording file and process it with enhanced features."""
    try:
        # Get form data
        meeting_id = request.POST.get('meeting_id')
        user_id = request.POST.get('user_id')
        
        # Get uploaded file
        if 'recording_file' not in request.FILES:
            return JsonResponse({"Error": "No recording file uploaded"}, status=400)
        
        uploaded_file = request.FILES['recording_file']
        
        if not meeting_id or not user_id:
            return JsonResponse({"Error": "Missing meeting_id or user_id"}, status=400)
        
        # Check if already processed
        existing = collection.find_one({
            "meeting_id": meeting_id, 
            "user_id": user_id, 
            "filename": uploaded_file.name
        })
        
        if existing:
            return JsonResponse({
                "status": "already_processed",
                "file": uploaded_file.name,
                "video_url": existing.get("video_url"),
                "transcript_url": existing.get("transcript_url"),
                "summary_url": existing.get("summary_url"),
                "summary_image_url": existing.get("image_url"),
                "subtitle_urls": existing.get("subtitles"),
                "message": "This video has already been processed."
            })
        
        # Validate file type
        allowed_extensions = ['.mp4', '.webm', '.mkv', '.avi']
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        if file_extension not in allowed_extensions:
            return JsonResponse({"Error": "Invalid file type. Only video files are allowed."}, status=400)
        
        # Save uploaded file temporarily and process
        with TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            
            with open(temp_file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            
            # Process the recording
            result = process_video_sync(temp_file_path, meeting_id, user_id)
            result["file"] = uploaded_file.name
            result["meeting_id"] = meeting_id
            
            return JsonResponse(result)
                
    except Exception as e:
        logger.error(f"[ERROR] Recording upload failed: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 10. UPLOAD RECORDING BLOB ===
@require_http_methods(["POST"])
@csrf_exempt
def upload_recording_blob(request):
    """Upload recording blob and process it with enhanced error handling."""
    try:
        # Debug request data
        logging.info(f"📥 Recording upload request received")
        logging.info(f"📄 Request FILES: {list(request.FILES.keys())}")
        logging.info(f"📄 Request POST: {dict(request.POST)}")
        
        # Validate required parameters
        if 'recording_blob' not in request.FILES:
            logging.error("❌ No recording_blob in request.FILES")
            return JsonResponse({
                "status": "error",
                "error": "No recording blob provided",
                "Error": "No recording blob provided"
            }, status=400)
        
        recording_file = request.FILES['recording_blob']
        meeting_id = request.POST.get('meeting_id')
        user_id = request.POST.get('user_id')
        recording_id = request.POST.get('recording_id')
        
        logging.info(f"📹 Processing upload: meeting_id={meeting_id}, user_id={user_id}, file_size={recording_file.size}")
        
        if not all([meeting_id, user_id]):
            return JsonResponse({
                "status": "error",
                "error": "Missing required parameters: meeting_id and user_id",
                "Error": "Missing required parameters"
            }, status=400)
        
        if recording_file.size == 0:
            return JsonResponse({
                "status": "error", 
                "error": "Recording file is empty",
                "Error": "Recording file is empty"
            }, status=400)
        
        # Update recording status to "uploading"
        try:
            if recording_id:
                collection.update_one(
                    {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                    {"$set": {
                        "recording_status": "uploading",
                        "upload_start_time": datetime.now(),
                        "file_size": recording_file.size
                    }}
                )
        except Exception as update_error:
            logging.warning(f"⚠ Failed to update recording status: {update_error}")
        
        # Process the recording
        with TemporaryDirectory() as temp_dir:
            # Determine file extension
            file_ext = '.webm'
            if recording_file.content_type:
                if 'mp4' in recording_file.content_type:
                    file_ext = '.mp4'
                elif 'webm' in recording_file.content_type:
                    file_ext = '.webm'
            
            temp_file_path = os.path.join(temp_dir, f"recording_{meeting_id}{file_ext}")
            
            logging.info(f"💾 Saving recording to: {temp_file_path}")
            
            # Save the uploaded file
            with open(temp_file_path, 'wb+') as destination:
                for chunk in recording_file.chunks():
                    destination.write(chunk)
            
            file_size = os.path.getsize(temp_file_path)
            logging.info(f"✅ File saved successfully, size: {file_size} bytes")
            
            if file_size == 0:
                return JsonResponse({
                    "status": "error",
                    "error": "Uploaded file is empty after saving",
                    "Error": "Upload failed - empty file"
                }, status=400)
            
            # Update status to "processing"
            try:
                if recording_id:
                    collection.update_one(
                        {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                        {"$set": {
                            "recording_status": "processing",
                            "processing_start_time": datetime.now()
                        }}
                    )
            except Exception:
                pass
            
            # Process the recording
            logging.info(f"🔄 Starting video processing...")
            result = process_video_sync(temp_file_path, meeting_id, user_id)
            logging.info(f"✅ Video processing completed: {result}")
            
            # Update final recording metadata
            if recording_id and result.get("status") == "success":
                try:
                    final_metadata = {
                        "recording_status": "completed",
                        "processing_completed_time": datetime.now(),
                        "video_url": result.get("video_url"),
                        "transcript_url": result.get("transcript_url"),
                        "summary_url": result.get("summary_url"),
                        "image_url": result.get("summary_image_url"),
                        "subtitles": result.get("subtitle_urls", {}),
                        "final_file_size": file_size,
                        "original_filename": recording_file.name,
                        "processing_notes": result.get("processing_notes", {}),
                        "upload_successful": True
                    }
                    
                    collection.update_one(
                        {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                        {"$set": final_metadata}
                    )
                    
                    logging.info(f"✅ Final metadata updated for recording: {recording_id}")
                    
                except Exception as final_update_error:
                    logging.warning(f"⚠ Failed to update final metadata: {final_update_error}")
            
            # Add additional metadata to result
            result.update({
                "meeting_id": meeting_id,
                "user_id": user_id,
                "original_filename": recording_file.name,
                "file_size": file_size,
                "upload_timestamp": datetime.now().isoformat(),
                "recording_id": recording_id,
                "upload_successful": True,
                "processing_completed": True
            })
            
            logging.info(f"✅ Recording upload and processing completed successfully")
            return JsonResponse(result)
                
    except Exception as e:
        logging.error(f"❌ Recording blob upload failed: {e}")
        import traceback
        logging.error(f"❌ Full traceback: {traceback.format_exc()}")
        
        # Update recording status to "failed" if recording_id available
        try:
            recording_id = request.POST.get('recording_id')
            if recording_id:
                collection.update_one(
                    {"_id": ObjectId(recording_id)} if len(recording_id) == 24 else {"custom_recording_id": recording_id},
                    {"$set": {
                        "recording_status": "failed",
                        "error_message": str(e),
                        "error_timestamp": datetime.now()
                    }}
                )
        except Exception:
            pass
        
        return JsonResponse({
            "status": "error",
            "error": str(e),
            "Error": f"Server error: {str(e)}",
            "upload_successful": False
        }, status=500)
              
# === 11. START RECORDING WITH METADATA ===
@require_http_methods(["POST"])
@csrf_exempt
def start_recording_with_metadata(request, id):
    """Start recording and store initial metadata."""
    try:
        # Check if meeting exists
        video = collection.find_one({"_id": ObjectId(id)})
        if not video:
            # Check if it's a meeting record instead
            video = collection.find_one({"meeting_id": id})
        
        if not video:
            logger.error(f"Meeting ID {id} not found")
            return JsonResponse({"Error": "Meeting not found"}, status=404)

        # Get user from request or use default
        data = json.loads(request.body) if request.body else {}
        user_id = data.get('user_id', video.get('user_id', 'unknown'))

        # Create initial recording metadata
        recording_metadata = {
            "meeting_id": id,
            "user_id": user_id,
            "recording_status": "active",
            "start_time": datetime.now(),
            "video_url": None,
            "transcript_url": None,
            "summary_url": None,
            "image_url": None,
            "subtitles": {},
            "file_size": 0,
            "duration": 0
        }

        # Insert or update recording metadata
        existing_recording = collection.find_one({"meeting_id": id, "recording_status": "active"})
        if existing_recording:
            collection.update_one(
                {"_id": existing_recording["_id"]},
                {"$set": recording_metadata}
            )
            recording_id = str(existing_recording["_id"])
        else:
            result = collection.insert_one(recording_metadata)
            recording_id = str(result.inserted_id)

        logger.info(f"Recording started for meeting {id} with metadata ID {recording_id}")
        
        return JsonResponse({
            "Message": "Recording started with metadata",
            "recording_id": recording_id,
            "meeting_id": id
        }, status=200)
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to start recording with metadata for meeting ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 12. STOP RECORDING AND FINALIZE ===
@require_http_methods(["POST"])
@csrf_exempt
def stop_recording_and_finalize(request, id):
    """Stop recording and prepare for upload."""
    try:
        # Find active recording
        recording = collection.find_one({"meeting_id": id, "recording_status": "active"})
        if not recording:
            logger.error(f"No active recording found for meeting ID {id}")
            return JsonResponse({"Error": "No active recording found"}, status=404)

        # Update recording status
        collection.update_one(
            {"_id": recording["_id"]},
            {"$set": {
                "recording_status": "stopped",
                "end_time": datetime.now()
            }}
        )

        logger.info(f"Recording stopped for meeting {id}")
        
        return JsonResponse({
            "Message": "Recording stopped successfully",
            "recording_id": str(recording["_id"]),
            "meeting_id": id,
            "ready_for_upload": True
        }, status=200)
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to stop recording for meeting ID {id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 13. GET RECORDING STATUS ===
@require_http_methods(["GET"])
def get_recording_status(request, meeting_id):
    """Get recording status for a meeting."""
    try:
        recording = collection.find_one({"meeting_id": meeting_id}, sort=[("start_time", -1)])
        
        if not recording:
            return JsonResponse({"Error": "No recording found for this meeting"}, status=404)
        
        # Convert ObjectId to string
        recording['_id'] = str(recording['_id'])
        if recording.get('start_time'):
            recording['start_time'] = recording['start_time'].isoformat()
        if recording.get('end_time'):
            recording['end_time'] = recording['end_time'].isoformat()
        if recording.get('upload_time'):
            recording['upload_time'] = recording['upload_time'].isoformat()
        
        return JsonResponse({
            "status": "success",
            "recording": recording
        })
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to get recording status for meeting {meeting_id}: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === 14. PROCESS MEETING RECORDING ===
def process_meeting_recording(video_blob, meeting_id, user_id):
    """Process meeting recording blob and upload to S3."""
    try:
        with TemporaryDirectory() as workdir:
            # Save blob to temporary file
            temp_video_path = os.path.join(workdir, f"meeting_recording_{meeting_id}.mp4")
            
            with open(temp_video_path, 'wb') as f:
                f.write(video_blob)
            
            # Use the synchronous process_video function
            result = process_video_sync(temp_video_path, meeting_id, user_id)
            
            logger.info(f"✅ Meeting recording processed successfully for meeting {meeting_id}")
            return result
            
    except Exception as e:
        logger.error(f"❌ Failed to process meeting recording: {e}")
        raise e

# === 15. START RECORDING (LEGACY) ===
@require_http_methods(["POST"])
@csrf_exempt
def Start_Recording(request, id):
    """Start LiveKit stream recording - UPDATED with duplicate prevention"""
    create_meetings_table()

    try:
        # Parse request body for additional settings
        recording_settings = {}
        if request.body:
            try:
                recording_settings = json.loads(request.body)
            except json.JSONDecodeError:
                pass

        with connection.cursor() as cursor:
            select_query = f"""
            SELECT Host_ID, Is_Recording_Enabled, Meeting_Name
            FROM {TBL_MEETINGS}
            WHERE ID = %s
            """
            cursor.execute(select_query, [id])
            row = cursor.fetchone()
            if not row:
                logging.error(f"Meeting ID {id} not found")
                return JsonResponse({"Error": "Meeting not found"}, status=404)

            host_id, is_recording_enabled, meeting_name = row
            
            # Check if recording is already active in DATABASE
            if is_recording_enabled:
                logging.info(f"Recording is already active in database for meeting {id}")
                return JsonResponse({
                    "Message": "Recording is already active",
                    "success": True,
                    "already_recording": True,
                    "is_recording": True,
                    "meeting_id": id,
                    "recording_type": "livekit_stream"
                }, status=200)

            # CLEAN UP ANY OLD FAILED RECORDINGS (older than 2 hours)
            try:
                cutoff_time = datetime.now() - timedelta(hours=2)
                cleanup_result = collection.update_many(
                    {
                        "meeting_id": id,
                        "recording_status": {"$in": ["starting", "active", "uploading"]},
                        "start_time": {"$lt": cutoff_time}
                    },
                    {
                        "$set": {
                            "recording_status": "failed",
                            "error": "Auto-cleaned up - stuck recording",
                            "cleanup_timestamp": datetime.now()
                        }
                    }
                )
                
                if cleanup_result.modified_count > 0:
                    logging.info(f"Auto-cleaned {cleanup_result.modified_count} old recordings for meeting {id}")
                    
            except Exception as cleanup_error:
                logging.warning(f"Recording cleanup failed: {cleanup_error}")

            # Start LiveKit stream recording
            try:
                from core.livekit_recording.recording_service import stream_recording_service
                
                room_name = recording_settings.get('room_name', f"meeting_{id}")
                result = stream_recording_service.start_stream_recording(id, str(host_id), room_name)
                
                if result.get("status") == "success":
                    # Update database to reflect recording state
                    started_at = timezone.now()
                    update_query = f"""
                    UPDATE {TBL_MEETINGS}
                    SET Is_Recording_Enabled = 1, Started_At = %s
                    WHERE ID = %s
                    """
                    cursor.execute(update_query, [started_at, id])
                    
                    logging.info(f"Stream recording started for meeting {id}")
                    
                    return JsonResponse({
                        "Message": "Stream recording started - capturing all participant streams",
                        "success": True,
                        "already_recording": False,
                        "is_recording": True,
                        "meeting_id": id,
                        "recording_id": result.get("recording_id"),
                        "recording_type": "livekit_stream",
                        "screen_share_required": False,
                        "user_interaction_required": False,
                        "bot_joining": True,
                        "captures": "all_video_audio_streams_and_screen_shares",
                        "room_name": room_name,
                        "recorder_identity": result.get("recorder_identity"),
                        "like_google_meet": True,
                        "records_all_participants": True,
                        "settings": recording_settings
                    })
                    
                elif result.get("status") in ["already_active", "already_exists"]:
                    # Recording already exists - sync database
                    started_at = timezone.now()
                    cursor.execute(update_query, [started_at, id])
                    
                    return JsonResponse({
                        "Message": "Recording was already active",
                        "success": True,
                        "already_recording": True,
                        "is_recording": True,
                        "meeting_id": id,
                        "recording_id": result.get("recording_id"),
                        "recording_type": "livekit_stream"
                    })
                    
                else:
                    return JsonResponse({
                        "Error": result.get("message", "Failed to start stream recording"),
                        "success": False,
                        "meeting_id": id,
                        "recording_type": "livekit_stream"
                    }, status=500)
                    
            except Exception as e:
                logging.error(f"Error starting stream recording: {e}")
                return JsonResponse({
                    "Error": f"Stream recording failed: {str(e)}",
                    "success": False,
                    "meeting_id": id
                }, status=500)
            
    except Exception as e:
        logging.error(f"Database error starting recording for meeting {id}: {e}")
        return JsonResponse({
            "Error": f"Database error: {str(e)}",
            "success": False,
            "meeting_id": id
        }, status=500)

# === 16. STOP RECORDING (LEGACY) ===
@require_http_methods(["POST"])
@csrf_exempt
def Stop_Recording(request, id):
    """Stop LiveKit stream recording - UPDATED with processing integration"""
    create_meetings_table()

    try:
        # Get meeting info BEFORE updating
        with connection.cursor() as cursor:
            cursor.execute(f"""
            SELECT Host_ID, Is_Recording_Enabled, Meeting_Name
            FROM {TBL_MEETINGS}
            WHERE ID = %s
            """, [id])
            
            row = cursor.fetchone()
            if not row:
                logging.error(f"Meeting ID {id} not found")
                return JsonResponse({"Error": "Meeting not found"}, status=404)

            host_id, is_recording_enabled, meeting_name = row
            
            if not is_recording_enabled:
                return JsonResponse({
                    "Message": "Recording was not active",
                    "meeting_id": id,
                    "is_recording": False,
                    "success": True,
                    "recording_type": "livekit_stream"
                }, status=200)

        # Stop LiveKit stream recording
        try:
            from core.livekit_recording.recording_service import stream_recording_service
            
            result = stream_recording_service.stop_stream_recording(id)
            
            # Update database to reflect stopped state FIRST
            ended_at = timezone.now()
            with connection.cursor() as cursor:
                cursor.execute(f"""
                UPDATE {TBL_MEETINGS}
                SET Is_Recording_Enabled = 0, Ended_At = %s
                WHERE ID = %s
                """, [ended_at, id])
            
            # Handle processing based on result
            if result and result.get("status") == "success":
                # Check if processing was completed
                processing_result = result.get("processing_result", {})
                
                if processing_result.get("status") == "success":
                    logging.info(f"Stream recording stopped AND PROCESSED for meeting {id}")
                    
                    return JsonResponse({
                        "Message": "Stream recording stopped and processed successfully",
                        "success": True,
                        "meeting_id": id,
                        "meeting_name": meeting_name,
                        "is_recording": False,
                        "recording_type": "livekit_stream",
                        "processing_completed": True,
                        "video_url": processing_result.get("video_url"),
                        "transcript_url": processing_result.get("transcript_url"),
                        "summary_url": processing_result.get("summary_url"),
                        "subtitle_urls": processing_result.get("subtitle_urls", {}),
                        "image_url": processing_result.get("image_url"),
                        "file_size": result.get("file_size", 0),
                        "transcription_available": bool(processing_result.get("transcript_url")),
                        "summary_available": bool(processing_result.get("summary_url")),
                        "streams_captured": "all_participant_streams",
                        "like_google_meet": True,
                        "captured_all_participants": True
                    })
                else:
                    # Recording stopped but processing failed
                    return JsonResponse({
                        "Message": "Recording stopped but processing failed",
                        "success": True,
                        "meeting_id": id,
                        "meeting_name": meeting_name,
                        "is_recording": False,
                        "recording_type": "livekit_stream",
                        "recording_stopped": True,
                        "processing_failed": True,
                        "processing_error": processing_result.get("error"),
                        "file_path": result.get("file_path"),
                        "file_size": result.get("file_size", 0),
                        "suggestion": "Raw file available for manual processing"
                    })
                    
            elif result and result.get("status") == "partial_success":
                return JsonResponse({
                    "Message": "Stream recording stopped but had processing issues",
                    "success": True,
                    "meeting_id": id,
                    "meeting_name": meeting_name,
                    "is_recording": False,
                    "recording_type": "livekit_stream",
                    "recording_stopped": True,
                    "processing_issues": True,
                    "error": result.get("error"),
                    "file_path": result.get("file_path")
                })
            else:
                error_msg = "Stream recording failed to produce valid output"
                if result:
                    error_msg = result.get("message", error_msg)
                
                logging.warning(f"Stream recording failed for meeting {id}: {error_msg}")
                
                return JsonResponse({
                    "Message": "Recording stopped but stream recording failed",
                    "success": True,
                    "meeting_id": id,
                    "meeting_name": meeting_name,
                    "is_recording": False,
                    "recording_type": "livekit_stream",
                    "recording_stopped": True,
                    "stream_recording_failed": True,
                    "error": error_msg,
                    "reason": "Stream recording bot failed to capture meeting content"
                })
                
        except Exception as e:
            logging.error(f"Error stopping stream recording: {e}")
            
            # Still update database to show recording stopped
            try:
                ended_at = timezone.now()
                with connection.cursor() as cursor:
                    cursor.execute(f"""
                    UPDATE {TBL_MEETINGS}
                    SET Is_Recording_Enabled = 0, Ended_At = %s
                    WHERE ID = %s
                    """, [ended_at, id])
            except Exception:
                pass
            
            return JsonResponse({
                "Error": f"Failed to stop stream recording: {str(e)}",
                "success": False,
                "meeting_id": id,
                "recording_type": "livekit_stream"
            }, status=500)

    except Exception as e:
        logging.error(f"Critical failure for meeting {id}: {e}")
        
        return JsonResponse({
            "Error": f"Critical error: {str(e)}", 
            "success": False,
            "meeting_id": id
        }, status=500)
              
# === 17. NEW: UPLOAD SINGLE FILE (FROM FASTAPI) ===
@require_http_methods(["POST"])
@csrf_exempt
def upload_single_file(request):
    """Upload and process a single video file (FastAPI equivalent)."""
    try:
        # Get form data
        meeting_id = request.POST.get('meeting_id')
        user_id = request.POST.get('user_id')
        
        if 'file' not in request.FILES:
            return JsonResponse({"error": "No file uploaded"}, status=400)
        
        uploaded_file = request.FILES['file']
        
        if not meeting_id or not user_id:
            return JsonResponse({"error": "Missing meeting_id or user_id"}, status=400)
        
        # Check if already processed
        existing = collection.find_one({
            "meeting_id": meeting_id, 
            "user_id": user_id, 
            "filename": uploaded_file.name
        })
        
        if existing:
            return JsonResponse({
                "status": "already_processed",
                "file": uploaded_file.name,
                "video_url": existing.get("video_url"),
                "transcript_url": existing.get("transcript_url"),
                "summary_url": existing.get("summary_url"),
                "summary_image_url": existing.get("image_url"),
                "subtitle_urls": existing.get("subtitles"),
                "message": "This video has already been processed."
            })
        
        # Save uploaded file temporarily and process
        with TemporaryDirectory() as temp_dir:
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            
            with open(temp_file_path, 'wb+') as destination:
                for chunk in uploaded_file.chunks():
                    destination.write(chunk)
            
            # Process the recording
            result = process_video_sync(temp_file_path, meeting_id, user_id)
            result["file"] = uploaded_file.name
            result["meeting_id"] = meeting_id
            
            return JsonResponse(result)
                
    except Exception as e:
        logger.error(f"[ERROR] Single file upload failed: {e}")
        return JsonResponse({"error": f"Server error: {str(e)}"}, status=500)

# === 18. NEW: LIST ALL VIDEOS WITH METADATA ===
@require_http_methods(["GET"])
def list_all_videos_detailed(request):
    """List all videos with detailed metadata."""
    try:
        # Query parameters
        page = int(request.GET.get('page', 1))
        limit = int(request.GET.get('limit', 20))
        meeting_id = request.GET.get('meeting_id')
        user_id = request.GET.get('user_id')
        
        # Build query filter
        query_filter = {}
        if meeting_id:
            query_filter['meeting_id'] = meeting_id
        if user_id:
            query_filter['user_id'] = user_id
        
        # Calculate skip value for pagination
        skip = (page - 1) * limit
        
        # Fetch videos
        videos = list(collection.find(query_filter).sort("timestamp", -1).skip(skip).limit(limit))
        total_count = collection.count_documents(query_filter)
        
        # Process videos
        for video in videos:
            video['_id'] = str(video['_id'])
            video['timestamp'] = video['timestamp'].isoformat() if video.get('timestamp') else None
            if video.get('start_time'):
                video['start_time'] = video['start_time'].isoformat()
            if video.get('end_time'):
                video['end_time'] = video['end_time'].isoformat()
            if video.get('upload_time'):
                video['upload_time'] = video['upload_time'].isoformat()
        
        return JsonResponse({
            "status": "success",
            "data": videos,
            "pagination": {
                "page": page,
                "limit": limit,
                "total": total_count,
                "total_pages": (total_count + limit - 1) // limit
            }
        })
        
    except Exception as e:
        logger.error(f"[ERROR] Failed to list videos: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === MEETING TRASH MANAGEMENT FUNCTIONS (USING EXISTING TABLES ONLY) ===
@require_http_methods(["DELETE"])
@csrf_exempt
def move_video_to_trash(request, id):
    """Move a video recording to trash."""
    try:
        # Validate ObjectId format
        try:
            if len(id) != 24:
                return JsonResponse({"Error": "Invalid video ID format"}, status=400)
            video_id = ObjectId(id)
        except Exception:
            return JsonResponse({"Error": "Invalid video ID format"}, status=400)

        # Find video in MongoDB
        video = collection.find_one({"_id": video_id})
        if not video:
            return JsonResponse({"Error": "Video not found"}, status=404)

        # Check if already trashed
        if video.get("is_trashed"):
            return JsonResponse({"Error": "Video is already in trash"}, status=400)

        # Check permissions
        user_id = request.GET.get("user_id")
        email = request.GET.get("email", "")
        
        if not user_id:
            return JsonResponse({"Error": "Missing user_id"}, status=400)

        meeting_id = video.get("meeting_id")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "Permission denied: Only the meeting host can delete this recording"}, status=403)

        # Move to trash
        trash_date = datetime.now()
        permanent_delete_date = trash_date + timedelta(days=TRASH_RETENTION_DAYS)
        
        collection.update_one(
            {"_id": video_id},
            {"$set": {
                "is_trashed": True,
                "trashed_at": trash_date,
                "permanent_delete_at": permanent_delete_date
            }}
        )

        logger.info(f"Video {id} moved to trash by user {user_id}")

        return JsonResponse({
            "Message": "Recording moved to trash successfully",
            "video_id": id,
            "trashed": True,
            "trashed_at": trash_date.isoformat(),
            "permanent_delete_at": permanent_delete_date.isoformat()
        })

    except Exception as e:
        logger.error(f"Failed to move video to trash: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

@require_http_methods(["GET"])
def list_trash_videos(request):
    """List trashed videos for the user - HOST ONLY ACCESS."""
    try:
        user_id = request.GET.get("user_id")
        
        if not user_id:
            return JsonResponse({"Error": "Missing user_id"}, status=400)

        # Get trashed videos where user is the HOST/OWNER only
        trashed_videos = list(collection.find({
            "is_trashed": True,
            "user_id": user_id  # Only show videos uploaded by this user (host)
        }).sort("trashed_at", -1))

        # Format response
        for video in trashed_videos:
            video["_id"] = str(video["_id"])
            for date_field in ["trashed_at", "permanent_delete_at", "timestamp"]:
                if date_field in video and video[date_field]:
                    video[date_field] = video[date_field].isoformat()

        return JsonResponse({
            "videos": trashed_videos,
            "total": len(trashed_videos),
            "page": 1,
            "pages": 1
        })

    except Exception as e:
        logger.error(f"Failed to list trash videos: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

@require_http_methods(["POST"])
@csrf_exempt  
def restore_video(request, id):
    """Restore video from trash."""
    try:
        video_id = ObjectId(id)
        video = collection.find_one({"_id": video_id})
        
        if not video:
            return JsonResponse({"Error": "Video not found"}, status=404)
            
        if not video.get("is_trashed"):
            return JsonResponse({"Error": "Video is not in trash"}, status=400)

        # Check permissions
        user_id = request.GET.get("user_id")
        email = request.GET.get("email", "")
        
        meeting_id = video.get("meeting_id")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "Permission denied"}, status=403)

        # Restore video
        collection.update_one(
            {"_id": video_id},
            {"$unset": {
                "is_trashed": "",
                "trashed_at": "",
                "permanent_delete_at": ""
            }}
        )

        return JsonResponse({
            "Message": "Recording restored successfully",
            "video_id": id,
            "restored": True
        })

    except Exception as e:
        logger.error(f"Failed to restore video: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

@require_http_methods(["DELETE"])
@csrf_exempt
def permanent_delete_video(request, id):
    """Permanently delete a video and its S3 files."""
    try:
        video_id = ObjectId(id)
        video = collection.find_one({"_id": video_id})
        
        if not video:
            return JsonResponse({"Error": "Video not found"}, status=404)
            
        if not video.get("is_trashed"):
            return JsonResponse({"Error": "Video must be in trash first"}, status=400)

        # Check permissions
        user_id = request.GET.get("user_id")
        email = request.GET.get("email", "")
        
        meeting_id = video.get("meeting_id")
        if not is_user_allowed(meeting_id, email=email, user_id=user_id):
            return JsonResponse({"Error": "Permission denied"}, status=403)

        # Delete S3 files
        s3_keys_to_delete = []
        for url_field in ['video_url', 'transcript_url', 'summary_url', 'image_url']:
            url = video.get(url_field)
            if url and AWS_S3_BUCKET in url:
                try:
                    s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    s3_keys_to_delete.append(s3_key)
                except:
                    pass

        # Delete subtitle files
        subtitles = video.get('subtitles', {})
        for lang, url in subtitles.items():
            if url and AWS_S3_BUCKET in url:
                try:
                    s3_key = url.split(f'{AWS_S3_BUCKET}.s3.')[1].split('/', 1)[1]
                    s3_keys_to_delete.append(s3_key)
                except:
                    pass

        # Delete S3 files
        deleted_count = 0
        for s3_key in s3_keys_to_delete:
            if delete_from_s3(s3_key):
                deleted_count += 1

        # Delete from MongoDB
        collection.delete_one({"_id": video_id})

        return JsonResponse({
            "Message": "Recording permanently deleted",
            "video_id": id,
            "deleted_s3_files": deleted_count,
            "total_s3_files": len(s3_keys_to_delete)
        })

    except Exception as e:
        logger.error(f"Failed to permanently delete video: {e}")
        return JsonResponse({"Error": f"Server error: {str(e)}"}, status=500)

# === NEW: STORE CUSTOM RECORDING NAME ===
@require_http_methods(["POST"])
@csrf_exempt
def store_custom_recording_name(request):
    """
    Store custom recording name immediately when user enters it.
    This happens while recording is still processing in background.
    """
    try:
        data = json.loads(request.body)
        meeting_id = data.get('meeting_id')
        custom_name = data.get('custom_name')
        user_id = data.get('user_id')  # For permission checking
        
        if not meeting_id or not custom_name:
            return JsonResponse({
                "status": "error",
                "error": "Missing meeting_id or custom_name"
            }, status=400)
        
        # Validate custom name
        custom_name = custom_name.strip()
        if len(custom_name) == 0:
            return JsonResponse({
                "status": "error",
                "error": "Recording name cannot be empty"
            }, status=400)
        
        if len(custom_name) > 200:
            return JsonResponse({
                "status": "error",
                "error": "Recording name too long (max 200 characters)"
            }, status=400)
        
        # Optional: Check if user has permission for this meeting
        if user_id:
            try:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT Host_ID FROM tbl_Meetings WHERE ID = %s", [meeting_id])
                    row = cursor.fetchone()
                    if row and str(row[0]) != str(user_id):
                        # Not the host, check if participant
                        cursor.execute(
                            "SELECT COUNT(*) FROM tbl_Participants WHERE Meeting_ID = %s AND User_ID = %s",
                            [meeting_id, user_id]
                        )
                        if cursor.fetchone()[0] == 0:
                            return JsonResponse({
                                "status": "error",
                                "error": "You don't have permission to name this recording"
                            }, status=403)
            except Exception as perm_error:
                logger.warning(f"Permission check failed: {perm_error}")
        
        # Store the custom name with a pending flag
        custom_name_document = {
            "meeting_id": meeting_id,
            "custom_recording_name": custom_name,
            "pending_name_update": True,
            "name_stored_at": datetime.now(),
            "user_id": user_id
        }
        
        # Check if already exists (user might change name before processing completes)
        existing = collection.find_one({
            "meeting_id": meeting_id,
            "pending_name_update": True
        })
        
        if existing:
            # Update existing pending name
            collection.update_one(
                {"_id": existing["_id"]},
                {"$set": {
                    "custom_recording_name": custom_name,
                    "name_stored_at": datetime.now()
                }}
            )
            logger.info(f"Updated pending custom name for meeting {meeting_id}")
        else:
            # Insert new pending name
            result = collection.insert_one(custom_name_document)
            logger.info(f"Stored custom name for meeting {meeting_id}: {custom_name}")
        
        return JsonResponse({
            "status": "success",
            "message": "Recording name stored successfully",
            "custom_name": custom_name,
            "meeting_id": meeting_id
        })
        
    except json.JSONDecodeError:
        return JsonResponse({
            "status": "error",
            "error": "Invalid JSON in request body"
        }, status=400)
    except Exception as e:
        logger.error(f"Error storing custom recording name: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return JsonResponse({
            "status": "error",
            "error": f"Server error: {str(e)}"
        }, status=500)

# === URLS CONFIGURATION ===
urlpatterns = [
    # Video CRUD operations (Original APIs converted to AWS)
    path('api/videos/lists', get_all_videos, name='get_all_videos'),
    path('api/videos/list-detailed', list_all_videos_detailed, name='list_all_videos_detailed'),
    path('api/videos/<str:id>', get_video_by_id, name='get_video_by_id'),
    path('api/videos/update/<str:id>', update_video, name='update_video'),
    path('api/videos/remove/<str:id>', delete_video, name='delete_video'),
    path('api/videos/stream/<str:id>', stream_video, name='stream_video'),
    path('api/videos/doc/<str:id>/<str:doc_type>', handle_document, name='handle_document'),
    path('api/videos/<str:id>/mindmap', view_mindmap, name='view_mindmap'),
    
    # NEW: Subtitle support
    path('api/videos/<str:id>/subtitles/<str:lang>', get_subtitles, name='get_subtitles'),
    
    # Meeting recording APIs (Original)
    path('api/meetings/<str:id>/start-recording', Start_Recording, name='Start_Recording'),
    path('api/meetings/<str:id>/stop-recording', Stop_Recording, name='Stop_Recording'),
    
    # Upload APIs (Original + Enhanced)
    path('api/videos/upload-recording', upload_recording, name='upload_recording'),
    path('api/recordings/upload-blob', upload_recording_blob, name='upload_recording_blob'),
    path('api/upload-single/', upload_single_file, name='upload_single_file'),
    
    # Recording management APIs (Original)
    path('api/recordings/start-with-metadata/<str:id>', start_recording_with_metadata, name='start_recording_with_metadata'),
    path('api/recordings/stop-and-finalize/<str:id>', stop_recording_and_finalize, name='stop_recording_and_finalize'),
    path('api/recordings/status/<str:meeting_id>', get_recording_status, name='get_recording_status'),

    # ===== VIDEO TRASH MANAGEMENT URLS (CORRECTED) =====
    path('api/recordings/trash/<str:id>', move_video_to_trash, name='move_video_to_trash'),
    path('api/videos/trash/list', list_trash_videos, name='list_trash_videos'),  # FIXED: was list_trash_meetings
    path('api/videos/restore/<str:id>', restore_video, name='restore_video'),
    path('api/videos/permanent-delete/<str:id>', permanent_delete_video, name='permanent_delete_video'),
    path('api/recordings/store-custom-name', store_custom_recording_name, name='store_custom_recording_name'),
]
